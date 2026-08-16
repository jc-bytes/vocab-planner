#!/usr/bin/env python3
"""Create the revised Grade 7 IIT Summative 5 and rhythm-game exam documents."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.table import _Cell


ROOT = Path(__file__).resolve().parents[1]
RUBRICS = ROOT / "plans" / "7th Grade Technology" / "Rubrics" / "IIT"
DAILY = RUBRICS / "Daily"
EXAMS = RUBRICS / "Exam Projects"
ARCHIVE = RUBRICS / "Archived Variants"

SUM_TEMPLATE = DAILY / "7th grade - IIT - Week 3 - Rubric for Summative 3.docx"
OLD_SUM = DAILY / "7th grade - IIT - Week 7 - Summative 5 - Scratch Debugging Check.docx"
NEW_SUM = DAILY / "7th grade - IIT - Week 8 - Summative 5 - Rhythm Game Code Walkthrough Presentation.docx"
OLD_EXAM = EXAMS / "7th grade - IIT - Final Project Packet and Rubric - Scratch Dance Game.docx"
NEW_EXAM = EXAMS / "7th grade - IIT - Final Project Packet and Rubric - Scratch Rhythm Game.docx"

FONT = "Arial"
SIZE = Pt(12)


def physical_cells(table, row_index: int) -> list[_Cell]:
    return [_Cell(tc, table) for tc in table._tbl.tr_lst[row_index].tc_lst]


def format_run(run, *, bold: bool | None = None) -> None:
    run.font.name = FONT
    run.font.size = SIZE
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, text: str, *, bold: bool = False, centered: bool = False) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    format_run(paragraph.add_run(text), bold=bold)


def set_cell(cell: _Cell, text: str, *, bold: bool = False, centered: bool = False) -> None:
    cell.text = ""
    set_paragraph(cell.paragraphs[0], text, bold=bold, centered=centered)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell: _Cell, width: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.get_or_add_tcW()
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for column, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(column, len(widths) - 1)])


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def prevent_row_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    trpr.append(cant_split)


def create_summative() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    for text, bold in (
        ("Academia Internacional David", True),
        ("Robotics and Technology", True),
        ("2nd Trimester — Summative #5 — Week 8 — 7th A & B", True),
        ("Reference Rhythm Game Code Walkthrough Presentation", True),
    ):
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, text, bold=bold, centered=True)

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [5700, 5700])
    set_cell(meta.cell(0, 0), "Name: ________________________________")
    set_cell(meta.cell(0, 1), "Date: __________________   Group: 7° A B")
    set_cell(meta.cell(1, 0), "Teacher: Porfirio Rios")
    set_cell(meta.cell(1, 1), "Score: _____ / 40 pts")

    task = doc.add_paragraph()
    set_paragraph(
        task,
        "Task: Demonstrate the teacher-provided D/F/J/K rhythm game, explain the assigned component and at least two connected scripts, connect the code to visible behavior, and answer one teacher question.",
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 2325, 2325, 2325, 2325])
    for cell, text in zip(table.rows[0].cells, ("Criteria", "9", "7", "5", "2")):
        set_cell(cell, text, bold=True, centered=True)
        shade_cell(cell, "D9EAD3")
    repeat_table_header(table.rows[0])
    criteria = (
        (
            "Live demonstration and component identification",
            "Game runs; assigned sprites, variables, and purpose are identified accurately.",
            "Game runs and most assigned parts are identified; one detail is unclear.",
            "Demonstration or identification is incomplete; several details are unclear.",
            "Game is not demonstrated or assigned parts are not identified.",
        ),
        (
            "Connected-script explanation",
            "Explains at least 2 connected scripts from event or broadcast to condition, variable change, and visible output.",
            "Explains 2 connected scripts with one missing or unclear connection.",
            "Explains part of 1 or 2 scripts, but the code-to-output connection is limited.",
            "Connected-script logic is missing or mostly incorrect.",
        ),
        (
            "Whole-game connection and vocabulary",
            "Explains how the component affects gameplay and correctly uses at least 4 terms: event, broadcast, target, lane, collision, score, time, tempo.",
            "Explains the gameplay connection and correctly uses 3 terms.",
            "Gives a partial gameplay connection and correctly uses 1 or 2 terms.",
            "Gameplay connection is missing or technical vocabulary is not used accurately.",
        ),
        (
            "Communication and teacher question",
            "Presents for about 2 minutes in understandable English, uses the live project or code map, and answers the teacher question accurately.",
            "Presentation is understandable and the question is answered with one minor omission.",
            "Presentation depends heavily on reading or the answer is incomplete.",
            "Presentation or teacher-question response is missing.",
        ),
    )
    for row_values in criteria:
        added_row = table.add_row()
        prevent_row_split(added_row)
        cells = added_row.cells
        for column, (cell, text) in enumerate(zip(cells, row_values)):
            set_cell(cell, text, bold=column == 0)
    points_row = table.add_row()
    prevent_row_split(points_row)
    for cell, text in zip(points_row.cells, ("Criteria", "4", "3", "2", "0")):
        set_cell(cell, text, bold=True, centered=True)
        shade_cell(cell, "D9EAD3")
    readiness = (
        "Punctuality, Readiness & Respect",
        "Presents on time; project and code map are ready; class time, equipment, and audience expectations are handled responsibly.",
        "Ready and on time with one minor preparation or responsibility issue.",
        "Presentation occurs, but readiness, timing, or respectful participation is inconsistent.",
        "Presentation evidence is not submitted or readiness and responsibility prevent evaluation.",
    )
    readiness_row = table.add_row()
    prevent_row_split(readiness_row)
    for column, (cell, text) in enumerate(zip(readiness_row.cells, readiness)):
        set_cell(cell, text, bold=column == 0)
    comments = table.add_row().cells
    merged = comments[0]
    for cell in comments[1:]:
        merged = merged.merge(cell)
    set_cell(merged, "Comments:\n\n")
    doc.save(NEW_SUM)
    print(f"created {NEW_SUM}")


def replace_paragraph_text(doc: Document, index: int, text: str, *, bold: bool = False) -> None:
    set_paragraph(doc.paragraphs[index], text, bold=bold)


def create_exam() -> None:
    exam_template = OLD_EXAM if OLD_EXAM.exists() else ARCHIVE / OLD_EXAM.name
    if not exam_template.exists():
        raise FileNotFoundError(f"Missing exam template: {exam_template}")
    shutil.copy2(exam_template, NEW_EXAM)
    doc = Document(NEW_EXAM)
    replacements = {
        2: ("Scratch Rhythm Game", True),
        3: ("7th Grade - 2nd Trimester - Final Project Packet and Rubric", True),
        4: ("Project Overview", True),
        5: (
            "Students create an original osu!-style Scratch rhythm game with four D/F/J/K lanes, moving targets, collision-based hit checks, hit/miss feedback, score, time, tempo, music, and a clear end condition.",
            False,
        ),
        6: ("Required Evidence", True),
        7: ("Project plan with stage, four lanes, keys, variables, broadcasts, hit/miss rules, timing, and roles.", False),
        8: ("Code map showing the target generator, D/F/J/K input checks, feedback, variables, and end condition.", False),
        9: ("Playable Scratch project with four target lanes, score, time, miss consequence, tempo, music, and feedback.", False),
        10: ("Test log with at least three tests, one bug, the code change, and the result after retesting.", False),
        11: ("Final reflection and project link or screenshot evidence.", False),
        12: ("Schedule and Evidence Expectations", True),
        13: ("Week 9: read the rubric, complete the code map and project plan, and test one starter lane.", False),
        14: ("Weeks 10-11: build all four lanes, complete the game systems, run three tests, fix one problem, and prepare the explanation.", False),
        15: ("Week 12: demonstrate the final game and submit the project, code map, test log, and reflection.", False),
        16: ("Presentation or Demonstration Requirements", True),
        17: ("Demonstrate the game from the start sequence through the end condition, including all four D/F/J/K lanes.", False),
        18: ("Explain target generation, key and collision logic, score or miss changes, time, tempo, and one bug fix.", False),
        19: ("Answer one teacher question and identify one realistic improvement.", False),
        20: ("Rubric", True),
    }
    for index, (text, bold) in replacements.items():
        replace_paragraph_text(doc, index, text, bold=bold)

    rubric = doc.tables[1]
    repeat_table_header(rubric.rows[0])
    rows = (
        (
            "Four-lane game design",
            "Stage, hit line, D/F/J/K indicators, four matching target paths, and visual theme are complete and clear.",
            "Most design elements are complete; one element is unclear or inconsistent.",
            "Two or more required design elements are incomplete.",
            "The four-lane design is mostly missing or cannot be checked.",
        ),
        (
            "Target generation and timing",
            "Targets reliably show or spawn, move to the correct hit areas, follow tempo, and reset or hide correctly.",
            "Target movement and timing mostly work with minor errors.",
            "Some targets move, but timing or resetting is unreliable.",
            "Target generation or movement does not work.",
        ),
        (
            "Input, collision, and feedback",
            "D/F/J/K checks match the lanes; collision and hit/miss rules work; feedback is immediate and clear.",
            "Most key and collision checks work with minor errors.",
            "Some input works, but hit/miss detection is inconsistent.",
            "Input or collision logic is mostly missing.",
        ),
        (
            "Variables and game flow",
            "Score, time, miss consequence, tempo, start sequence, loop, music, and end condition work correctly.",
            "Most variables and game-flow features work with minor errors.",
            "Several variables or game-flow features are incomplete.",
            "Variables or game flow cannot be demonstrated.",
        ),
        (
            "Testing and improvement",
            "Includes 3 tests, 1 specific bug, a relevant code change, a retest result, and a clear improvement.",
            "Includes most testing evidence and a reasonable improvement.",
            "Testing evidence is limited or the improvement is unclear.",
            "Little or no testing evidence is submitted.",
        ),
        (
            "Presentation and responsibility",
            "Demo is complete; explanations and teacher answer are accurate; evidence is timely; class time and equipment are used responsibly.",
            "Presentation and submission are mostly complete with minor omissions.",
            "Presentation or evidence is incomplete, unclear, or late.",
            "Presentation is missing or the project cannot be evaluated.",
        ),
    )
    for row_index, row_values in enumerate(rows, start=1):
        prevent_row_split(rubric.rows[row_index])
        for column, text in enumerate(row_values):
            set_cell(rubric.cell(row_index, column), text, bold=column == 0, centered=column == 0)

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                format_run(run)

    doc.save(NEW_EXAM)
    print(f"created {NEW_EXAM}")


def archive_old_files() -> None:
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    for path in (OLD_SUM, OLD_EXAM):
        if path.exists():
            destination = ARCHIVE / path.name
            if destination.exists():
                destination.unlink()
            path.rename(destination)
            print(f"archived {path} -> {destination}")


def main() -> None:
    create_summative()
    create_exam()
    archive_old_files()


if __name__ == "__main__":
    main()
