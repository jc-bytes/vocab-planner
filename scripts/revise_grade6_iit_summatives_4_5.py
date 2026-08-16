#!/usr/bin/env python3
"""Rebuild the Grade 6 IIT Summative 4 and 5 rubrics with objective evidence.

The script intentionally preserves the existing school header, logo, week
numbers, filenames, page setup, and 40-point scoring structure. It repairs the
rubric table geometry, replaces broad descriptors with countable evidence, and
mirrors the finished files to Generated Rubrics 2026.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell


ROOT = Path(__file__).resolve().parents[1]
DAILY_ROOT = ROOT / "plans" / "6th Grade Technology" / "Rubrics" / "IIT" / "Daily"
GENERATED_ROOT = (
    ROOT
    / "plans"
    / "Generated Rubrics 2026"
    / "6th Grade Technology"
    / "2nd Trimester"
)

FONT_NAME = "Arial"
FONT_SIZE = Pt(12)
TABLE_WIDTHS = [1900, 2360, 2360, 2360, 2360]


@dataclass(frozen=True)
class RubricRevision:
    filename: str
    title: str
    task_line: str
    criteria: tuple[tuple[str, str, str, str, str], ...]
    submission: tuple[str, str, str, str, str]


REVISIONS = (
    RubricRevision(
        filename="6th grade - IIT - Week 7 - Rubric for Summative 4.docx",
        title="mBot Sensor Condition Flowchart",
        task_line=(
            "Task: Make one sensor flowchart with a condition, True and False responses, "
            "arrows and labels, and one English sentence. Use 3 words: sensor, condition, "
            "True, False, response."
        ),
        criteria=(
            (
                "Sensor condition",
                "All 3: mBot sensor, what it detects, and an if/then condition.",
                "Includes 2 of the 3 required parts.",
                "Includes 1 of the 3 required parts.",
                "The sensor condition is missing or has none of the required parts.",
            ),
            (
                "True and False responses",
                "All 4: True action, False action, different actions, and both connected to the condition.",
                "Includes 3 of the 4 required parts.",
                "Includes 2 of the 4 required parts.",
                "Includes 0 or 1 of the 4 required parts.",
            ),
            (
                "Flowchart structure",
                "All 5: Start, decision, True label, False label, and arrows.",
                "Includes 4 of the 5 required parts.",
                "Includes 2 or 3 of the 5 required parts.",
                "Includes 0 or 1 of the 5 required parts.",
            ),
            (
                "English explanation",
                "All 5: English sentence, sensor and condition, True response, False response, and 3 correct technical words.",
                "Includes 4 of the 5 required parts.",
                "Includes 2 or 3 of the 5 required parts.",
                "Includes 0 or 1 of the 5 required parts or no explanation.",
            ),
        ),
        submission=(
            "Submission evidence",
            "All 4: name, group, date, and a file that opens or paper that is readable.",
            "3 of the 4 submission checks are present.",
            "1 or 2 of the 4 submission checks are present.",
            "No flowchart is submitted.",
        ),
    ),
    RubricRevision(
        filename="6th grade - IIT - Week 12 - Rubric for Summative 5.docx",
        title="mBot Rescue Robot Challenge",
        task_line=(
            "Challenge: From START, complete 3 movements and 1 turn, stop in the RESCUE "
            "ZONE, and activate an LED or buzzer. Word bank: rescue, route, movement, "
            "turn, signal, test, improve."
        ),
        criteria=(
            (
                "Rescue challenge plan",
                "All 5: START, RESCUE ZONE, 3 ordered movements, 1 turn, and an LED or buzzer signal.",
                "Includes 4 of the 5 required parts.",
                "Includes 2 or 3 of the 5 required parts.",
                "Includes 0 or 1 of the 5 required parts or no plan.",
            ),
            (
                "Rescue robot run",
                "All 4: starts at START, follows the planned route, stops in the RESCUE ZONE, and activates the signal.",
                "Completes 3 of the 4 required actions.",
                "Completes 2 of the 4 required actions.",
                "Completes 0 or 1 of the 4 required actions.",
            ),
            (
                "Testing and improvement",
                "All 4: first test result, one problem, one program change, and final test result.",
                "Records 3 of the 4 required items.",
                "Records 2 of the 4 required items.",
                "Records 0 or 1 of the 4 required items.",
            ),
            (
                "One-minute English explanation",
                "All 5: rescue goal, route, signal, one improvement, and 3 correct challenge words.",
                "Includes 4 of the 5 required parts.",
                "Includes 2 or 3 of the 5 required parts.",
                "Includes 0 or 1 of the 5 required parts or no explanation.",
            ),
        ),
        submission=(
            "Submission evidence",
            "All 4: name and group, challenge plan, test log, and saved code or screenshot.",
            "3 of the 4 submission items are submitted.",
            "1 or 2 of the 4 submission items are submitted.",
            "No project evidence is submitted.",
        ),
    ),
)


def physical_cells(table, row_index: int) -> list[_Cell]:
    """Return physical cells without python-docx's grid-span duplicates."""
    return [_Cell(tc, table) for tc in table._tbl.tr_lst[row_index].tc_lst]


def set_run_font(run, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    for tag in ("w:sz", "w:szCs"):
        size_element = r_pr.find(qn(tag))
        if size_element is None:
            size_element = OxmlElement(tag)
            r_pr.append(size_element)
        size_element.set(qn("w:val"), "24")
    if bold is not None:
        run.bold = bold


def set_cell_text(cell: _Cell, text: str, *, bold: bool = False, centered: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell: _Cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.get_or_add_tcW()
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_grid_span(cell: _Cell, span: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in list(tc_pr.findall(qn("w:gridSpan"))):
        tc_pr.remove(existing)
    if span > 1:
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), str(span))
        tc_pr.append(grid_span)


def set_cell_margins(cell: _Cell, *, top: int = 70, start: int = 80, bottom: int = 70, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def repair_table_geometry(table) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in TABLE_WIDTHS:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl_width = tbl.tblPr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl.tblPr.insert(0, tbl_width)
    tbl_width.set(qn("w:w"), str(sum(TABLE_WIDTHS)))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_layout = tbl.tblPr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl.tblPr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    for row_index in range(len(tbl.tr_lst)):
        tr_pr = tbl.tr_lst[row_index].get_or_add_trPr()
        for tag in ("w:gridBefore", "w:wBefore", "w:gridAfter", "w:wAfter"):
            for stale in list(tr_pr.findall(qn(tag))):
                tr_pr.remove(stale)
        cells = physical_cells(table, row_index)
        if row_index in (0, 8):
            set_grid_span(cells[0], 5)
            set_cell_width(cells[0], sum(TABLE_WIDTHS))
        else:
            if len(cells) != 5:
                raise ValueError(f"Expected 5 physical cells in row {row_index + 1}, found {len(cells)}")
            for cell, width in zip(cells, TABLE_WIDTHS):
                set_grid_span(cell, 1)
                set_cell_width(cell, width)
                set_cell_margins(cell)


def replace_header_title(table, title: str, task_line: str) -> None:
    header = physical_cells(table, 0)[0]
    paragraphs = header.paragraphs
    if len(paragraphs) < 8:
        raise ValueError("The rubric header does not contain the expected title paragraph.")

    title_paragraph = paragraphs[7]
    for run in list(title_paragraph.runs):
        title_paragraph._element.remove(run._element)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.line_spacing = 1.0
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, bold=True)
    title_run.add_break()
    task_run = title_paragraph.add_run(task_line)
    set_run_font(task_run, bold=False)

    for paragraph in header.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run)


def strip_undeclared_ignorable_prefixes(path: Path) -> None:
    """Remove stale mc:Ignorable attributes that python-docx cannot preserve."""
    temporary = path.with_suffix(".compat.docx")
    pattern = re.compile(rb'\s+mc:Ignorable="[^"]*"')
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                data = pattern.sub(b"", data, count=1)
            target.writestr(info, data)
    temporary.replace(path)


def rebuild_rubric(revision: RubricRevision) -> tuple[Path, Path]:
    source = DAILY_ROOT / revision.filename
    document = Document(source)
    if len(document.tables) != 1:
        raise ValueError(f"Expected one rubric table in {source}")

    table = document.tables[0]
    repair_table_geometry(table)
    replace_header_title(table, revision.title, revision.task_line)

    for cell, text in zip(physical_cells(table, 1), ("Criteria", "9", "7", "5", "2")):
        set_cell_text(cell, text, bold=True, centered=True)

    for row_index, criterion in enumerate(revision.criteria, start=2):
        cells = physical_cells(table, row_index)
        for column_index, (cell, text) in enumerate(zip(cells, criterion)):
            set_cell_text(cell, text, bold=column_index == 0)

    for cell, text in zip(physical_cells(table, 6), ("Criteria", "4", "3", "2", "0")):
        set_cell_text(cell, text, bold=True, centered=True)

    for column_index, (cell, text) in enumerate(zip(physical_cells(table, 7), revision.submission)):
        set_cell_text(cell, text, bold=column_index == 0)

    set_cell_text(physical_cells(table, 8)[0], "Comments:", bold=True)

    # Normalize all visible table text after editing while leaving the logo intact.
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)

    document.save(source)
    strip_undeclared_ignorable_prefixes(source)
    GENERATED_ROOT.mkdir(parents=True, exist_ok=True)
    generated = GENERATED_ROOT / source.name
    shutil.copy2(source, generated)
    return source, generated


def main() -> None:
    for revision in REVISIONS:
        source, generated = rebuild_rubric(revision)
        print(source)
        print(generated)


if __name__ == "__main__":
    main()
