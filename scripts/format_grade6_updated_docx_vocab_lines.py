from pathlib import Path
import re

from docx import Document

from update_grade6_updated_docx_vocab_in_place import DEFINITIONS, DOCX_DIR, PRACTICE


def words_only(words):
    return ", ".join(words)


def definition_lines(label):
    return "\n".join(f"{word}: {DEFINITIONS[word]}" for word in PRACTICE[label])


def review_block(label):
    return f"Review {label} practice vocabulary.\n{definition_lines(label)}"


def format_vocabulary_text(text):
    original = text

    for label, words in PRACTICE.items():
        word_list = words_only(words)
        old_defs = "; ".join(f"{word} = {DEFINITIONS[word]}" for word in words)
        text = text.replace(
            f"Practice vocabulary ({label}): {old_defs}.",
            review_block(label),
        )

        text = re.sub(
            rf"([^.\n]*{re.escape(label)} practice vocabulary[^:\n]*): {re.escape(word_list)}\.",
            lambda match, current_label=label: f"{match.group(1)}.\n{definition_lines(current_label)}",
            text,
        )

    return text, text != original


def update_paragraph(paragraph):
    text, changed = format_vocabulary_text(paragraph.text)
    if not changed:
        return 0

    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return 1


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def main():
    total = 0
    for path in sorted(DOCX_DIR.glob("**/6° Technology - *.docx")):
        if path.name.startswith("~$"):
            continue
        doc = Document(path)
        count = sum(update_paragraph(paragraph) for paragraph in iter_paragraphs(doc))
        if count:
            doc.save(path)
        total += count
        print(f"{path.name}: {count} paragraphs updated")
    print(f"total paragraphs updated: {total}")


if __name__ == "__main__":
    main()
