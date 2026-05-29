from pathlib import Path
import json
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = ROOT / "plans/9th grade/9th Grade Technology - Updated"
VOCAB_DIR = ROOT / "vocabularies/grade9"


TOPIC_WEEKS = {
    "9° Technology - March.docx": {
        "Digital and analog signals": ("Week 2", "grade9_it_march_week2_blink_signals.json"),
        "LED blink and core vocabulary and program structure": ("Week 2", "grade9_it_march_week2_blink_signals.json"),
        "Vocabulary and debugging routine": ("Week 3", "grade9_it_march_week3_lab_routine.json"),
        "Mini LED control practice": ("Week 3", "grade9_it_march_week3_lab_routine.json"),
    },
    "9° Technology - April.docx": {
        "Button input": ("Week 4", "grade9_it_april_week4_button_toggle.json"),
        "Toggle control": ("Week 4", "grade9_it_april_week4_button_toggle.json"),
        "Serial monitor": ("Week 5", "grade9_it_april_week5_serial_analog.json"),
        "Analog input with potentiometer": ("Week 5", "grade9_it_april_week5_serial_analog.json"),
        "Sensor threshold quiz-style check": ("Week 6", "grade9_it_april_week6_sensor_rgb.json"),
        "RGB output and state colors": ("Week 6", "grade9_it_april_week6_sensor_rgb.json"),
        "Addressable LEDs and animation": ("Week 7", "grade9_it_april_week7_led_teamwork.json"),
        "Team robotics practice": ("Week 7", "grade9_it_april_week7_led_teamwork.json"),
        "State logic chart": ("Week 8", "grade9_it_april_week8_state_prototype.json"),
        "Integrated input-output prototype": ("Week 8", "grade9_it_april_week8_state_prototype.json"),
    },
    "9° Technology - May.docx": {
        "Robotics project planning": ("Week 9", "grade9_it_may_week9_robot_design.json"),
        "Prototype rehearsal": ("Week 9", "grade9_it_may_week9_robot_design.json"),
        "Exam project setup": ("Week 10", "grade9_it_may_week10_project_build.json"),
        "Exam project build": ("Week 10", "grade9_it_may_week10_project_build.json"),
        "Exam project testing": ("Week 11", "grade9_it_may_week11_reliability.json"),
        "Exam project improvement": ("Week 11", "grade9_it_may_week11_reliability.json"),
        "Presentation rehearsal": ("Week 12", "grade9_it_may_week12_demonstration.json"),
        "Robotics project presentations": ("Week 12", "grade9_it_may_week12_demonstration.json"),
        "Exam/admin buffer": ("Week 13", "grade9_it_may_week13_archive.json"),
    },
    "9° Technology - June.docx": {
        "While loops": ("Week 2", "grade9_iit_june_week2_list_methods.json"),
        "For loops with lists": ("Week 2", "grade9_iit_june_week2_list_methods.json"),
        "Loop debugging quiz-style check": ("Week 3", "grade9_iit_june_week3_loops_strings.json"),
        "Strings as data": ("Week 3", "grade9_iit_june_week3_loops_strings.json"),
    },
    "9° Technology - July.docx": {
        "Mini program with sequences": ("Week 4", "grade9_iit_july_week4_data_science.json"),
        "From programming to data science": ("Week 4", "grade9_iit_july_week4_data_science.json"),
        "Data collection plan": ("Week 5", "grade9_iit_july_week5_collection_plan.json"),
        "Collect and organize data": ("Week 5", "grade9_iit_july_week5_collection_plan.json"),
        "Data cleaning habits": ("Week 6", "grade9_iit_july_week6_cleaning_model.json"),
        "Tinkercad data model intro": ("Week 6", "grade9_iit_july_week6_cleaning_model.json"),
        "Charts and patterns": ("Week 7", "grade9_iit_july_week7_charts_patterns.json"),
        "Correlation and outliers": ("Week 7", "grade9_iit_july_week7_charts_patterns.json"),
    },
    "9° Technology - August.docx": {
        "Data conclusions": ("Week 8", "grade9_iit_august_week8_data_product.json"),
        "Data product prototype": ("Week 8", "grade9_iit_august_week8_data_product.json"),
        "Project readiness": ("Week 9", "grade9_iit_august_week9_project_readiness.json"),
        "Data project planning": ("Week 9", "grade9_iit_august_week9_project_readiness.json"),
        "Exam project setup": ("Week 10", "grade9_iit_august_week10_visualization.json"),
        "Exam project build": ("Week 10", "grade9_iit_august_week10_visualization.json"),
        "Exam project testing": ("Week 11", "grade9_iit_august_week11_revision.json"),
        "Exam project finalization": ("Week 11", "grade9_iit_august_week11_revision.json"),
        "Presentation rehearsal": ("Week 12", "grade9_iit_august_week12_presentation.json"),
        "Data project presentations": ("Week 12", "grade9_iit_august_week12_presentation.json"),
    },
    "9° Technology - September.docx": {
        "Sound representation": ("Week 2", "grade9_iiit_september_week2_sound_quality.json"),
        "Sonic playground": ("Week 2", "grade9_iiit_september_week2_sound_quality.json"),
    },
    "9° Technology - October.docx": {
        "You and your data": ("Week 3", "grade9_iiit_october_week3_personal_data.json"),
        "Digital responsibility habits": ("Week 3", "grade9_iiit_october_week3_personal_data.json"),
        "Social engineering": ("Week 4", "grade9_iiit_october_week4_social_malware.json"),
        "Script kiddies and malware": ("Week 4", "grade9_iiit_october_week4_social_malware.json"),
        "Cyber security scenario quiz": ("Week 5", "grade9_iiit_october_week5_security_risk.json"),
        "Risk and impact": ("Week 5", "grade9_iiit_october_week5_security_risk.json"),
        "STEM project problem selection": ("Week 6", "grade9_iiit_october_week6_project_problem.json"),
        "STEM prototype planning": ("Week 6", "grade9_iiit_october_week6_project_problem.json"),
        "Risk map final": ("Week 7", "grade9_iiit_october_week7_risk_map.json"),
        "STEM design refinement": ("Week 7", "grade9_iiit_october_week7_risk_map.json"),
    },
    "9° Technology - November.docx": {
        "STEM proposal": ("Week 8", "grade9_iiit_november_week8_proposal.json"),
        "Prototype practice": ("Week 8", "grade9_iiit_november_week8_proposal.json"),
        "Project habits reflection": ("Week 9", "grade9_iiit_november_week9_project_habits.json"),
        "Exam project planning": ("Week 9", "grade9_iiit_november_week9_project_habits.json"),
        "Exam project setup": ("Week 10", "grade9_iiit_november_week10_build_evidence.json"),
        "Exam project build": ("Week 10", "grade9_iiit_november_week10_build_evidence.json"),
        "Exam project testing": ("Week 11", "grade9_iiit_november_week11_testing_improvement.json"),
        "Exam project improvement": ("Week 11", "grade9_iiit_november_week11_testing_improvement.json"),
    },
    "9° Technology - December.docx": {
        "Presentation rehearsal": ("Week 12", "grade9_iiit_december_week12_final_demo.json"),
        "STEM project presentations": ("Week 12", "grade9_iiit_december_week12_final_demo.json"),
    },
}


def load_vocab(file_name):
    data = json.loads((VOCAB_DIR / file_name).read_text())
    return [(item["word"], item["definition"].rstrip(".")) for item in data["words"]]


def split_words(words):
    midpoint = (len(words) + 1) // 2
    return words[:midpoint], words[midpoint:]


def definition_lines(words):
    return "\n".join(f"{word}: {definition}" for word, definition in words)


def get_topic(text):
    match = re.search(r"Topic:\n+([^\n]+)", text)
    return match.group(1).strip() if match else ""


def clean_old_vocab_prompt(text):
    replacements = [
        r"Match these words with definitions: debug, test, upload, sketch, circuit\.",
        r"Play a vocabulary review game using sound and image terms\.",
    ]
    for pattern in replacements:
        text = re.sub(pattern, "", text)
    return re.sub(r"\n{3,}", "\n\n", text)


def insert_vocab_block(text, week, words, part_number):
    if definition_lines(words) in text:
        return text, False

    sentence = f"Review {week} practice vocabulary (part {part_number})."
    block = f"{sentence}\n{definition_lines(words)}"
    marker = "Pre-activities\n"

    if marker not in text:
        return text, False

    text = clean_old_vocab_prompt(text)
    if f"{week} practice vocabulary" in text:
        pattern = re.compile(
            rf"Review {re.escape(week)} practice vocabulary \(part [12]\)\.\n"
            r"(?:[^\n]+: [^\n]+\n?)+",
            re.I,
        )
        text = pattern.sub("", text).replace(marker + "\n", marker)

    return text.replace(marker, marker + block + "\n", 1), True


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_docx(doc_name):
    direct = DOCX_DIR / doc_name
    if direct.exists():
        return direct
    matches = sorted(DOCX_DIR.glob(f"*/{doc_name}"))
    if matches:
        return matches[0]
    return direct


def main():
    total = 0
    for doc_name, topic_weeks in TOPIC_WEEKS.items():
        path = find_docx(doc_name)
        doc = Document(path)
        paragraphs = list(iter_paragraphs(doc))
        week_seen = {}
        doc_count = 0

        for paragraph in paragraphs:
            topic = get_topic(paragraph.text)
            if topic not in topic_weeks:
                continue
            week, vocab_file = topic_weeks[topic]
            words = load_vocab(vocab_file)
            first_half, second_half = split_words(words)
            occurrence = week_seen.get(week, 0) + 1
            week_seen[week] = occurrence
            part_number = 1 if occurrence == 1 else 2
            part_words = first_half if occurrence == 1 else second_half
            text, changed = insert_vocab_block(paragraph.text, week, part_words, part_number)
            if changed:
                set_paragraph_text(paragraph, text)
                doc_count += 1

        if doc_count:
            doc.save(path)
        total += doc_count
        print(f"{doc_name}: {doc_count} paragraphs updated")

    print(f"total paragraphs updated: {total}")


if __name__ == "__main__":
    main()
