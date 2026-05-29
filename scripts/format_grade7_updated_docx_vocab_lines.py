from pathlib import Path
import json
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = ROOT / "plans/7th grade/7th Grade Technology - Updated"
VOCAB_DIR = ROOT / "vocabularies/grade7"


WEEK_FILES = {
    "7° Technology - March.docx": {
        "Week 2": "grade7_it_march_week2_mblock_workspace.json",
        "Week 3": "grade7_it_march_week3_movement_commands.json",
        "Week 4": "grade7_it_march_week4_robot_systems.json",
    },
    "7° Technology - April.docx": {
        "Week 5": "grade7_it_april_week5_calibration.json",
        "Week 6": "grade7_it_april_week6_route_planning.json",
        "Week 7": "grade7_it_april_week7_feedback_signals.json",
        "Week 8": "grade7_it_april_week8_sensor_conditions.json",
    },
    "7° Technology - May.docx": {
        "Week 9": "grade7_it_may_week9_maze_planning.json",
        "Week 10": "grade7_it_may_week10_project_plan.json",
        "Week 11": "grade7_it_may_week11_testing_debugging.json",
        "Week 12": "grade7_it_may_week12_presentation.json",
    },
    "7° Technology - June.docx": {
        "Week 2": "grade7_iit_june_week2_branding.json",
    },
    "7° Technology - July.docx": {
        "Week 3": "grade7_iit_july_week3_scratch_intro.json",
        "Week 4": "grade7_iit_july_week4_variables.json",
        "Week 5": "grade7_iit_july_week5_operators.json",
        "Week 6": "grade7_iit_july_week6_loops.json",
        "Week 7": "grade7_iit_july_week7_debugging_roles.json",
    },
    "7° Technology - August.docx": {
        "Week 8": "grade7_iit_august_week8_dance_prep.json",
        "Week 9": "grade7_iit_august_week9_project_planning.json",
        "Week 10": "grade7_iit_august_week10_build.json",
        "Week 11": "grade7_iit_august_week11_improve.json",
        "Week 12": "grade7_iit_august_week12_demo.json",
    },
    "7° Technology - September.docx": {
        "Week 2": "grade7_iiit_september_week2_formulas_charts.json",
    },
    "7° Technology - October.docx": {
        "Week 3": "grade7_iiit_october_week3_data_analysis.json",
        "Week 4": "grade7_iiit_october_week4_scratch_decomposition.json",
        "Week 5": "grade7_iiit_october_week5_sources_lists.json",
        "Week 6": "grade7_iiit_october_week6_blog_media.json",
        "Week 7": "grade7_iiit_october_week7_sensor_systems.json",
    },
    "7° Technology - November.docx": {
        "Week 8": "grade7_iiit_november_week8_test_table.json",
        "Week 9": "grade7_iiit_november_week9_review_setup.json",
        "Week 10": "grade7_iiit_november_week10_project_plan.json",
        "Week 11": "grade7_iiit_november_week11_debug_reliability.json",
        "Week 12": "grade7_iiit_november_week12_presentation.json",
    },
}


MISSING_TOPIC_WEEKS = {
    "7° Technology - May.docx": {
        "Maze prototype and final project preparation": "Week 10",
        "Build and program maze sections": "Week 11",
        "Final testing and documentation": "Week 12",
    },
    "7° Technology - August.docx": {
        "Project rubric and planning": "Week 9",
        "Begin Exam Project: Scratch Dance Game": "Week 10",
        "Final game improvement and presentation prep": "Week 11",
        "Final dance game demonstration": "Week 12",
    },
    "7° Technology - November.docx": {
        "Threshold testing practice": "Week 8",
        "Project rubric and readiness": "Week 9",
        "Build and test the sensor logic": "Week 10",
        "Presentation rehearsal": "Week 12",
    },
}


def load_vocab(file_name):
    data = json.loads((VOCAB_DIR / file_name).read_text())
    return [(item["word"], item["definition"].rstrip(".")) for item in data["words"]]


def split_words(words):
    midpoint = (len(words) + 1) // 2
    return words[:midpoint], words[midpoint:]


def definition_lines(words):
    return "\n".join(f"{word}: {definition}" for word, definition in words)


def sentence_for(week, part_number=None):
    suffix = f" (part {part_number})" if part_number else ""
    return f"Review {week} practice vocabulary{suffix}."


def insert_after_pre_activities(text, week):
    marker = "Pre-activities\n"
    if marker not in text:
        return text, False
    insertion = sentence_for(week) + "\n"
    return text.replace(marker, marker + insertion, 1), True


def get_topic(text):
    match = re.search(r"Topic:\n+([^\n]+)", text)
    return match.group(1).strip() if match else ""


def update_paragraph_text(text, week, words, part_number=None):
    if definition_lines(words) in text:
        return text, False

    pattern = re.compile(rf"([^\n.]*\b{re.escape(week)} practice vocabulary[^\n.]*\.)", re.I)
    match = pattern.search(text)
    if match:
        replacement = f"{sentence_for(week, part_number)}\n{definition_lines(words)}"
        return text[:match.start(1)] + replacement + text[match.end(1):], True

    return text, False


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_docx(doc_name):
    direct = DOCX_DIR / doc_name
    if direct.exists():
        return direct
    matches = sorted(DOCX_DIR.glob(f"*/{doc_name}"))
    if matches:
        return matches[0]
    return direct


def main():
    total = 0
    for doc_name, week_files in WEEK_FILES.items():
        path = find_docx(doc_name)
        doc = Document(path)
        paragraphs = list(iter_paragraphs(doc))
        doc_count = 0

        for topic, week in MISSING_TOPIC_WEEKS.get(doc_name, {}).items():
            for paragraph in paragraphs:
                if get_topic(paragraph.text) != topic:
                    continue
                if f"{week} practice vocabulary" in paragraph.text:
                    continue
                text, changed = insert_after_pre_activities(paragraph.text, week)
                if changed:
                    set_paragraph_text(paragraph, text)
                    doc_count += 1
                break

        for week, vocab_file in week_files.items():
            words = load_vocab(vocab_file)
            matching = [
                paragraph for paragraph in paragraphs
                if f"{week} practice vocabulary" in paragraph.text
            ]
            if not matching:
                continue

            if len(matching) >= 2:
                first_half, second_half = split_words(words)
                splits = [first_half, second_half]
                for index, paragraph in enumerate(matching[:2]):
                    text, changed = update_paragraph_text(paragraph.text, week, splits[index], index + 1)
                    if changed:
                        set_paragraph_text(paragraph, text)
                        doc_count += 1
                for paragraph in matching[2:]:
                    text, changed = update_paragraph_text(paragraph.text, week, words)
                    if changed:
                        set_paragraph_text(paragraph, text)
                        doc_count += 1
            else:
                text, changed = update_paragraph_text(matching[0].text, week, words)
                if changed:
                    set_paragraph_text(matching[0], text)
                    doc_count += 1

        if doc_count:
            doc.save(path)
        total += doc_count
        print(f"{doc_name}: {doc_count} paragraphs updated")
    print(f"total paragraphs updated: {total}")


if __name__ == "__main__":
    main()
