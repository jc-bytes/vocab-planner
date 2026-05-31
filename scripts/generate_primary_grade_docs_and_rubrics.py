#!/usr/bin/env python3
"""Generate primary-grade monthly planning DOCX files and summative rubrics.

The generator is intentionally scoped to the newer TECHNOLOGY 3° and
TECHNOLOGY 4° folder layout. It uses the Markdown monthly drafts as the
planning source and leaves existing official DOCX/PDF files untouched.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PLANS_ROOT = ROOT / "plans"
MONTHLY_TEMPLATE = PLANS_ROOT / "Templates" / "X\u00b0 Technology - Month.docx"
GENERATED_RUBRICS_ROOT = PLANS_ROOT / "Generated Rubrics 2026"

GRADES = {
    "3": {
        "folder": PLANS_ROOT / "3rd Grade Technology",
        "ordinal": "3rd",
        "display": "3\u00b0",
        "groups_text": "3rd A & B & C & D",
        "group_line": "3\u00b0 A B C D",
        "section_line": "3A, 3B, 3C, 3D",
        "teacher": "Edgar Pitty",
    },
    "4": {
        "folder": PLANS_ROOT / "4th Grade Technology",
        "ordinal": "4th",
        "display": "4\u00b0",
        "groups_text": "4th A & B & C",
        "group_line": "4\u00b0 A B C",
        "section_line": "4A, 4B, 4C",
        "teacher": "Edgar Pitty",
    },
}

TRIMESTERS = {
    "1st Trimester": ("IT", "1st"),
    "2nd Trimester": ("IIT", "2nd"),
    "3rd Trimester": ("IIIT", "3rd"),
}

MONTH_NUMBERS = {
    "January": "01",
    "February": "02",
    "March": "03",
    "April": "04",
    "May": "05",
    "June": "06",
    "July": "07",
    "August": "08",
    "September": "09",
    "October": "10",
    "November": "11",
    "December": "12",
}

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_NS = NS["w"]


@dataclass
class LessonRow:
    week: str
    class_label: str
    topic: str
    pre: str
    while_: str
    post: str


@dataclass
class MonthlyDraft:
    path: Path
    grade_number: str
    grade_ordinal: str
    trimester_label: str
    trimester_code: str
    trimester_short: str
    month: str
    title: str
    sections: dict[str, str]
    rows: list[LessonRow]


@dataclass
class Summative:
    grade_number: str
    grade_ordinal: str
    grade_display: str
    trimester_label: str
    trimester_code: str
    week_number: int
    summative_number: int
    title: str
    activity_text: str
    evidence_text: str
    groups_text: str
    group_line: str
    teacher: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate DOCX monthly plans and rubrics for primary grades.")
    parser.add_argument("--grades", nargs="+", choices=sorted(GRADES), default=sorted(GRADES))
    parser.add_argument("--monthly-only", action="store_true", help="Generate monthly DOCX files only.")
    parser.add_argument("--rubrics-only", action="store_true", help="Generate rubric DOCX files only.")
    return parser.parse_args()


def clean_text(value: str) -> str:
    value = value.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"`([^`]*)`", r"\1", value)
    value = value.replace("&amp;", "&")
    return "\n".join(line.strip() for line in value.splitlines() if line.strip())


def markdown_sections(text: str) -> tuple[str, dict[str, str]]:
    title = "Technology"
    sections: dict[str, list[str]] = {}
    current: str | None = None

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
            continue
        if current is not None:
            sections[current].append(line)

    return title, {key: "\n".join(value).strip() for key, value in sections.items()}


def parse_markdown_table(section: str) -> list[LessonRow]:
    rows: list[LessonRow] = []
    table_lines = [line.strip() for line in section.splitlines() if line.strip().startswith("|")]
    if len(table_lines) < 3:
        return rows

    for line in table_lines[2:]:
        parts = [part.strip() for part in line.strip("|").split("|")]
        if len(parts) < 6:
            continue
        rows.append(
            LessonRow(
                week=clean_text(parts[0]),
                class_label=clean_text(parts[1]),
                topic=clean_text(parts[2]),
                pre=clean_text(parts[3]),
                while_=clean_text(parts[4]),
                post=clean_text(parts[5]),
            )
        )
    return rows


def find_month_from_title(title: str, path: Path) -> str:
    haystack = f"{title} {path.name}"
    for month in MONTH_NUMBERS:
        if re.search(rf"\b{month}\b", haystack, re.IGNORECASE):
            return month
    raise ValueError(f"Could not determine month for {path}")


def monthly_draft_paths(grade_folder: Path) -> list[Path]:
    root = grade_folder / "Draft Planning" / "Monthly Planning"
    return sorted(root.glob("* Trimester/*/*.md"))


def load_monthly_drafts(grade_number: str) -> list[MonthlyDraft]:
    config = GRADES[grade_number]
    drafts: list[MonthlyDraft] = []
    for path in monthly_draft_paths(config["folder"]):
        trimester_label = next((part for part in path.parts if part in TRIMESTERS), None)
        if trimester_label is None:
            continue
        trimester_code, trimester_short = TRIMESTERS[trimester_label]
        text = path.read_text(encoding="utf-8", errors="ignore")
        title, sections = markdown_sections(text)
        rows = parse_markdown_table(sections.get("Monthly Plan", ""))
        drafts.append(
            MonthlyDraft(
                path=path,
                grade_number=grade_number,
                grade_ordinal=config["ordinal"],
                trimester_label=trimester_label,
                trimester_code=trimester_code,
                trimester_short=trimester_short,
                month=find_month_from_title(title, path),
                title=title,
                sections=sections,
                rows=rows,
            )
        )
    return drafts


def bullet_items(section: str) -> list[str]:
    items: list[str] = []
    for raw_line in section.splitlines():
        line = raw_line.strip()
        if line.startswith("- "):
            items.append(clean_text(line[2:]))
    return items


def topics_for_month(draft: MonthlyDraft) -> list[str]:
    topics: list[str] = []
    for row in draft.rows:
        if row.topic not in topics:
            topics.append(row.topic)
    return topics


def fallback_competences(draft: MonthlyDraft) -> list[str]:
    return [
        f"Follow safe Grade {draft.grade_number} technology routines while working with monthly tools, vocabulary, and class materials.",
        "Identify, label, sort, create, test, and submit visible evidence using guided class steps.",
        "Use technology vocabulary responsibly in drawings, worksheets, notebooks, or digital products.",
    ]


def fallback_objectives(draft: MonthlyDraft) -> list[str]:
    return [
        "Identify key vocabulary and examples from the monthly topics.",
        "Complete the guided class evidence named in each lesson with teacher support.",
        "Check work for clear labels, correct steps, safe handling, and a simple reflection.",
    ]


def fallback_outcomes(draft: MonthlyDraft) -> list[str]:
    summatives = bullet_items(draft.sections.get("Evaluation", ""))
    summative_lines = [item for item in summatives if "Summative Activity" in item]
    if not summative_lines:
        summative_lines = ["Complete the formal summative activity or activities scheduled for this month."]
    return summative_lines + ["Complete formative notebook, worksheet, drawing, digital, or reflection evidence."]


def month_week_count(draft: MonthlyDraft) -> int:
    weekly_time = draft.sections.get("Weekly Time", "")
    match = re.search(r"Month plan weeks:\s*(\d+)", weekly_time)
    if match:
        return int(match.group(1))
    return len({row.week for row in draft.rows})


def clear_block(parent) -> None:
    for paragraph in list(parent.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def add_run(paragraph, text: str, *, bold: bool = False, size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    if size:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{key}"), "Arial")
    if size:
        half_points = str(size * 2)
        for tag in ("sz", "szCs"):
            node = r_pr.find(qn(f"w:{tag}"))
            if node is None:
                node = OxmlElement(f"w:{tag}")
                r_pr.append(node)
            node.set(qn("w:val"), half_points)


def add_labeled_lines(cell, label: str, lines: Iterable[str], *, bullet: bool = False, font_size: int = 9) -> None:
    clear_block(cell)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    add_run(paragraph, label, bold=True, size=font_size)
    paragraph.paragraph_format.space_after = Pt(1)
    for line in lines:
        if not line:
            continue
        p = cell.add_paragraph()
        prefix = "\u2022 " if bullet else ""
        add_run(p, prefix + line, size=font_size)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9, align=None) -> None:
    clear_block(cell)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    add_run(paragraph, text, bold=bold, size=size)
    paragraph.paragraph_format.space_after = Pt(1)
    if align is not None:
        paragraph.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def remove_all_rows(table) -> None:
    for row in list(table.rows):
        table._tbl.remove(row._tr)


def split_lines(text: str) -> list[str]:
    return [line.strip() for line in clean_text(text).splitlines() if line.strip()]


def write_activity_cell(cell, row: LessonRow) -> None:
    clear_block(cell)
    blocks = [
        ("Topic:", [row.topic]),
        ("Pre-Activities:", split_lines(row.pre)),
        ("While Activities:", split_lines(row.while_)),
        ("Post Activities:", split_lines(row.post)),
    ]
    first = True
    for label, lines in blocks:
        if not first:
            spacer = cell.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)
        first = False
        p = cell.add_paragraph()
        add_run(p, label, bold=True, size=8)
        p.paragraph_format.space_after = Pt(0)
        for line in lines:
            item = cell.add_paragraph()
            add_run(item, line, size=8)
            item.paragraph_format.space_after = Pt(1)
            item.paragraph_format.line_spacing = 1.0
    for paragraph in cell.paragraphs:
        if not paragraph.text:
            p = paragraph._element
            p.getparent().remove(p)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def style_document_defaults(doc: Document, font_size: int = 10) -> None:
    for style_name in ("Normal", "Table Grid"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Arial"
        style.font.size = Pt(font_size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def merge_row_across(table, row_index: int):
    cells = table.rows[row_index].cells
    return cells[0].merge(cells[-1])


def fill_monthly_meta_table(doc: Document, draft: MonthlyDraft) -> None:
    config = GRADES[draft.grade_number]
    table = doc.tables[0]
    info = [
        "Subject: Technology",
        f"Teacher: {config['teacher']}",
        f"Grade: {config['ordinal']}",
        f"Trimester:{draft.trimester_short}",
        "Year: 2026",
        "Weekly Hours: 3",
    ]
    for cell, text in zip(table.rows[0].cells, info):
        set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    competences = bullet_items(draft.sections.get("Competences", "")) or fallback_competences(draft)
    objectives = bullet_items(draft.sections.get("Learning Objectives", "")) or fallback_objectives(draft)
    outcomes = bullet_items(draft.sections.get("Learning Outcomes", "")) or fallback_outcomes(draft)

    for row_index, label, values in (
        (1, "Competences:", competences),
        (2, "Learning Objectives:", objectives),
        (3, "Learning Outcomes:", outcomes),
    ):
        cell = merge_row_across(table, row_index)
        add_labeled_lines(cell, label, values, bullet=True, font_size=9)

    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 80, 80, 80, 80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def rows_by_week(draft: MonthlyDraft) -> list[tuple[str, LessonRow | None, LessonRow | None]]:
    grouped: list[tuple[str, list[LessonRow]]] = []
    index: dict[str, list[LessonRow]] = {}
    for row in draft.rows:
        if row.week not in index:
            index[row.week] = []
            grouped.append((row.week, index[row.week]))
        index[row.week].append(row)

    pairs: list[tuple[str, LessonRow | None, LessonRow | None]] = []
    for week, rows in grouped:
        class1 = next((row for row in rows if "Class 1" in row.class_label), rows[0] if rows else None)
        class2 = next((row for row in rows if "Class 2" in row.class_label), rows[1] if len(rows) > 1 else None)
        pairs.append((week, class1, class2))
    return pairs


def fill_monthly_lesson_table(doc: Document, draft: MonthlyDraft) -> None:
    config = GRADES[draft.grade_number]
    table = doc.tables[1]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    remove_all_rows(table)
    usable_width = 14400
    set_table_width(table, usable_width)
    col_width = usable_width // 2

    for week, class1, class2 in rows_by_week(draft):
        header = table.add_row()
        detail = table.add_row()
        for cell in header.cells + detail.cells:
            set_cell_width(cell, col_width)
            set_cell_margins(cell, 90, 100, 90, 100)

        for cell in header.cells:
            shade_cell(cell, "D9EAF7")
        set_cell_text(
            header.cells[0],
            f"{config['section_line']} - 45 min - {week}",
            bold=True,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(
            header.cells[1],
            f"{config['section_line']} - 90 min - {week}",
            bold=True,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        if class1:
            write_activity_cell(detail.cells[0], class1)
        else:
            set_cell_text(detail.cells[0], "No class scheduled.", size=8)
        if class2:
            write_activity_cell(detail.cells[1], class2)
        else:
            set_cell_text(detail.cells[1], "No class scheduled.", size=8)


def generate_monthly_docx(draft: MonthlyDraft) -> Path:
    config = GRADES[draft.grade_number]
    doc = Document(str(MONTHLY_TEMPLATE))
    style_document_defaults(doc, font_size=9)
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    headings = [
        "MINISTERIO DE EDUCACI\u00d3N DE PANAM\u00c1",
        "ACADEMIA INTERNACIONAL DAVID",
        "WEEKLY/MONTHLY PLAN",
        f"MONTH: {draft.month.upper()}",
        f"WEEKS: {month_week_count(draft)}",
    ]
    for idx, text in enumerate(headings):
        if idx >= len(doc.paragraphs):
            doc.add_paragraph()
        paragraph = doc.paragraphs[idx]
        paragraph.clear()
        add_run(paragraph, text, bold=idx in {0, 1, 2}, size=11 if idx < 3 else 10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)

    fill_monthly_meta_table(doc, draft)
    fill_monthly_lesson_table(doc, draft)

    output_dir = config["folder"] / f"{config['ordinal']} Grade Monthly Planning" / draft.trimester_label
    output_dir.mkdir(parents=True, exist_ok=True)
    suffix = " September Closure" if "Closure" in draft.path.stem else f" {draft.month}"
    output_path = output_dir / f"{config['display']} Technology -{suffix}.docx"
    doc.save(output_path)
    return output_path


def extract_summative_text(row: LessonRow) -> tuple[int | None, str, str]:
    combined = "\n".join([row.while_, row.post])
    match = re.search(r"Complete Summative Activity #(\d+) \(Summative\):\s*(.*)", combined, re.IGNORECASE | re.DOTALL)
    if not match:
        return None, "", ""
    number = int(match.group(1))
    after = clean_text(match.group(2))
    parts = after.splitlines()
    activity = parts[0].strip() if parts else after
    evidence_lines = [line for line in parts[1:] if "evidence" in line.lower() or "include" in line.lower()]
    evidence = " ".join(evidence_lines).strip() or activity
    return number, activity, evidence


def collect_summatives(drafts: list[MonthlyDraft]) -> list[Summative]:
    items: list[Summative] = []
    for draft in drafts:
        config = GRADES[draft.grade_number]
        for row in draft.rows:
            number, activity, evidence = extract_summative_text(row)
            if number is None:
                continue
            week_match = re.search(r"W(\d+)", row.week)
            if not week_match:
                raise ValueError(f"Could not parse week number from {row.week} in {draft.path}")
            title = row.topic.strip() or activity.split(".")[0]
            items.append(
                Summative(
                    grade_number=draft.grade_number,
                    grade_ordinal=config["ordinal"],
                    grade_display=config["display"],
                    trimester_label=draft.trimester_label,
                    trimester_code=draft.trimester_code,
                    week_number=int(week_match.group(1)),
                    summative_number=number,
                    title=title,
                    activity_text=activity,
                    evidence_text=evidence,
                    groups_text=config["groups_text"],
                    group_line=config["group_line"],
                    teacher=config["teacher"],
                )
            )
    return sorted(items, key=lambda item: (item.grade_number, item.trimester_code, item.week_number, item.summative_number))


def short_evidence(text: str, max_len: int = 120) -> str:
    text = re.sub(r"\s+", " ", clean_text(text)).strip()
    if len(text) <= max_len:
        return text
    return text[: max_len - 3].rsplit(" ", 1)[0] + "..."


def criterion_rows(summative: Summative) -> list[tuple[str, str, str, str, str]]:
    return [
        (
            "Required evidence",
            "All required activity parts are complete and easy to check.",
            "Most parts are complete; one small part is missing or unclear.",
            "Some parts are missing, unclear, or hard to check.",
            "Most evidence is missing, not submitted, or off task.",
        ),
        (
            "Correct content",
            "Answers, labels, matches, steps, or work are accurate.",
            "Most content is accurate, with only minor mistakes.",
            "Several answers, labels, matches, or steps are partly incorrect.",
            "Content is mostly incorrect, copied, or incomplete.",
        ),
        (
            "Organization and care",
            "Work is neat, organized, named correctly, and follows order.",
            "Work is readable, with minor naming, order, or neatness issues.",
            "Work can be checked, but organization needs improvement.",
            "Work is disorganized, hard to read, or misplaced.",
        ),
        (
            "Explanation or check",
            "Reflection, explanation, test, or teacher check supports the work.",
            "The check is mostly clear and connected to the work.",
            "The check is too short, unclear, or weakly connected.",
            "The check is missing or not understandable.",
        ),
    ]


def set_rubric_cell(cell, text: str, *, bold: bool = False, align=None, shade: str | None = None) -> None:
    set_cell_text(cell, text, bold=bold, size=12, align=align)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(1)
    if shade:
        shade_cell(cell, shade)
    set_cell_margins(cell, 100, 100, 100, 100)


def generate_rubric_docx(summative: Summative) -> tuple[Path, Path]:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style_document_defaults(doc, font_size=12)

    table = doc.add_table(rows=9, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 10560)
    widths = [1900, 2500, 2350, 2200, 1610]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    header_cell = table.cell(0, 0).merge(table.cell(0, 4))
    header_lines = [
        "Academia Internacional David",
        "Robotics and Technology",
        summative.trimester_label,
        f"Summative #{summative.summative_number}",
        summative.groups_text,
        f"Name: ________________________________   Date: __________________   Group: {summative.group_line}",
        f"Teacher: {summative.teacher}                  Score: _____ / 40pts",
        summative.title,
    ]
    set_rubric_cell(header_cell, "\n".join(header_lines), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="EAF4FB")

    for idx, text in enumerate(["Criteria", "9", "7", "5", "2"]):
        set_rubric_cell(table.cell(1, idx), text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="D9EAF7")

    for row_idx, values in enumerate(criterion_rows(summative), start=2):
        for col_idx, text in enumerate(values):
            set_rubric_cell(table.cell(row_idx, col_idx), text, bold=col_idx == 0)

    for idx, text in enumerate(["Criteria", "4", "3", "2", "0"]):
        set_rubric_cell(table.cell(6, idx), text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="D9EAF7")

    readiness = [
        "Punctuality,\nReadiness & Respect",
        "Submitted on time; ready to work; class routines followed.",
        "Submitted on time, with one minor readiness or routine issue.",
        "Submitted, but lateness, readiness, or routines affected work.",
        "Not submitted on time, or not ready to complete the activity.",
    ]
    for idx, text in enumerate(readiness):
        set_rubric_cell(table.cell(7, idx), text, bold=idx == 0)

    comments = table.cell(8, 0).merge(table.cell(8, 4))
    set_rubric_cell(comments, "Comments:", bold=True)

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            p = paragraph._element
            p.getparent().remove(p)

    file_name = (
        f"{summative.grade_ordinal} grade - {summative.trimester_code} - Week {summative.week_number} - "
        f"Rubric for Summative {summative.summative_number}.docx"
    )
    grade_folder = GRADES[summative.grade_number]["folder"] / "Rubrics"
    generated_folder = GENERATED_RUBRICS_ROOT / f"{summative.grade_ordinal} Grade Technology" / summative.trimester_label
    grade_folder.mkdir(parents=True, exist_ok=True)
    generated_folder.mkdir(parents=True, exist_ok=True)

    grade_path = grade_folder / file_name
    generated_path = generated_folder / file_name
    doc.save(grade_path)
    shutil.copy2(grade_path, generated_path)
    return grade_path, generated_path


def docx_xml_parts_parse(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.endswith(".xml"):
                try:
                    ET.fromstring(archive.read(name))
                except Exception as exc:  # pragma: no cover - CLI report
                    errors.append(f"{name}: {exc}")
    return errors


def file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def visible_rubric_font_errors(path: Path) -> list[str]:
    errors: list[str] = []
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    for idx, run in enumerate(root.findall(".//w:r", NS), start=1):
        text = "".join(node.text or "" for node in run.findall(".//w:t", NS)).strip()
        if not text:
            continue
        r_pr = run.find("w:rPr", NS)
        fonts = r_pr.find("w:rFonts", NS) if r_pr is not None else None
        size = r_pr.find("w:sz", NS) if r_pr is not None else None
        size_cs = r_pr.find("w:szCs", NS) if r_pr is not None else None
        ascii_font = fonts.get(f"{{{W_NS}}}ascii") if fonts is not None else None
        hansi_font = fonts.get(f"{{{W_NS}}}hAnsi") if fonts is not None else None
        size_value = size.get(f"{{{W_NS}}}val") if size is not None else None
        size_cs_value = size_cs.get(f"{{{W_NS}}}val") if size_cs is not None else None
        if ascii_font != "Arial" or hansi_font != "Arial" or size_value != "24" or size_cs_value != "24":
            errors.append(f"run {idx} ({text[:30]!r}) font={ascii_font}/{hansi_font} size={size_value}/{size_cs_value}")
    return errors


def write_report(monthly_paths: list[Path], rubric_pairs: list[tuple[Path, Path]]) -> None:
    lines = [
        "# 3rd and 4th Grade Generated Docs and Rubrics Report",
        "",
        f"Monthly DOCX files generated: {len(monthly_paths)}",
        f"Rubric DOCX files generated in grade folders: {len(rubric_pairs)}",
        f"Rubric DOCX files mirrored to Generated Rubrics 2026: {len(rubric_pairs)}",
        "",
        "## Monthly DOCX Files",
        "",
    ]
    for path in monthly_paths:
        lines.append(f"- `{path.relative_to(ROOT)}`")
    lines.extend(["", "## Rubric DOCX Files", ""])
    for grade_path, generated_path in rubric_pairs:
        match = "match" if file_hash(grade_path) == file_hash(generated_path) else "DIFFER"
        lines.append(f"- `{grade_path.relative_to(ROOT)}` -> `{generated_path.relative_to(ROOT)}` ({match})")

    xml_errors: list[str] = []
    font_errors: list[str] = []
    for path in monthly_paths:
        for error in docx_xml_parts_parse(path):
            xml_errors.append(f"{path.relative_to(ROOT)}: {error}")
    for grade_path, generated_path in rubric_pairs:
        for path in (grade_path, generated_path):
            for error in docx_xml_parts_parse(path):
                xml_errors.append(f"{path.relative_to(ROOT)}: {error}")
        for error in visible_rubric_font_errors(grade_path):
            font_errors.append(f"{grade_path.relative_to(ROOT)}: {error}")
        for error in visible_rubric_font_errors(generated_path):
            font_errors.append(f"{generated_path.relative_to(ROOT)}: {error}")

    lines.extend(["", "## Structural QA", ""])
    lines.append(f"- XML parse errors: {len(xml_errors)}")
    lines.append(f"- Rubric visible Arial 12 errors: {len(font_errors)}")
    if xml_errors:
        lines.extend(["", "### XML Errors", ""])
        lines.extend(f"- {error}" for error in xml_errors[:100])
    if font_errors:
        lines.extend(["", "### Font Errors", ""])
        lines.extend(f"- {error}" for error in font_errors[:100])

    report_path = GENERATED_RUBRICS_ROOT / "generation-report-3rd-4th.md"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"Report: {report_path}")


def main() -> int:
    args = parse_args()
    if args.monthly_only and args.rubrics_only:
        raise SystemExit("Choose either --monthly-only or --rubrics-only, not both.")

    monthly_paths: list[Path] = []
    rubric_pairs: list[tuple[Path, Path]] = []
    all_drafts: list[MonthlyDraft] = []

    for grade in args.grades:
        drafts = load_monthly_drafts(grade)
        all_drafts.extend(drafts)
        if not args.rubrics_only:
            for draft in drafts:
                monthly_paths.append(generate_monthly_docx(draft))

    if not args.monthly_only:
        summatives = collect_summatives(all_drafts)
        for summative in summatives:
            rubric_pairs.append(generate_rubric_docx(summative))

    write_report(monthly_paths, rubric_pairs)
    print(f"Monthly DOCX generated: {len(monthly_paths)}")
    print(f"Rubrics generated: {len(rubric_pairs)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
