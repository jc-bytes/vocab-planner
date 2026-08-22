from __future__ import annotations

from copy import deepcopy
from html import escape
import os
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[6]


def b(week, month, short, long, t45, o45, a45, t90, o90, a90):
    return dict(week=week, month=month, short=short, long=long,
                t45=t45, o45=o45, a45=a45, t90=t90, o90=o90, a90=a90)


G8 = [
 b(1,'September','8A Wed Sep 16 / 8B Mon Sep 14 — 45 minutes','8A and 8B Thu Sep 17 — 90 minutes',
   'App purpose and decomposition practice','Break a school app idea into clear tasks before the assessment.','Use MOD-APP-DESIGN-01 START-01 and USER-LEARN-01; create a five-step task list and one user need.',
   'App vocabulary and task-decomposition practice','Complete the fixed ten-word table and decompose the club-signup task.','Practice only; correct the table and task list before the formal app assessment.'),
 b(2,'September','8A Wed Sep 23 / 8B Mon Sep 21 — 45 minutes','8A and 8B Thu Sep 24 — 90 minutes',
   'Screen and event-map practice','Trace screen changes, inputs, outputs, and events.','Use USER-PRACTICE-01 and EVENTS-LEARN-01; correct one guided event map.',
   'D1 Club Signup App Screen and Event Map','Create the supplied three-screen plan and event map.','Formal Daily Grade 1; submit the named file or equal paper template.'),
 b(3,'September','8A Wed Sep 30 / 8B Mon Sep 28 — 45 minutes','8A and 8B Thu Oct 1 — 90 minutes',
   'A1 App Development Work Habits and File Responsibility','Organize the app file and show observable independent work habits.','Formal Appreciation Grade 1; use the supplied checklist, keep one correctly named file, follow the task order, and record one help request or independent fix.',
   'A2 App Testing, Revision, and Reflection','Run the three fixed tests, revise one item, and explain the change.','Formal Appreciation Grade 2; record expected, actual, pass/fail, revision, and reflection using the supplied cases.'),
 b(4,'October','8A Wed Oct 7 / 8B Mon Oct 5 — 45 minutes','8A and 8B Thu Oct 8 — 90 minutes',
   'STEAM preparation 1 — ungraded','Continue the single assigned group project with its responsible teacher.','Use the shared STEAM kit and group log; Technology does not create or score a second project.',
   'STEAM preparation 1 — ungraded','Complete the assigned group task and update the single group log.','No Technology grade; record group progress only.'),
 b(5,'October','8A Wed Oct 14 / 8B Mon Oct 12 — 45 minutes','8A and 8B Thu Oct 15 — 90 minutes',
   'STEAM preparation 2 — ungraded','Test or refine the single assigned group project.','Use the shared STEAM kit; keep the same group product and source of truth.',
   'D4 STEAM Project Preparation and Work Process','Complete one group test and improvement while showing individual Technology process evidence.','Formal Daily Grade 4: teacher observation plus each student’s role, dated contribution, test row, file/material routine, and short explanation. The project topic does not change the grade.'),
 b(6,'October','8A Wed Oct 21 / 8B Mon Oct 19 — 45 minutes','8A and 8B Thu Oct 22 — 90 minutes',
   'Official STEAM Week — ungraded','Prepare the assigned project explanation and setup.','Follow the responsible teacher and shared kit; no appreciation or daily grade.',
   'D3 STEAM Expo Participation and Closure','Present or support the assigned project and complete individual Technology closure evidence.','Formal Daily Grade 3: individual role/explanation, teacher question, cleanup, and reflection. The group keeps one shared product; the project subject does not change the Technology grade.'),
 b(7,'October','8A Wed Oct 28 / 8B Mon Oct 26 — 45 minutes','8A and 8B Thu Oct 29 — 90 minutes',
   'Python arithmetic practice','Enter, run, and correct arithmetic/output examples from the taught module.','Use MOD-PYTHON-FUNDAMENTALS-01 exact arithmetic/output sections and the supplied two-number starter.',
   'Python arithmetic and output practice','Complete the fixed arithmetic/output program and tests.','Practice only; save the .py file and test record for feedback before debugging assessment.'),
 b(8,'November','8A Wed Nov 4 — No class / 8B Mon Nov 2 — 45 minutes','8A and 8B Thu Nov 5 — No class',
   '8B Python debugging preparation / 8A closure','Prepare E1–E5 corrections without completing the graded file.','8B uses debug_five_errors.py; 8A receives the same preparation on Nov 9. Nov 4 is a closure.',
   'School closure','No class for either group.','Nov 5 closure; no replacement evidence.'),
 b(9,'November','8A Wed Nov 11 / 8B Mon Nov 9 — 45 minutes','8A and 8B Thu Nov 12 — 90 minutes',
   'Python debugging preparation and parity','Give both groups the same fixed-error preparation.','8A completes E1–E5 guided practice; 8B reviews the same five error types. No graded evidence.',
   'D2 Python Debugging: Find and Fix Five Errors','Correct exactly E1–E5 and prove the fixed program with three tests.','Formal Daily Grade 2; submit corrected file and error/test record.'),
 b(10,'November','8A Wed Nov 18 / 8B Mon Nov 16 — 45 minutes','8A and 8B Thu Nov 19 — 90 minutes',
   'Binary and representation practice','Use the 8-4-2-1 row to convert six values and explain representation.','Use MOD-DIGITAL-REPRESENTATION-01 taught sections and the fixed practice sheet.',
   'D5 Binary Numbers and Digital Representation','Complete the fixed vocabulary, conversions, and explanations.','Formal Daily Grade 5; submit the named response file or paper check.'),
 b(11,'November','8A Wed Nov 25 / 8B Mon Nov 23 — 45 minutes','8A and 8B Thu Nov 26 — 90 minutes',
   'Exam launch: Catch the Star','Open the supplied Scratch design and assign the fixed controls and variables.','Use the exam brief and supplied boundary; save the correctly named project and first log entry.',
   'Catch the Star build and test','Build the required game behavior and record the fixed tests.','Continue the exam project; collect milestone and test-log evidence, not a grade yet.'),
 b(12,'December','8A Wed Dec 2 / 8B Mon Nov 30 — 45 minutes','8A and 8B Thu Dec 3 — 90 minutes',
   'Catch the Star final checks','Run the supplied tests and rehearse the required explanation.','Finish the test record, one improvement, and the demo checklist.',
   'Exam — Catch the Star: Scratch + micro:bit Game','Demonstrate the game and submit the project evidence.','Formal exam; submit .sb3, log, tests, demo, and reflection.'),
 b(13,'December','8A Wed Dec 9 / 8B Mon Dec 7 — school exams or make-up','8A and 8B Thu Dec 10 — school exams or make-up',
   'School examinations and make-up only','Complete only an assigned equal-evidence make-up.','No new regular lesson or grade; Dec 8 closure does not affect Grade 8 timetable.',
   'School examinations and make-up only','Complete only an assigned equal-evidence make-up.','No new regular lesson or grade.'),
]

G9 = [
 b(1,'September','9A and 9B Mon Sep 14 — 45 minutes','9A Tue Sep 15 / 9B Wed Sep 16 — 90 minutes',
   'Digital representation vocabulary practice','Connect ten taught terms to fixed examples.','Use MOD-DIGITAL-REPRESENTATION-01 START-01 and IMAGE-LEARN-01; complete the guided match only.',
   'Digital representation vocabulary practice','Complete the fixed ten-term table with meanings and examples.','Practice only; correct the table before the image and sound assessment.'),
 b(2,'September','9A and 9B Mon Sep 21 — 45 minutes','9A Tue Sep 22 / 9B Wed Sep 23 — 90 minutes',
   'A1 Module Work Habits and Evidence Organization','Complete assigned module work and keep reconstructable evidence organized.','Formal Appreciation Grade 1; use the supplied checklist: task order, saved attempt, readable filename, correction, and export/printed record.',
   'D1 Image and Sound Data Calculations','Complete the fixed image and sound calculations using supplied values.','Formal Daily Grade 1; show units and working; submit the named file or paper check.'),
 b(3,'September','9A and 9B Mon Sep 28 — 45 minutes','9A Tue Sep 29 / 9B Wed Sep 30 — 90 minutes',
   'A2 Cybersecurity Work Process and Reflection','Follow the fixed scenario-analysis process and reflect on one correction.','Formal Appreciation Grade 2; use THREAT-LEARN-01 and the supplied source, revise one response, and write the fixed reflection.',
   'D2 Cybersecurity Threats and Protections','Identify the threat and best protection in the fixed scenarios.','Formal Daily Grade 2; submit the named response file or equal paper task.'),
 b(4,'October','9A and 9B Mon Oct 5 — 45 minutes','9A Tue Oct 6 / 9B Wed Oct 7 — 90 minutes',
   'STEAM preparation 1 — ungraded','Continue the single assigned group project with its responsible teacher.','Use the shared STEAM kit and single group log; Technology assigns no second product.',
   'STEAM preparation 1 — ungraded','Complete the assigned group task.','No Technology grade; record group progress only.'),
 b(5,'October','9A and 9B Mon Oct 12 — 45 minutes','9A Tue Oct 13 / 9B Wed Oct 14 — 90 minutes',
   'STEAM preparation 2 — ungraded','Test or refine the assigned group project.','Use the shared kit and one group source of truth.',
   'D4 STEAM Project Preparation and Work Process','Complete one group test and improvement while showing individual Technology process evidence.','Formal Daily Grade 4: teacher observation plus each student’s role, dated contribution, test row, file/material routine, and short explanation. The project topic does not change the grade.'),
 b(6,'October','9A and 9B Mon Oct 19 — 45 minutes','9A Tue Oct 20 / 9B Wed Oct 21 — 90 minutes',
   'Official STEAM Week — ungraded','Prepare the assigned project explanation and setup.','Follow the responsible teacher; no appreciation or daily grade.',
   'D3 STEAM Expo Participation and Closure','Present or support the assigned group project and complete individual Technology closure evidence.','Formal Daily Grade 3: individual role/explanation, teacher question, cleanup, and reflection. The group keeps one shared product; the project subject does not change the Technology grade.'),
 b(7,'October','9A and 9B Mon Oct 26 — 45 minutes','9A Tue Oct 27 / 9B Wed Oct 28 — 90 minutes',
   'Six-risk analysis preparation','Practice probability, impact, protection, and priority with three records.','Use MOD-CYBERSECURITY-RISK-01 RISK-LEARN-01, RISK-PRACTICE-01, and PROTECTION-LEARN-01.',
   'Six-risk analysis practice','Complete a guided six-risk table using the fixed source.','Correct ratings and justifications before D5; no formal grade.'),
 b(8,'November','9A and 9B Mon Nov 2 — 45 minutes','9A Tue Nov 3 / 9B Wed Nov 4 — No class',
   'D5 Six-Risk Cybersecurity Analysis','Rate R1–R6, choose protections, and justify the highest priority.','Formal Daily Grade 5; submit the fixed table and paragraph.',
   'School closures','No class for either group.','Nov 3 and Nov 4 closures; no replacement evidence.'),
 b(9,'November','9A and 9B Mon Nov 9 — 45 minutes','9A Tue Nov 10 — No class / 9B Wed Nov 11 — 90 minutes',
   'STEM proposal launch','Select the assigned option and begin the fixed proposal template.','Record problem, users, system goal, constraints, and three success criteria.',
   '9B supported proposal work / 9A closure','Keep both groups on the same required evidence.','9B may work and receive feedback; no extra required evidence. 9A receives equal support Nov 16. Nov 10 is a closure.'),
 b(10,'November','9A and 9B Mon Nov 16 — 45 minutes','9A Tue Nov 17 / 9B Wed Nov 18 — 90 minutes',
   'STEM system proposal milestone','Complete the assigned system proposal and three test criteria.','Ungraded exam milestone; teacher approves the fixed proposal before building.',
   'Exam launch: STEM Prototype','Begin the approved proposal, first version, and test log.','Use the exam brief; save the named product and first milestone.'),
 b(11,'November','9A and 9B Mon Nov 23 — 45 minutes','9A Tue Nov 24 / 9B Wed Nov 25 — 90 minutes',
   'STEM prototype final test and rehearsal','Complete controlled tests, one improvement, and demo notes.','Finish the fixed test record and reflection draft.',
   'Exam — STEM Prototype: Build, Test, Improve, and Demonstrate','Demonstrate the prototype and submit the complete evidence set.','Formal exam; collect product, log, tests, improvement, demo, and reflection.'),
 b(12,'December','9A and 9B — Graduand cutoff after Fri Nov 27','9A and 9B — Graduand cutoff after Fri Nov 27',
   'No regular Grade 9 class — graduand cutoff','No regular Technology class is planned.','Tentative pending AID confirmation; no new evidence or grade.',
   'No regular Grade 9 class — graduand cutoff','No regular Technology class is planned.','Tentative pending AID confirmation; no new evidence or grade.'),
 b(13,'December','9A and 9B — Graduand cutoff; Dec 8 school closure','9A and 9B — Graduand cutoff after Fri Nov 27',
   'No regular Grade 9 class — graduand cutoff','No regular Technology class is planned.','Dec 8 is also a school closure. Tentative pending AID confirmation.',
   'No regular Grade 9 class — graduand cutoff','No regular Technology class is planned.','No new evidence or grade. Tentative pending AID confirmation.'),
]


def grade_root(grade):
    return ROOT / 'plans' / f'{grade}th Grade Technology'


def note_path(grade, x, dur):
    monnum={'September':'09','October':'10','November':'11','December':'12'}[x['month']]
    folder=grade_root(grade)/'Materials'/'Class Notes'/f"T3 2026-{monnum} {x['month']}"
    return folder/f"{grade}th Grade Technology - T3 - 2026-{monnum} - {x['month']} - Week {x['week']} - {dur} minutes.md"


def write_note(grade, x, dur):
    title=x[f't{dur}']; obj=x[f'o{dur}']; activity=x[f'a{dur}']; label=x['short' if dur==45 else 'long']
    p=note_path(grade,x,dur); p.parent.mkdir(parents=True,exist_ok=True)
    assess='Ungraded class block'
    if title.startswith(('D1 ','D2 ','D3 ','D4 ','D5 ')): assess='Formal daily grade'
    elif title.startswith(('A1 ','A2 ')): assess='Formal regular Technology appreciation grade'
    elif title.startswith('Exam '): assess='Formal trimester exam'
    p.write_text(f'''---
title: "{grade}th Grade Technology — {label}"
trimester: 3
year: 2026
week: {x['week']}
schedule: "{label}"
duration: "{dur} minutes"
topic: "{title}"
---

# {grade}th Grade Technology — {label}

## Class Snapshot

| Item | Details |
|---|---|
| Scheduled group/date/duration | {label} |
| Topic | {title} |
| Assessment | {assess} |
| Evidence rule | Saved file, typed response, module export, teacher observation, or printed equivalent. No phone, photo, or screenshot evidence. |

## Objective

{obj}

## Before

- Project the matching anchor in `Class Visuals.html`.
- Open the linked module section, supplied source, starter, or task sheet named in the package index.
- Confirm the fallback before class. A documented school device, network, or login failure receives the equal fallback with no penalty.

## During

{activity}

## Check and collection

- Collect only the evidence named in the package index and student directions.
- Readiness/file responsibility may count for no more than 10% of a formal grade; a student's own working computer and login are valid readiness evidence.
- Do not deduct for a documented school equipment, network, or account failure when the student completes the equal fallback.

## After

- Verify the filename or student name on the paper fallback.
- Record missing evidence and assign the same-objective make-up.
''',encoding='utf-8')


def write_month(grade, month, items):
    draft=grade_root(grade)/'Planning'/'Drafts'/'3rd Trimester'/f'{grade}° Technology - {month}.md'
    notes=[]; rows=[]
    for x in items:
        for dur,labelkey in ((45,'short'),(90,'long')):
            p=note_path(grade,x,dur)
            rel=p.relative_to(ROOT/'plans').with_suffix('')
            notes.append(f"- {x[labelkey]}: [[{rel.as_posix()}|{x[f't{dur}']}]]")
            rows.append(f"| Week {x['week']} | {x[labelkey]} | {x[f't{dur}']} | {x[f'o{dur}']} | Open the package visual, source, module section, and fallback. | {x[f'a{dur}']} | Collect/check only the named evidence; use equal make-up if needed. |")
    cutoff='\n- Grade 9 graduand cutoff after Friday, November 27 is tentative pending AID confirmation. No regular Grade 9 Technology class is planned after the cutoff.\n' if grade==9 else ''
    steam='October 5–23 uses one assigned group project and one group source of truth. Daily Grades 3 and 4 use individual Technology process, testing, explanation, and reflection evidence regardless of project subject.' if month=='October' else 'Formal assessments occur only on the dates named below.'
    draft.write_text(f'''# {grade}° Technology - {month} 2026

## Monthly Focus

{steam}{cutoff}

## Weekly Time and Calendar

- Grade {grade} timetable is group-specific; each row names the exact group, date, and duration.
- Trimester 3 runs September 14–December 11. There are no planning rows after December 11.
- Closures reflected in the rows: November 3, 4, 5, 10, and December 8. November 28 is a Saturday and creates no class row.

## Device and Evidence Rules

- Students may use their own charged computer and login.
- No phone, photo, or screenshot evidence.
- Readiness/file responsibility is at most 10% of a formal grade.
- A documented school device, network, or login failure receives an equal offline/local fallback with no penalty.

## Evaluation

- Exactly five daily grades, two regular Technology appreciation grades, and one exam are scheduled across the trimester.
- Daily Grades 3 and 4 are individual Technology STEAM grades regardless of the assigned project subject.
- The group keeps one shared project/product and one group source of truth; Technology does not assign a second project.

## Class Notes (Generated)

{chr(10).join(notes)}

## Monthly Plan

| Week | Actual group/date/duration | Topic | Class Objective | Pre-Activities | While Activities | Post-Activities |
|---|---|---|---|---|---|---|
{chr(10).join(rows)}
''',encoding='utf-8')


def write_index(grade, blocks):
    package=grade_root(grade)/'Materials'/'Lesson Packages'/'T3 2026'
    rows=[]
    for x in blocks:
        for dur,key in ((45,'short'),(90,'long')):
            anchor=f"w{x['week']}-{dur}"
            note=note_path(grade,x,dur)
            relnote=Path(os.path.relpath(note, package)).as_posix()
            title=x[f't{dur}']; evidence='No new graded evidence'
            if title.startswith(('D1 ','D2 ','D3 ','D4 ','D5 ')): evidence='Named daily-grade file/task and rubric'
            elif title.startswith(('A1 ','A2 ')): evidence='Named checklist/reflection plus teacher observation'
            elif title.startswith('Exam '): evidence='Final product, log, tests, improvement, demo, reflection'
            rows.append(f"| {x[key]} | [{title}]({escape(relnote).replace(' ','%20')}) | [Visual](Class%20Visuals.html#{anchor}) | {x[f'a{dur}']} | {evidence}; equal printed/local fallback |")
    cutoff='\n> Grade 9 assumption: the graduand cutoff after Friday, November 27 is tentative and pending AID confirmation. The exam is therefore completed November 24/25.\n' if grade==9 else ''
    mapping=('8A continues its assigned Science project. 8B uses the fixed Technology list: Native Panamanian Language Platform, Obstacle-Detecting Walking Stick, Biomimetic Dexterous Hand, or Wind Turbine Generator.' if grade==8 else '9A continues its assigned Arts project and 9B continues its assigned Science project.')
    package.joinpath('00 Trimester Materials Index.md').write_text(f'''# Grade {grade} Technology — Trimester 3 2026 Materials Index

Trimester window: September 14–December 11, 2026. This map is the source of truth for the group-specific calendar.{cutoff}

## Assessment schedule

{'| Label | Official assessment | Group date(s) |' if grade==8 else '| Label | Official assessment | Group date(s) |'}
|---|---|---|
{assessment_rows(grade)}

Use the [Shared T3 2026 STEAM Kit](../../../../Shared%20Materials/T3%202026%20STEAM%20Kit/) as the source of truth. {mapping} Students receive individual Technology Daily Grades 3 and 4 from observable process and evidence regardless of project subject. The group keeps one shared product and log; Technology creates no second project.

## Class-block map

| Actual group/date/duration | Teacher guide and title | Projected visual | Student task | Evidence, check, make-up, offline |
|---|---|---|---|---|
{chr(10).join(rows)}

## Evidence and readiness rule

Students may use their own charged computer and login. Readiness/file responsibility may contribute no more than 10% of a formal grade. A documented school device, network, or account failure is not penalized when the student completes the equal local or printed fallback. No phone, photo, or screenshot evidence is accepted or required.
''',encoding='utf-8')


def assessment_rows(grade):
    if grade==8:
        vals=[('D1','Club Signup App Screen and Event Map','8A/8B Sep 24'),('A1','App Development Work Habits and File Responsibility','8A Sep 30; 8B Sep 28'),('A2','App Testing, Revision, and Reflection','8A/8B Oct 1'),('D4','STEAM Project Preparation and Work Process','8A/8B Oct 15'),('D3','STEAM Expo Participation and Closure','8A/8B Oct 22'),('D2','Python Debugging: Find and Fix Five Errors','8A/8B Nov 12'),('D5','Binary Numbers and Digital Representation','8A/8B Nov 19'),('Exam','Catch the Star: Scratch + micro:bit Game','8A/8B Dec 3')]
    else:
        vals=[('A1','Module Work Habits and Evidence Organization','9A/9B Sep 21'),('D1','Image and Sound Data Calculations','9A Sep 22; 9B Sep 23'),('A2','Cybersecurity Work Process and Reflection','9A/9B Sep 28'),('D2','Cybersecurity Threats and Protections','9A Sep 29; 9B Sep 30'),('D4','STEAM Project Preparation and Work Process','9A Oct 13; 9B Oct 14'),('D3','STEAM Expo Participation and Closure','9A Oct 20; 9B Oct 21'),('D5','Six-Risk Cybersecurity Analysis','9A/9B Nov 2'),('Exam','STEM Prototype: Build, Test, Improve, and Demonstrate','9A Nov 24; 9B Nov 25')]
    return '\n'.join(f'| {a} | {n} | {d} |' for a,n,d in vals)


def write_visuals(grade, blocks):
    cards=[]
    for x in blocks:
        for dur,key in ((45,'short'),(90,'long')):
            cards.append(f'''<section id="w{x['week']}-{dur}"><p class="eyebrow">{escape(x[key])}</p><h2>{escape(x[f't{dur}'])}</h2><p>{escape(x[f'o{dur}'])}</p><ol><li>Open the named source, module section, or starter.</li><li>{escape(x[f'a{dur}'])}</li><li>Save or hand in only the named evidence. No phone, photo, or screenshot.</li></ol></section>''')
    p=grade_root(grade)/'Materials'/'Lesson Packages'/'T3 2026'/'Class Visuals.html'
    p.write_text(f'''<!doctype html><html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Grade {grade} T3 Class Visuals</title><style>body{{font:20px/1.5 system-ui;margin:0;background:#f7f4ea;color:#123}}header,section{{max-width:1000px;margin:auto;padding:3rem}}header{{background:#123d55;color:white;max-width:none}}section{{min-height:70vh;border-bottom:4px solid #e3a008}}h1{{font-size:3rem}}h2{{font-size:2.4rem}}.eyebrow{{font-weight:800;color:#087f6b}}li{{margin:.7rem 0}}@media print{{section{{break-after:page}}}}</style></head><body><header><h1>Grade {grade} Technology — T3 2026</h1><p>Use the URL anchor shown in the package index. Each screen names the real group/date/duration.</p></header>{''.join(cards)}</body></html>''',encoding='utf-8')


def update_docx(grade, month, items):
    path=grade_root(grade)/'Planning'/'Monthly'/'3rd Trimester'/f'{grade}° Technology - {month}.docx'
    d=Document(path); table=d.tables[1]
    proto=[deepcopy(r._tr) for r in table.rows[:3]]
    for r in list(table._tbl.tr_lst): table._tbl.remove(r)
    for x in items:
        for tr in proto: table._tbl.append(deepcopy(tr))
        rows=table.rows[-3:]
        rows[0].cells[0].text=x['short']; rows[0].cells[1].text=x['long']
        for col,dur in ((0,45),(1,90)):
            rows[1].cells[col].text=f"Topic:\n{x[f't{dur}']}\nClass Objective:\n{x[f'o{dur}']}\nPre-Activities:\nOpen the package visual, named source/module section, and fallback.\nWhile Activities:\n{x[f'a{dur}']}\nPost-Activities:\nCollect/check only the named evidence; use the equal make-up if needed."
            rows[2].cells[col].text="Resources:\nStudent computer or school fallback, projected visual, named module/source/starter, printed equivalent. No phone, photo, or screenshot evidence."
    if month == 'October':
        replacements = ([
            ('Complete Appreciation Grades #1 and #2 through teacher observation of planning responsibility and expo participation.',
             'Complete Daily Grade 4 through preparation-process evidence and Daily Grade 3 through expo participation and closure.'),
            ('Submit the Python arithmetic program and debugging check after the expo.',
             'Complete the Python arithmetic program as practice, then submit the Daily Grade 2 debugging check.'),
        ] if grade == 8 else [
            ('Complete setup, the Thursday, October 22 expo, cleanup, and reflection during October 19-23.',
             'Complete setup, the October 20-21 expo, cleanup, and reflection during October 19-23.'),
            ('Complete Appreciation Grade #1 for the October 5-16 work process and Appreciation Grade #2 for expo participation and closure.',
             'Complete Daily Grade 4 for preparation work process and Daily Grade 3 for expo participation and closure.'),
            ('Submit the cybersecurity scenario quiz and fixed six-risk map after the expo.',
             'Keep the completed Daily Grade 2 scenario check and prepare the fixed Daily Grade 5 six-risk map.'),
        ])
        for row in d.tables[0].rows:
            for cell in row.cells:
                for old, new in replacements:
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)
    d.save(path)


for grade,blocks in ((8,G8),(9,G9)):
    for x in blocks:
        write_note(grade,x,45); write_note(grade,x,90)
    for month in ('September','October','November','December'):
        items=[x for x in blocks if x['month']==month]
        write_month(grade,month,items); update_docx(grade,month,items)
    write_index(grade,blocks); write_visuals(grade,blocks)

print('date-aware Grade 8 and Grade 9 plans rebuilt')
