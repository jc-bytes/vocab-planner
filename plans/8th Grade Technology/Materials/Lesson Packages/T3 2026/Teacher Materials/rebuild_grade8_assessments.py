from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

ROOT=Path(__file__).resolve().parents[6]
G8=ROOT/'plans/8th Grade Technology'
DOC=G8/'Assessments/Rubrics/IIIT'
MIR=ROOT/'plans/Shared/Generated Outputs/Rubrics 2026/8th Grade Technology/3rd Trimester'

def c(title, desc, levels, points=(9,6,3,0)):
    return {'title':title,'description':desc,'points':list(points),'levels':['Exceeds','Meets','Developing','Not demonstrated'],'levelDescriptions':levels}

specs=[
('Daily','8th grade - IIIT - 2026-09-17 - Daily Grade 1 - App Vocabulary and Task Decomposition',40,[
c('Required entries','Completeness of the fixed ten-term table.', ['All 10 terms each have a definition and example.','8–9 terms each have both entries.','5–7 terms each have both entries.','0–4 terms have both entries.']),
c('Definition accuracy','Accuracy of the ten definitions against the supplied reference.', ['9–10 definitions are accurate.','7–8 are accurate.','4–6 are accurate.','0–3 are accurate.']),
c('Example accuracy','How clearly each example shows the named app concept.', ['9–10 examples fit the term.','7–8 fit the term.','4–6 fit the term.','0–3 fit the term.']),
c('Decomposition','Five ordered, distinct action steps for the fixed homework-app goal.', ['5 distinct ordered action steps cover the full goal.','4 distinct steps cover most of the goal.','2–3 usable steps are present.','0–1 usable step is present.']),
c('Checkable evidence','The named submission contains the fixed table and breakdown.', ['Both parts are complete and checkable.','Both parts are present; one small field is unclear.','Only one part is checkable or several fields are unclear.','No checkable submission evidence.'],(4,3,2,0))]),
('Daily','8th grade - IIIT - 2026-09-24 - Daily Grade 2 - Club Signup App Screen and Event Map',40,[
c('Screens and elements','Use of the three supplied screens and five supplied elements.', ['All 3 screens and all 5 elements are correctly placed.','All screens and 4 elements are correctly placed.','2 screens or 2–3 elements are usable.','Fewer than 2 screens or 2 elements are usable.']),
c('Events and navigation','Four event-to-action pairs and directional connections.', ['4 correct pairs and all navigation arrows are usable.','3 correct pairs and most arrows are usable.','2 correct pairs or unclear arrows.','0–1 correct pair.']),
c('Stored data','Use of the supplied selected_club value.', ['Value is captured, stored, and used in confirmation.','Value is captured and stored but its final use is partly unclear.','Value is named but flow is incomplete.','Value is absent or unrelated.']),
c('Fixed tests','Expected results for the three supplied test inputs.', ['All 3 expected results are complete and consistent with the map.','2 are complete and consistent.','1 is complete and consistent.','No consistent test result.']),
c('Checkable evidence','The named map opens or the printed template is complete.', ['All required sections are checkable.','One small field is unclear.','Several fields are unclear or one section is missing.','No checkable map.'],(4,3,2,0))]),
('Daily','8th grade - IIIT - 2026-10-29 - Daily Grade 3 - Python Arithmetic and Output',40,[
c('Inputs and variables','Two supplied inputs are converted with int and stored clearly.', ['Both inputs use int and meaningful variables correctly.','Both inputs work; one name or conversion detail is weak.','Only one input is converted and stored correctly.','No working numeric input.']),
c('Calculations','Required sum, difference, and product.', ['All 3 calculations use the required operators correctly.','2 are correct.','1 is correct.','No required calculation is correct.']),
c('Labeled output','Three required results are printed with clear labels.', ['All 3 labels and values are correct.','2 labels and values are correct.','1 label and value is correct.','No required labeled result.']),
c('Fixed tests','Recorded results for 8/3, 0/5, and 12/4.', ['All 3 tests show exact actual output and correct pass/fail.','2 tests are complete and accurate.','1 test is complete and accurate.','No complete accurate test.']),
c('Checkable evidence','Named source file and three-row test record.', ['Both open and match.','Both are present; one small field is unclear.','Only one item is checkable.','No checkable evidence.'],(4,3,2,0))]),
('Daily','8th grade - IIIT - 2026-11-12 - Daily Grade 4 - Python Debugging Find and Fix Five Errors',40,[
c('Error identification','Identification of the five supplied practiced errors E1–E5.', ['All 5 causes are identified accurately.','4 causes are accurate.','2–3 causes are accurate.','0–1 cause is accurate.']),
c('Corrections','Exact corrected lines for E1–E5.', ['All 5 corrections are accurate.','4 are accurate.','2–3 are accurate.','0–1 is accurate.']),
c('Explanations','Why each correction fixes the named problem.', ['All 5 explanations connect cause to correction.','4 do.','2–3 do.','0–1 does.']),
c('Fixed tests','Runs for 6/3, 4/0, and 5/2 with required status behavior.', ['All 3 tests are recorded and correct.','2 are recorded and correct.','1 is recorded and correct.','No correct recorded test.']),
c('Checkable evidence','Named corrected file or paper trace plus fix/test record.', ['Both parts are complete and checkable.','Both are present; one small field is unclear.','Only one part is checkable.','No checkable evidence.'],(4,3,2,0))]),
('Daily','8th grade - IIIT - 2026-11-19 - Daily Grade 5 - Binary Numbers and Digital Representation',40,[
c('Vocabulary','Matching of six supplied representation terms.', ['All 6 matches are accurate.','5 are accurate.','3–4 are accurate.','0–2 are accurate.']),
c('Decimal to binary','Four-bit conversions for 2, 6, and 10.', ['All 3 conversions and place values are correct.','2 are correct.','1 is correct.','No correct conversion.']),
c('Binary to decimal','Conversions for 0100, 1011, and 1111.', ['All 3 conversions and place values are correct.','2 are correct.','1 is correct.','No correct conversion.']),
c('Representation explanations','Two fixed explanations about shared rules and leading zeros.', ['Both explanations are accurate and use the supplied examples.','Both are mostly accurate; one detail is unclear.','One explanation is accurate.','Neither explanation is accurate.']),
c('Checkable evidence','Fixed check shows the 8,4,2,1 work and responses.', ['All sections are readable and checkable.','One small field is unclear.','Several fields are unclear or one section is missing.','No checkable evidence.'],(4,3,2,0))]),
('Appreciation','8th grade - IIIT - 2026-09-28 and 2026-09-30 - Appreciation Grade 1 - App Development Work Habits and File Responsibility',40,[
c('Task order','Uses the supplied app-development steps in order.', ['All steps are completed in order without prompting.','All steps are completed with one prompt.','Some steps are skipped or need repeated prompts.','No usable task sequence is shown.']),
c('Independent correction or help request','Records one specific correction or clear help request.', ['Record names the issue, action, and result.','Record names the issue and action.','Record is vague or incomplete.','No usable record.']),
c('File organization','Opens, names, saves, and locates the correct app file.', ['All four actions are correct and independent.','Three actions are correct.','One or two actions are correct.','No checkable app file or equal fallback.']),
c('Assigned app edit','Completes the supplied screen or event edit.', ['Edit is complete, accurate, and checkable.','Edit works with one minor issue.','Edit is partly complete.','No usable edit.']),
c('Readiness and file responsibility','4/40: own computer/login or school fallback, on-time start, and saved evidence.', ['Ready and evidence saved independently.','One issue corrected after one reminder.','Repeated reminders or partial fallback evidence.','No readiness evidence and no fallback completed.'],(4,3,2,0))]),
('Appreciation','8th grade - IIIT - 2026-10-01 - Appreciation Grade 2 - App Testing Revision and Reflection',40,[
c('Three fixed tests','Expected, actual, and pass/fail results for all three supplied cases.', ['All 3 records are complete and accurate.','2 are complete and accurate.','1 is complete and accurate.','No accurate complete test.']),
c('Evidence-based revision','One screen, event, label, or arrow is revised from a test result.', ['Revision is specific and directly fixes the evidence.','Revision fits the evidence with one unclear detail.','Revision is present but weakly connected.','No usable revision.']),
c('Change explanation','States exactly what changed and why it improves the app.', ['Both parts are specific and accurate.','Both are present; one is general.','Only one part is usable.','No usable explanation.']),
c('Fixed reflection','Completes all four supplied reflection sentences.', ['All 4 are specific and supported by the work.','3 are complete and specific.','1–2 are usable.','No usable reflection.']),
c('Readiness and file responsibility','4/40: own computer/login or school fallback, on-time start, and saved evidence.', ['Ready and evidence saved independently.','One issue corrected after one reminder.','Repeated reminders or partial fallback evidence.','No readiness evidence and no fallback completed.'],(4,3,2,0))]),
('Exam Projects','8th grade - IIIT - 2026-12-03 - Exam - Catch the Star Scratch + microbit Game',90,[
c('Supplied game design','Catch the Star sprites, starts, goal, controls, score, and obstacle.', ['All 8 supplied elements are implemented accurately.','6–7 are accurate.','3–5 are accurate.','0–2 are accurate.'],(15,10,5,0)),
c('Input mapping','A/left and B/right mappings and labeled hardware fallback.', ['Both directions work and mapping is explained accurately.','Both work; explanation has one gap.','One direction works or mapping is incomplete.','No working required input.'],(15,10,5,0)),
c('Scratch logic','Reset, movement, forever loop, conditions, score, and feedback behaviors.', ['All 6 required behaviors work together.','4–5 work.','2–3 work.','0–1 works.'],(15,10,5,0)),
c('Testing and debugging','Five fixed tests, peer feedback, one correction, and same-condition retest.', ['All 5 tests plus feedback, correction, and retest are complete.','4 tests plus correction/retest are complete.','2–3 tests or partial correction evidence.','0–1 test and no usable correction.'],(15,10,5,0)),
c('Documentation and reflection','Build log, speaking notes, and five reflection responses.', ['All three records are complete and specific.','All are present; one has gaps.','Only one or two are usable.','No usable documentation.'],(15,10,5,0)),
c('Live demonstration','90–120 second demo covers goal, controls, logic, test correction, and future step.', ['All 5 points are accurate within time and project opens.','4 points are accurate or time is slightly outside range.','2–3 points are usable.','0–1 point is usable or no demo.'],(15,10,5,0))]),
]

# Active T3 structure: five daily grades, two regular Technology appreciation
# grades, and one exam. STEAM preparation/expo are Daily Grades 4 and 3.
by_name={name:(folder,name,total,criteria) for folder,name,total,criteria in specs}
app=next(x for x in specs if 'Club Signup App' in x[1])
debug=next(x for x in specs if 'Python Debugging Find and Fix' in x[1])
binary=next(x for x in specs if 'Binary Numbers and Digital Representation' in x[1])
appreciation1=next(x for x in specs if 'App Development Work Habits and File Responsibility' in x[1])
appreciation2=next(x for x in specs if 'App Testing Revision and Reflection' in x[1])
exam=next(x for x in specs if x[0]=='Exam Projects')
app=('Daily','8th grade - IIIT - 2026-09-24 - Daily Grade 1 - Club Signup App Screen and Event Map',40,app[3])
debug=('Daily','8th grade - IIIT - 2026-11-12 - Daily Grade 2 - Python Debugging Find and Fix Five Errors',40,debug[3])
binary=('Daily','8th grade - IIIT - 2026-11-19 - Daily Grade 5 - Binary Numbers and Digital Representation',40,binary[3])
steam1=('Daily','8th grade - IIIT - 2026-10-15 - Daily Grade 4 - STEAM Project Preparation and Work Process',40,[
c('Role and first task','Assigned project subject, role, first task, and success rule in the shared log.', ['All 4 fields are accurate and specific.','3 fields are accurate.','1–2 fields are usable.','No usable individual role record.']),
c('Dated individual contribution','One preparation contribution names action, result, and next step.', ['All 3 parts are specific and teacher-observed.','All parts are present; one is general.','Only 1–2 parts are usable.','No usable contribution evidence.']),
c('Test and improvement','One test row records condition, expected result, actual result, and next improvement.', ['All 4 fields are accurate and connected.','3 fields are accurate.','1–2 fields are usable.','No usable test evidence.']),
c('Safe file and material routine','Observable responsible handling, file organization, and respectful collaboration.', ['All 3 routines are independent.','All are shown with one reminder.','Repeated reminders are needed.','No safe/responsible routine is demonstrated.']),
c('Readiness and individual evidence','4/40: own computer/login or equal fallback, on-time start, identifiable evidence.', ['Ready and evidence is identifiable.','One issue corrected after one reminder.','Partial fallback or repeated reminders.','No identifiable evidence or fallback.'],(4,3,2,0))])
steam2=('Daily','8th grade - IIIT - 2026-10-22 - Daily Grade 3 - STEAM Expo Participation and Closure',40,[
c('Assigned expo role','Completes the individual speaking, demonstration, setup, or support role.', ['Role is complete and independent.','Role is complete with one prompt.','Role is partly complete.','No demonstrated role.']),
c('Five-point explanation','Explains problem, goal, own contribution, one test result, and one improvement.', ['All 5 points are accurate.','4 are accurate.','2–3 are usable.','0–1 is usable.']),
c('Evidence response','Answers one teacher question using the shared product or log.', ['Answer cites specific evidence.','Answer is accurate but general.','Answer is partly accurate.','No usable answer.']),
c('Closure and reflection','Completes cleanup/file return and the fixed individual reflection.', ['Both are complete and specific.','Both are complete; one detail is general.','Only one is usable.','Neither is usable.']),
c('Readiness and individual evidence','4/40: own materials/login or equal fallback, on-time role, identifiable evidence.', ['Ready and evidence is identifiable.','One issue corrected after one reminder.','Partial fallback or repeated reminders.','No identifiable evidence or fallback.'],(4,3,2,0))])
specs=[app,debug,steam2,steam1,binary,appreciation1,appreciation2,exam]

RESPONSIBILITY_40=c(
    'Punctuality, responsibility, and readiness',
    '4/40: charged computer and login; on-time start; responsible use; deadline. School-controlled failures use the fallback without deduction.',
    [
        'Device and login ready; on time; responsible use; deadline met.',
        'One issue corrected after one reminder; work remains checkable.',
        'Device or login missing, late or unprepared, or repeated reminders; some fallback work.',
        'No device or login and no fallback, no checkable work, or responsible-use rules refused.'
    ],(4,3,2,0))
RESPONSIBILITY_90=c(
    'Punctuality, responsibility, and readiness',
    '9 of 90 points: checkpoint readiness, punctual starts, responsible file and material care, milestones, and final handoff. Documented school-controlled failures use the fallback without deduction.',
    [
        'Device, login, and materials are ready at every checkpoint; starts and milestones are on time; files are organized; final evidence is submitted.',
        'One readiness, punctuality, or milestone issue is corrected after one reminder; final evidence remains checkable.',
        'Readiness or milestone issues repeat, or one session lacks the device or login; some fallback project work is completed.',
        'Readiness is not shown across checkpoints and the approved fallback or final evidence is not completed.'
    ],(9,7,4,0))

for folder, name, total, criteria in specs:
    if total == 40:
        criteria[-1]=RESPONSIBILITY_40
    else:
        for criterion, maximum in zip(criteria, (14,14,14,13,13,13)):
            criterion['points']=[maximum, 10 if maximum == 14 else 9, 5, 0]
        criteria.append(RESPONSIBILITY_90)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def build_rubric(folder,name,total,criteria):
    d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.55); sec.left_margin=sec.right_margin=Inches(.55)
    sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width,sec.page_height=Inches(11),Inches(8.5)
    styles=d.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(12)
    display=name
    dates=''
    if 'App Vocabulary' in name: dates='Scheduled: 8A and 8B — September 17, 2026'
    elif 'App Screen' in name: dates='Scheduled: 8A and 8B — September 24, 2026'
    elif 'App Development Work Habits' in name: display='Appreciation Grade 1 — App Development Work Habits and File Responsibility'; dates='Scheduled: 8A September 30; 8B September 28, 2026'
    elif 'App Testing Revision' in name: display='Appreciation Grade 2 — App Testing, Revision, and Reflection'; dates='Scheduled: 8A and 8B — October 1, 2026'
    elif 'Daily Grade 4 - STEAM Project Preparation' in name: display='Daily Grade 4 — STEAM Project Preparation and Work Process'; dates='Scheduled: 8A and 8B — October 15, 2026'
    elif 'Daily Grade 3 - STEAM Expo' in name: display='Daily Grade 3 — STEAM Expo Participation and Closure'; dates='Scheduled: 8A and 8B — October 22, 2026'
    elif 'Arithmetic' in name: dates='Scheduled: 8A and 8B — October 29, 2026'
    elif 'Python Debugging' in name: display='Daily Grade 2 — Python Debugging: Find and Fix Five Errors'; dates='Scheduled: 8A and 8B — November 12, 2026'
    elif 'Digital Representation' in name: display='Daily Grade 5 — Binary Numbers and Digital Representation'; dates='Scheduled: 8A and 8B — November 19, 2026'
    elif folder == 'Exam Projects': display='Exam — Catch the Star: Scratch + micro:bit Game'; dates='Scheduled: 8A and 8B — December 3, 2026'
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(display); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x0B,0x3B,0x55)
    p=d.add_paragraph(f'Grade 8 Technology · Third Trimester 2026 · Total: {total} points'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if dates:
        p=d.add_paragraph(dates); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph('Student: ______________________________  Class: ______  Date: __________')
    if folder == 'Exam Projects':
        d.add_heading('Fixed project: Catch the Star',1)
        d.add_paragraph('Build the supplied Scratch game. Player starts at (0, -120), Star at (0, 80), Obstacle at (120, -120), and score at 0. A/left moves x by -20; B/right moves x by 20. Star contact adds 1, gives feedback, and relocates Star. Obstacle contact says “Try again,” resets score, and returns Player to start. Green flag resets all supplied values.')
        d.add_paragraph('Submit: 8_Class_LastName_Exam_CatchTheStar.sb3, build/test log, peer-feedback sheet, reflection, and 90–120 second live demo. Fixed tests: reset; left; right; Star contact; Obstacle contact. Keyboard or simulator input receives equal credit when hardware is unavailable. Decoration and advanced blocks are not scored.')
        d.add_paragraph('Allowed resources: supplied package, assigned module sections, class demonstrations, Scratch help, and teacher feedback. Offline route: local .sb3 or paper block plan and five traces followed by a scheduled live check.')
        d.add_page_break()
    t=d.add_table(rows=1, cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=['Criterion','Highest','Meets','Developing','Not demonstrated']
    for i,x in enumerate(hdr): t.rows[0].cells[i].text=x; shade(t.rows[0].cells[i],'0B3B55')
    for c0 in t.rows[0].cells:
        for run in c0.paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for cr in criteria:
        row=t.add_row().cells; row[0].text=f"{cr['title']}\n{cr['description']}"
        row[0]._tc.getparent().get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        for j in range(4): row[j+1].text=f"{cr['points'][j]} points\n{cr['levelDescriptions'][j]}"
        for c0 in row: c0.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    d.add_paragraph('Score: ______ / '+str(total)+'    Teacher evidence note: ______________________________________________')
    d.add_paragraph("A forgotten device or login may lower only the readiness criterion. A documented school network, platform, teacher-provided account, hardware, or internet failure outside the student's control does not lower the score when the stated fallback is completed.")
    for paragraph in d.paragraphs:
        for run in paragraph.runs: run.font.name='Arial'; run.font.size=Pt(12)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs: run.font.name='Arial'; run.font.size=Pt(12)
    out=DOC/folder/(name+'.docx'); out.parent.mkdir(parents=True,exist_ok=True); d.save(out)
    mout=MIR/folder/(name+'.docx'); mout.parent.mkdir(parents=True,exist_ok=True); copy2(out,mout)

for spec in specs: build_rubric(*spec)

def assessment(path,title,paras,tables):
    d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.6); sec.left_margin=sec.right_margin=Inches(.65)
    d.styles['Normal'].font.name='Arial'; d.styles['Normal'].font.size=Pt(10)
    h=d.add_heading(title,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph('Student: ______________________________  Class: ______  Date: __________')
    for p in paras: d.add_paragraph(p)
    for heading,headers,rows in tables:
        d.add_heading(heading,1); t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
        for i,x in enumerate(headers): t.rows[0].cells[i].text=x; shade(t.rows[0].cells[i],'D8EAF0')
        for vals in rows:
            c0=t.add_row().cells
            for i,x in enumerate(vals): c0[i].text=str(x)
    path.parent.mkdir(parents=True,exist_ok=True); d.save(path)

daily=DOC/'Daily'
debug_rows=[('E1','Missing parenthesis','print("Order total")','Completes the print call.'),('E2','Misspelled variable','quantity = int(input("Quantity: "))','Uses one consistent variable name.'),('E3','Missing integer conversion','price = int(input("Price: "))','Makes price numeric.'),('E4','Wrong operator','total = price * quantity','Multiplies price by quantity.'),('E5','Missing indentation','    print("Order recorded")','Places the status line inside the supplied if.')]
assessment(daily/'8th grade - IIIT - Week 8 - Summative 4 - Python Selection and Loop Debugging Check.docx','Daily Grade 2: Python debugging check',['Use the supplied debug_five_errors.py. Correct only E1–E5. Record the cause, exact corrected line, and reason.','Fixed tests: 6 and 3 → 18 plus status; 4 and 0 → 0 without status; 5 and 2 → 10 plus status.'],[('Five corrections',['Error','Cause','Exact corrected line','Why it works'],[(x[0],'','','') for x in debug_rows]),('Test record',['Inputs','Expected','Actual','Pass/fail'],[('6, 3','18; Order recorded','',''),('4, 0','0; no status','',''),('5, 2','10; Order recorded','','')])])
assessment(daily/'8th grade - IIIT - Week 8 - Summative 4 - Python Selection and Loop Debugging Check - Teacher Key.docx','Teacher key: Daily Grade 2 Python debugging',['Expected corrections and fixed tests. Accept equivalent quotation style and spacing. Do not score loops or student-created selection.'],[('Corrections',['Error','Cause','Correct line','Reason'],debug_rows),('Tests',['Inputs','Expected total','Expected status'],[('6, 3','18','Order recorded'),('4, 0','0','No status line'),('5, 2','10','Order recorded')])])
binary_rows=[('representation','an agreed way to stand for information'),('bit','one binary digit, 0 or 1'),('binary','a base-two number system'),('decimal','the base-ten number system'),('place value','the value a position contributes'),('conversion','changing the form without changing the value')]
assessment(daily/'8th grade - IIIT - Week 9 - Summative 5 - Representation and Binary Check.docx','Daily Grade 5: binary numbers and digital representation',['Use the 8, 4, 2, 1 row. Only the content on this page is assessed.'],[('Vocabulary',['Term','Matching definition letter'],[(x[0],'') for x in binary_rows]),('Conversions',['Question','8','4','2','1','Answer'],[(x,'','','','','') for x in ['2 to binary','6 to binary','10 to binary','0100 to decimal','1011 to decimal','1111 to decimal']]),('Explanations',['Question','Response'],[('Why does a bit pattern need an agreed representation rule?',''),('Why do leading zeros not change the value of 0011?','')])])
assessment(daily/'8th grade - IIIT - Week 9 - Summative 5 - Representation and Binary Check - Teacher Key.docx','Teacher key: Daily Grade 5 representation and binary',['Expected answers: 2=0010, 6=0110, 10=1010, 0100=4, 1011=11, 1111=15.','An agreed rule is needed because the same bits can stand for different kinds of information. Leading zeros add zero in unused place values, so 0011 still equals 3.'],[('Vocabulary',['Term','Expected meaning'],binary_rows)])

spec_json=G8/'Materials/Lesson Packages/T3 2026/Teacher Materials/grade8_classroom_rubric_specs.json'
spec_json.write_text(json.dumps([{'basename':n+' - Google Classroom Rubric.xlsx','criteria':cr} for _,n,_,cr in specs],indent=2),encoding='utf-8')
print('built',len(specs),'rubrics and four checks')
