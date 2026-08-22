#!/usr/bin/env python3
import json, shutil
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[6]
PKG = Path(__file__).resolve().parent.parent
SPEC = json.loads((Path(__file__).with_name("grade7-assessment-spec.json")).read_text())
CANON = ROOT / "plans/7th Grade Technology/Assessments/Rubrics/IIIT"
MIRROR = ROOT / "plans/Shared/Generated Outputs/Rubrics 2026/7th Grade Technology/3rd Trimester"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=12):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text)); r.bold = bold; r.font.name = "Arial"; r.font.size = Pt(12)
    if color: r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def build(item, out):
    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(.42); sec.left_margin = sec.right_margin = Inches(.45)
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.space_before = Pt(0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("ACADEMIA INTERNACIONAL DAVID\nTECHNOLOGY · GRADE 7 · TRIMESTER 3"); r.bold=True; r.font.name="Arial"; r.font.size=Pt(12); r.font.color.rgb=RGBColor(11,71,96)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(item["title"]); r.bold=True; r.font.name="Arial"; r.font.size=Pt(12)
    p=doc.add_paragraph("Name: ______________________________  Class: 7A / 7B  Date: __________  Score: ____ / %d" % item["max"])
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    note=doc.add_paragraph("Use the announced fixed directions and supplied assets. Readiness is 4 of 40 points or 9 of 90 points. A forgotten device or login may lower only that criterion. A documented school network, platform, teacher-provided account, connection, or hardware failure outside the student's control does not lower the score when the approved fallback is completed.")
    note.paragraph_format.space_after=Pt(5)
    levels=["Complete evidence","Minor gap","Partial evidence","Not demonstrated"]
    widths=[1.6,2.12,2.12,2.12,2.12]
    def new_table():
        table=doc.add_table(rows=1, cols=5); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style="Table Grid"; table.autofit=False
        for c,w in zip(table.rows[0].cells,widths): c.width=Inches(w)
        for i,h in enumerate(["Criterion"]+levels): set_cell(table.rows[0].cells[i],h,True,(255,255,255),12); shade(table.rows[0].cells[i],"0B4760")
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        return table
    table=new_table()
    for criterion_index,(title,maxpts,descs) in enumerate(item["criteria"]):
        if item["kind"] == "exam" and criterion_index == 4:
            doc.add_page_break(); table=new_table()
        row=table.add_row(); pts=[15,10,5,0] if maxpts==15 else ([4,3,2,0] if maxpts==4 else [9,7,4,0])
        if maxpts == 14: pts=[14,10,5,0]
        if maxpts == 13: pts=[13,9,5,0]
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        set_cell(row.cells[0],f"{title}\n({maxpts} pts)",True,(11,71,96),8)
        for i,(pt,desc) in enumerate(zip(pts,descs),1): set_cell(row.cells[i],f"{pt} — {desc}",False,None,7.4)
        if len(table.rows)%2==1:
            for c in row.cells: shade(c,"EAF3F6")
    if item["kind"] != "exam":
        p=doc.add_paragraph("Teacher comments: ________________________________________________________________________________________________")
        p.paragraph_format.space_before=Pt(4)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=footer.add_run("Grade 7 T3 · Closed-world assessment rubric · Same criteria for approved fallback"); run.font.name="Arial"; run.font.size=Pt(12)
    out.parent.mkdir(parents=True, exist_ok=True); doc.save(out)

expected=set()
for item in SPEC:
    name=f"7th grade - IIIT - {item['label']}.docx"
    expected.add(name)
    out=CANON/item["folder"]/name
    build(item,out)
    shutil.copyfile(out,MIRROR/name)
    print(out)

# Keep replaced assessment rubrics out of the active folders and mirrors.
for folder in ("Daily","Appreciation","Exam Projects"):
    legacy=CANON/folder/"Legacy - do not use"; legacy.mkdir(parents=True,exist_ok=True)
    for p in (CANON/folder).glob("7th grade - IIIT - *.docx"):
        if p.name not in expected and (" - 2026-" in p.name or "Rubric" in p.name or "Final Project Packet" in p.name):
            target=legacy/p.name
            if target.exists(): target.unlink()
            p.replace(target)
for p in MIRROR.glob("7th grade - IIIT - *.docx"):
    if p.name not in expected: p.unlink()
