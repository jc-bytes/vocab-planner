#!/usr/bin/env python3
"""Update the four official Grade 7 T3 DOCX plans without changing their template."""
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[6]
OUT=ROOT/'plans/7th Grade Technology/Planning/Monthly/3rd Trimester'
schedule={
'September':[
(1,'7A — Sep 16 — 45 min / 7B — Sep 17 — 45 min','Spreadsheet cell references','Practice finding rows, columns, cells, and six fixed references.','Project the labeled grid; students trace references.','Complete fixed module/task-sheet practice.','Correct six references and save/return practice.','7A and 7B — Sep 18 — 90 min','Spreadsheet Vocabulary and Cell References','Complete Daily Grade 1 using the fixed ten-term and six-reference template.','Review one completed model row and the fixed cell grid.','Complete only the announced fixed assessment.','Submit the working file or equal printed fallback.'),
(2,'7A — Sep 23 — 45 min / 7B — Sep 24 — 45 min','Spreadsheet formulas and chart practice','Practice SUM, AVERAGE, MAX, column charts, and two conclusions.','Open the fixed dataset and model one formula.','Complete assigned module practice.','Check results against the data.','7A and 7B — Sep 25 — 90 min','Spreadsheet Formulas and Column Chart','Complete Daily Grade 2 using the fixed dataset and build sheet.','Review exact range, formulas, title, and labels.','Build the fixed spreadsheet and column chart.','Submit spreadsheet/build sheet or equal paper fallback.')],
'October':[
(3,'7A — Sep 30 — 45 min / 7B — Oct 1 — 45 min','Scratch custom-block practice','Plan and trace one custom block inside repeat 4.','Model the fixed Robot delivery sequence.','Practice custom block, loop, and trace.','Correct one failed trace.','7A and 7B — Oct 2 — 90 min','Scratch Custom Block and Repeat Loop','Complete Daily Grade 5 and three fixed tests.','Review exact block name and sequence.','Build/trace the fixed task.','Submit .sb3/test sheet or equal paper trace.'),
(4,'7A — Oct 7 — 45 min / 7B — Oct 8 — 45 min','STEAM preparation','Continue the assigned cross-subject project.','Open the shared STEAM checkpoint.','Complete assigned role work.','Record result and next step.','7A and 7B — Oct 9 — 90 min','STEAM development','Complete the assigned project checkpoint.','Confirm role and milestone.','Build or investigate with assigned teacher.','Update the shared-kit work log.'),
(5,'7A — Oct 14 — 45 min / 7B — Oct 15 — 45 min','STEAM test planning','Prepare one fair test, improvement, and individual Technology log.','Review shared-kit test record.','Write expected result and setup.','Confirm one controlled variable.','7A and 7B — Oct 16 — 90 min','STEAM Preparation and Work Process','Complete Daily Grade 4 through the assigned project and individual Technology evidence.','Prepare materials and individual log.','Complete role work, test, and three log entries.','Submit Daily Grade 4 individual Technology log.'),
(6,'7A — Oct 21 — 45 min / 7B — Oct 22 — 45 min','STEAM Week preparation','Rehearse assigned explanation and organize materials.','Confirm expo/closure role.','Practice concise explanation.','Complete handoff checklist.','7A and 7B — Oct 23 — 90 min','STEAM Expo Participation and Closure','Complete Daily Grade 3 through the assigned role and individual Technology reflection.','Review role, explanation, safety, and closure.','Present/support and complete individual sheet.','Submit Daily Grade 3 individual expo/closure sheet.'),
(7,'7A — Oct 28 — 45 min / 7B — Oct 29 — 45 min','Source and image-credit practice','Practice author, date, evidence, purpose, and four-part credit checks.','Distribute fixed source cards.','Complete fixed-card practice only.','Correct decisions using the key.','7A and 7B — Oct 30 — 90 min','Source Credibility and Image Credit practice','Complete formative fixed-card practice; no new Daily grade.','Review the four checks and credit pattern.','Complete the supplied practice.','Save or return formative response.')],
'November':[
(8,'7A — Nov 4 — CLOSED / 7B — Nov 5 — CLOSED','Closure — no 45-minute class','Record actual losses: 7A Nov 4 and 7B Nov 5.','No activity.','No activity.','No evidence.','7A and 7B — Nov 6 — 90 min','Digital Work Habits and File Responsibility','Complete Appreciation Grade 1 using fixed file rules.','Review supplied folder and filename rules.','Complete six files, decisions, and verification.','Submit the fixed task sheet.'),
(9,'7A — Nov 11 — 45 min / 7B — Nov 12 — 45 min','Sensor-system and threshold practice','Trace input-process-signal-output and 17/18/19 cm boundary cases.','Open Sensor Systems practice.','Trace fixed >18 rule.','Correct boundary predictions.','7A and 7B — Nov 13 — 90 min','Mandrake Detection System Design Plan milestone','Complete the fixed plan as exam-process evidence, not a Daily grade.','Review user, chain, components, flow, tests.','Complete fixed plan.','Add plan to exam packet.'),
(10,'7A — Nov 18 — 45 min / 7B — Nov 19 — 45 min','Mandrake project planning','Confirm fixed purpose, roles, logic, and test setup.','Open project brief/checklist.','Complete plan and first log entry.','Teacher checks readiness for build.','7A and 7B — Nov 20 — 90 min','Mandrake build and first tests','Build or simulate and record first comparable trials.','Set one consistent test setup.','Run fixed trials and record actual results.','Save project/simulation and test record.'),
(11,'7A — Nov 25 — 45 min / 7B — Nov 26 — 45 min','Mandrake reliability and peer review','Prepare milestone evidence and fixed peer checks.','Review current tests and boundary.','Complete peer review and choose one change.','Prepare same-condition retest.','7A and 7B — Nov 27 — 90 min','Mandrake Project Process and Peer Feedback','Complete Appreciation Grade 2.','Confirm milestone checks.','Complete log, peer review, revision, and retest.','Submit fixed process/peer sheet.')],
'December':[
(12,'7A — Dec 2 — 45 min / 7B — Dec 3 — 45 min','Mandrake rehearsal and final checks','Use the submission checklist and rehearse the fixed explanation.','Open final checklist.','Trace demo and check all evidence.','Resolve one checkable gap.','7A and 7B — Dec 4 — 90 min','Mandrake Obstacle-Detection System — Exam','Submit and demonstrate the fixed exam system or simulation.','Confirm files/log/tests/reflection.','Demonstrate and explain.','Submit all fixed evidence.'),
(13,'7A — Dec 9 — 45 min / 7B — Dec 10 — 45 min','School exams and Technology make-up','Complete only an approved make-up or administrative task.','Dec 8 is a closure.','Use assigned equal fallback or organize files.','No new Technology assessment.','7A and 7B — Dec 11 — 90 min','School exams and Technology make-up','Complete only approved make-up/return/archive work.','Check missing-evidence list.','Complete assigned make-up only.','No new Technology assessment.')]
}
overview={
'September':(
'Use spreadsheet vocabulary, cell references, formulas, and a column chart accurately.\nInterpret a fixed dataset and preserve unchanged source values.',
'Identify fixed table parts and cell references.\nUse SUM, AVERAGE, and MAX on B2:B9.\nBuild and explain one labeled column chart.',
'Submit Daily Grades 1 and 2 as working files or equal printed fallbacks with the announced evidence.'),
'October':(
'Use a Scratch custom block and repeat loop to organize a program.\nApply fixed source-credibility and image-credit checks.\nProduce individual Technology evidence through the assigned STEAM project.',
'Build and test the fixed Scratch sequence.\nComplete Technology Daily Grade 4 on Oct 16 and Daily Grade 3 on Oct 23.\nEvaluate supplied source cards without open-web searching.',
'Submit Daily Grades 5, 4, and 3; preserve shared project records in the shared kit.'),
'November':(
'Organize digital files responsibly and explain system behavior with input, process, signal, output, and threshold logic.\nUse comparable tests and peer feedback to improve the Mandrake system.',
'Complete Appreciation Grade 1, formative sensor work, and Appreciation Grade 2.\nBuild or simulate the announced >18 cm system.',
'Submit the file-responsibility task, design plan, process/peer sheet, project log, and test evidence. Closures remove only 7A Nov 4 and 7B Nov 5.'),
'December':(
'Demonstrate and explain the fixed Mandrake obstacle-detection system or equal simulation.\nOrganize make-up evidence during the school-exam buffer.',
'Rehearse and complete the exam on Dec 4.\nUse Dec 7-11 only for assigned make-up or administration; Dec 8 is closed.',
'Submit the Dec 4 exam evidence. No new Technology assessment is assigned after Dec 4.')}

def fill(cell, title, topic, objective, pre, during, post):
    cell.text=''
    for label,text in ((title,''),('Topic:',topic),('Class Objective:',objective),('Pre-activities:',pre),('While activities:',during),('Post activities:',post)):
        p=cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        r=p.add_run(label); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10)
        if text:
            r=p.add_run('\n'+text); r.font.name='Arial'; r.font.size=Pt(10)

def header(cell, text):
    cell.text=''
    p=cell.paragraphs[0]; r=p.add_run(text); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10)

def section(cell,label,text):
    cell.text=''; p=cell.paragraphs[0]
    r=p.add_run(label+'\n'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10)
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(10)

for month,weeks in schedule.items():
    path=OUT/f'7° Technology - {month}.docx'; doc=Document(path); table=doc.tables[1]
    competences,objectives,outcomes=overview[month]
    section(doc.tables[0].rows[1].cells[0],'Competences:',competences)
    section(doc.tables[0].rows[2].cells[0],'Learning Objectives:',objectives)
    section(doc.tables[0].rows[3].cells[0],'Learning Outcomes:',outcomes)
    templates=[deepcopy(r._tr) for r in table.rows[:3]]
    for r in list(table.rows): table._tbl.remove(r._tr)
    for w in weeks:
        for tr in templates: table._tbl.append(deepcopy(tr))
        base=(len(table.rows)-3)
        (num,l45,t45,o45,p45,d45,a45,l90,t90,o90,p90,d90,a90)=w
        header(table.rows[base].cells[0],l45); header(table.rows[base].cells[1],l90)
        fill(table.rows[base+1].cells[0],'',t45,o45,p45,d45,a45); fill(table.rows[base+1].cells[1],'',t90,o90,p90,d90,a90)
        header(table.rows[base+2].cells[0],'Resources:\nComputer/login or equal print fallback; phones, photos, and screenshots prohibited.')
        header(table.rows[base+2].cells[1],'Resources:\nComputer/login or equal print fallback; phones, photos, and screenshots prohibited.')
    for p in doc.paragraphs:
        if p.text.startswith('MONTH:'): p.text=f'MONTH: {month.upper()}'
        if p.text.startswith('WEEKS:'): p.text=f'WEEKS: {len(weeks)}'
    doc.save(path); print(path)
