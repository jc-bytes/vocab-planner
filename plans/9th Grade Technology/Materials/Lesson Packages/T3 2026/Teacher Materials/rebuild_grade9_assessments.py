from pathlib import Path
from shutil import copy2
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).resolve().parents[6]
G9=ROOT/'plans/9th Grade Technology'
DOC=G9/'Assessments/Rubrics/IIIT'
GC=G9/'Assessments/Rubrics/Google Classroom'
MIR=ROOT/'plans/Shared/Generated Outputs/Rubrics 2026/9th Grade Technology/3rd Trimester'

def c(title,desc,high,meet,dev,points=(9,6,3,0)):
    return {'title':title,'description':desc,'points':list(points),'levels':['Exceeds','Meets','Developing','Not demonstrated'],'levelDescriptions':[high,meet,dev,'No checkable evidence for this criterion.']}

ready=c('Readiness and evidence responsibility','4/40: own computer/login or equal school fallback, on-time start, correct filename, and checkable handoff. Documented school failures use the fallback without deduction.','Ready, correctly named evidence is handed in independently.','One issue is corrected after one reminder.','Repeated reminders or only partial fallback evidence.',(4,3,2,0))

items=[
('Daily','9th grade - IIIT - 2026-09-15 and 2026-09-16 - Daily Grade 1 - Digital Representation Vocabulary','Daily Grade 1 — Digital Representation Vocabulary','9A Sep 15; 9B Sep 16',40,[
c('Ten fixed terms','Completes the supplied ten-term table.','All 10 rows are complete.','8–9 rows are complete.','5–7 rows are complete.'),c('Meanings','Meanings match the supplied vocabulary source.','9–10 meanings are accurate.','7–8 are accurate.','4–6 are accurate.'),c('Examples','Examples accurately show each representation term.','9–10 examples fit.','7–8 fit.','4–6 fit.'),c('Source use','Responses stay within the supplied source and taught module.','Every response is source-supported.','One response is weakly supported.','Several responses are unsupported.'),ready]),
('Appreciation','9th grade - IIIT - 2026-09-21 - Appreciation Grade 1 - Module Work Habits and Evidence Organization','Appreciation Grade 1 — Module Work Habits and Evidence Organization','9A/9B Sep 21',40,[
c('Assigned attempts','Completes the assigned image and sound practice in order.','All assigned attempts are complete.','One attempt is incomplete.','Several attempts are incomplete.'),c('Correction','Corrects one missed response using feedback.','Correction names old answer, new answer, and reason.','Correction is accurate but one detail is missing.','Correction is partly accurate.'),c('Reconstructable report','Export or printed record preserves prompts, answers, and results.','All practice can be reconstructed.','One field is unclear.','Only part can be reconstructed.'),c('Evidence organization','Name, group, filename, and checklist are complete.','All four are accurate.','Three are accurate.','One or two are accurate.'),ready]),
('Daily','9th grade - IIIT - 2026-09-22 and 2026-09-23 - Daily Grade 2 - Image and Sound Data Calculations','Daily Grade 2 — Image and Sound Data Calculations','9A Sep 22; 9B Sep 23',40,[
c('Image calculations','Uses the supplied image values and formula.','All image calculations and units are accurate.','One calculation or unit is incorrect.','Some setup is correct but several results are wrong.'),c('Sound calculations','Uses the supplied sample rate, bit depth, channels, and time.','All sound calculations and units are accurate.','One calculation or unit is incorrect.','Some setup is correct but several results are wrong.'),c('Working','Shows substitutions and operations for every required answer.','Every answer has checkable working.','One answer lacks working.','Several answers lack working.'),c('Interpretation','Explains the fixed image/sound trade-off using the results.','Both explanations are accurate and evidence-based.','Both are present; one detail is weak.','Only one explanation is usable.'),ready]),
('Appreciation','9th grade - IIIT - 2026-09-28 - Appreciation Grade 2 - Cybersecurity Work Process and Reflection','Appreciation Grade 2 — Cybersecurity Work Process and Reflection','9A/9B Sep 28',40,[
c('Analysis process','Uses evidence, threat, protection, reason in order.','All assigned responses follow all four steps.','One response misses one step.','Several steps are missing.'),c('Source evidence','Uses facts from the fixed cybersecurity source.','Every decision cites a relevant fact.','One decision has general support.','Several decisions lack support.'),c('Revision','Corrects one response after feedback.','Change and reason are exact and accurate.','Change is accurate with a general reason.','Change is only partly accurate.'),c('Reflection','Answers both fixed reflection parts.','Both parts are specific and evidence-based.','Both are present; one is general.','Only one part is usable.'),ready]),
('Daily','9th grade - IIIT - 2026-09-29 and 2026-09-30 - Daily Grade 3 - Cybersecurity Threats and Protections','Daily Grade 3 — Cybersecurity Threats and Protections','9A Sep 29; 9B Sep 30',40,[
c('Threat identification','Identifies the threat in each fixed scenario.','All assigned threats are accurate.','One is incorrect.','Several are incorrect.'),c('Protection selection','Selects the best supplied protection for each scenario.','All protections fit the evidence.','One protection is weak.','Several protections do not fit.'),c('Justifications','Uses scenario evidence to justify the required answers.','All required justifications cite exact evidence.','One is general.','Several are incomplete.'),c('Safe response','Applies the taught safe-response rule without adding unsafe actions.','Every response is safe and source-aligned.','One detail is unclear.','Several details are unsafe or unsupported.'),ready]),
('Daily','9th grade - IIIT - 2026-11-02 - Daily Grade 4 - Six-Risk Cybersecurity Analysis','Daily Grade 4 — Six-Risk Cybersecurity Analysis','9A/9B Nov 2',40,[
c('R1–R6 ratings','Rates probability and impact for all six fixed records.','All 12 ratings are evidence-supported.','10–11 are supported.','6–9 are supported.'),c('Protections','Chooses one feasible protection for each record.','All 6 fit the fixed source.','5 fit.','3–4 fit.'),c('Risk reasoning','Explains ratings using facts from each record.','All 6 explanations use record facts.','5 use record facts.','3–4 are usable.'),c('Priority recommendation','Identifies and justifies the highest-priority risk.','Choice, evidence, and action are specific.','All parts are present; one is general.','Only one or two parts are usable.'),ready]),
('Daily','9th grade - IIIT - 2026-11-16 - Daily Grade 5 - STEM System Proposal','Daily Grade 5 — STEM System Proposal','9A/9B Nov 16',40,[
c('Problem and users','Uses the assigned option card to state problem and users.','Both are exact and source-aligned.','Both are present; one detail is weak.','Only one is usable.'),c('System design','Includes goal, platform/materials, and labeled design.','All required design fields are complete.','One field is incomplete.','Several fields are incomplete.'),c('C1–C3 criteria','Writes three measurable success criteria from the assigned card.','All 3 are measurable and aligned.','2 are measurable and aligned.','1 is usable.'),c('Evidence plan','Names versions, log, three tests, comparison, and live demo.','All 6 evidence items are named.','5 are named.','3–4 are named.'),ready]),
]

exam=[
c('Approved system and design','Uses the assigned proposal, users, goal, constraints, and labeled design.','All required elements remain aligned.','One element is unclear.','Several elements are incomplete.',(14,10,5,0)),
c('Prototype function','Produces the checkable product behavior required by the option card.','All required behavior is checkable.','Most behavior works.','Only part works.',(14,10,5,0)),
c('Controlled testing','Records the three fixed inputs, expected/actual results, and pass/fail.','All three complete records are accurate.','Two are complete and accurate.','One is complete.',(14,10,5,0)),
c('Evidence-based improvement','Names a weakness, changes it, retests, and compares before/after.','All four parts are specific.','One part is general.','Only one or two parts are usable.',(13,9,5,0)),
c('Log and reflection','Maintains dated decisions and answers the fixed reflection.','Log and reflection are complete and specific.','Both are present with small gaps.','Only one is usable.',(13,9,5,0)),
c('Live demonstration','Shows the product and explains problem, design, tests, improvement, and limitation.','All five points are accurate and checkable.','Four are accurate.','Two or three are usable.',(13,9,5,0)),
c('Readiness and evidence responsibility','9/90: own computer/login or equal fallback, milestones, version names, and final handoff. Documented school failures use the fallback without deduction.','All checkpoints and files are ready and checkable.','One issue is corrected after one reminder.','Repeated issues or partial fallback evidence.',(9,7,4,0))]
items.append(('Exam Projects','9th grade - IIIT - 2026-11-17 and 2026-11-18 - Exam - STEM Prototype Build Test Improve and Demonstrate','Exam — STEM Prototype: Build, Test, Improve, and Demonstrate','9A Nov 17; 9B Nov 18',90,exam))

# Active T3 structure: five daily grades, two regular Technology appreciation
# grades, and one exam. STEAM preparation/expo are Daily Grades 4 and 3.
image=next(x for x in items if 'Image and Sound Data Calculations' in x[2])
threat=next(x for x in items if 'Cybersecurity Threats and Protections' in x[2])
risk=next(x for x in items if 'Six-Risk Cybersecurity Analysis' in x[2])
appreciation1=next(x for x in items if 'Module Work Habits and Evidence Organization' in x[2])
appreciation2=next(x for x in items if 'Cybersecurity Work Process and Reflection' in x[2])
exam_item=next(x for x in items if x[0]=='Exam Projects')
image=('Daily','9th grade - IIIT - 2026-09-22 and 2026-09-23 - Daily Grade 1 - Image and Sound Data Calculations','Daily Grade 1 — Image and Sound Data Calculations','9A Sep 22; 9B Sep 23',40,image[5])
threat=('Daily','9th grade - IIIT - 2026-09-29 and 2026-09-30 - Daily Grade 2 - Cybersecurity Threats and Protections','Daily Grade 2 — Cybersecurity Threats and Protections','9A Sep 29; 9B Sep 30',40,threat[5])
risk=('Daily','9th grade - IIIT - 2026-11-02 - Daily Grade 5 - Six-Risk Cybersecurity Analysis','Daily Grade 5 — Six-Risk Cybersecurity Analysis','9A/9B Nov 2',40,risk[5])
steam1=('Daily','9th grade - IIIT - 2026-10-13 and 2026-10-14 - Daily Grade 4 - STEAM Project Preparation and Work Process','Daily Grade 4 — STEAM Project Preparation and Work Process','9A Oct 13; 9B Oct 14',40,[
c('Role and first task','Assigned project subject, role, first task, and success rule in the shared log.','All 4 fields are accurate and specific.','3 fields are accurate.','1–2 fields are usable.'),
c('Dated individual contribution','One preparation contribution names action, result, and next step.','All 3 parts are specific and teacher-observed.','All are present; one is general.','Only 1–2 are usable.'),
c('Test and improvement','One test row records condition, expected result, actual result, and next improvement.','All 4 fields are accurate and connected.','3 are accurate.','1–2 are usable.'),
c('Safe file and material routine','Observable responsible handling, file organization, and respectful collaboration.','All 3 routines are independent.','All are shown with one reminder.','Repeated reminders are needed.'),ready])
steam2=('Daily','9th grade - IIIT - 2026-10-20 and 2026-10-21 - Daily Grade 3 - STEAM Expo Participation and Closure','Daily Grade 3 — STEAM Expo Participation and Closure','9A Oct 20; 9B Oct 21',40,[
c('Assigned expo role','Completes the individual speaking, demonstration, setup, or support role.','Role is complete and independent.','Role is complete with one prompt.','Role is partly complete.'),
c('Five-point explanation','Explains problem, goal, own contribution, one test result, and one improvement.','All 5 points are accurate.','4 are accurate.','2–3 are usable.'),
c('Evidence response','Answers one teacher question using the shared product or log.','Answer cites specific evidence.','Answer is accurate but general.','Answer is partly accurate.'),
c('Closure and reflection','Completes cleanup/file return and the fixed individual reflection.','Both are complete and specific.','Both are complete; one detail is general.','Only one is usable.'),ready])
items=[appreciation1,image,appreciation2,threat,steam2,steam1,risk,exam_item]

def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pr.append(shd)

def build(folder,basename,title,date,total,criteria):
    d=Document(); s=d.sections[0]; s.orientation=WD_ORIENT.LANDSCAPE; s.page_width,s.page_height=Inches(11),Inches(8.5); s.top_margin=s.bottom_margin=Inches(.5); s.left_margin=s.right_margin=Inches(.5)
    d.styles['Normal'].font.name='Arial'; d.styles['Normal'].font.size=Pt(9 if total==90 else 10)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(11,59,85)
    p=d.add_paragraph(f'Grade 9 Technology · Third Trimester 2026 · {date} · Total: {total} points'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph('Student: ______________________________  Class: ______  Date: __________')
    d.add_heading('Directions and evidence boundary',1)
    d.add_paragraph('Use only the fixed source, starter, template, module sections, test inputs, and response questions named in the matching package direction. Submit the named file or the equal printed/local fallback. No phone, photo, screenshot, open-web research, or unannounced product is required. A documented school device, network, or account failure receives the equal fallback without penalty.')
    if total==90: d.add_paragraph('Required evidence: approved proposal, prototype/product, versioned files, dated log, three fixed tests, weakness, improvement and comparison, live demonstration, and reflection.')
    t=d.add_table(rows=1,cols=5); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,v in enumerate(['Criterion','Highest','Meets','Developing','Not demonstrated']): t.rows[0].cells[i].text=v; shade(t.rows[0].cells[i],'0B3B55')
    for cell in t.rows[0].cells:
        for run in cell.paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
    for cr in criteria:
        row=t.add_row().cells; row[0].text=cr['title']+'\n'+cr['description']
        for j in range(4): row[j+1].text=f"{cr['points'][j]} points\n{cr['levelDescriptions'][j]}"
    d.add_paragraph(f'Score: ______ / {total}    Teacher evidence note: __________________________________________')
    out=DOC/folder/(basename+'.docx'); out.parent.mkdir(parents=True,exist_ok=True); d.save(out)
    mp=MIR/folder/(basename+'.docx'); mp.parent.mkdir(parents=True,exist_ok=True); copy2(out,mp)

for item in items: build(*item)

def relabel_check(path, replacements):
    d=Document(path)
    for p in d.paragraphs:
        for old,new in replacements:
            if old in p.text:
                for run in p.runs: run.text=run.text.replace(old,new)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old,new in replacements:
                        if old in p.text:
                            for run in p.runs: run.text=run.text.replace(old,new)
    d.save(path)

daily=DOC/'Daily'
for p in daily.glob('*Summative 2 - Image and Sound Representation Check*.docx'):
    relabel_check(p,[('Summative 2','Daily Grade 1'),('summative 2','Daily Grade 1')])
for p in daily.glob('*Summative 3 - Cyber Security Scenario Quiz*.docx'):
    relabel_check(p,[('Summative 3','Daily Grade 2'),('summative 3','Daily Grade 2')])

spec=[]
for folder,basename,title,date,total,criteria in items:
    spec.append({'basename':basename+' - Google Classroom Rubric.xlsx','title':title,'date':date,'total':total,'criteria':criteria})
(Path(__file__).parent/'grade9_classroom_rubric_specs.json').write_text(json.dumps(spec,indent=2),encoding='utf-8')
print('built',len(items),'Grade 9 DOCX rubrics and specs')
