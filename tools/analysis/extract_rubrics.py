from docx import Document
import json
import sys

def cell_text(cell):
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())

out = []
for path in sys.argv[1:]:
    doc = Document(path)
    item = {
        "path": path,
        "paragraphs": [p.text.strip() for p in doc.paragraphs if p.text.strip()],
        "tables": [],
    }
    for table in doc.tables:
        item["tables"].append([[cell_text(cell) for cell in row.cells] for row in table.rows])
    out.append(item)
print(json.dumps(out, indent=2, ensure_ascii=False))
