#!/usr/bin/env python3
"""Create aligned Grade 7-9 T3 rubric DOCX files from the school template."""

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "plans/6th Grade Technology/Assessments/Rubrics/IIIT/6th grade - IIIT - Week 1 - Rubric for Summative 1.docx"
FONT = "Arial"

READY = (
    "Evidence completion and submission check",
    "Every named part is complete, identifiable, and accessible in the required file or approved paper fallback.",
    "One named part, label, or filename detail needs correction; the work remains checkable.",
    "Several named parts or labels are missing, but some checkable evidence is submitted.",
    "No checkable digital or approved paper evidence is submitted.",
)

STEAM_1 = [
    ("Preparation", "Uses the assigned-project brief in both STEAM weeks, states the project and role, locates the shared product, and begins the named task.", "Completes all four actions with one reminder.", "Completes two or three actions or needs repeated reminders.", "Completes zero or one action and provides no checkable preparation evidence."),
    ("Participation and role", "Completes the assigned role and contributes useful work at every checkpoint, regardless of project subject.", "Completes the role with one incomplete checkpoint or reminder.", "Completes only part of the role and needs frequent support.", "Provides no usable contribution evidence."),
    ("Responsible work", "Uses files, materials, class time, and group communication safely and responsibly.", "Works responsibly with one corrected reminder.", "Needs repeated reminders about materials, files, time, or conduct.", "Uses resources unsafely or prevents group progress."),
    ("Individual work log", "Records at least 3 dated entries naming the task, result, and next step.", "Records 2 complete entries and one partial entry.", "Records 1 complete entry or unclear fragments.", "Provides no checkable individual log."),
]
STEAM_2 = [
    ("Assigned expo role", "Completes the assigned speaking, support, or demonstration role at the expo.", "Completes the role with one prompt or small missing part.", "Completes only part of the role or needs several prompts.", "Does not complete the role or approved make-up."),
    ("Project explanation", "Explains the problem, product, process, test result, and improvement, even when the project belongs to another subject.", "Explains 4 of the 5 points clearly.", "Explains 2-3 points or several are unclear.", "Explains 0-1 point or gives no checkable explanation."),
    ("Respect and cleanup", "Listens respectfully, follows expo directions, and completes the assigned cleanup task.", "Meets all expectations with one corrected reminder.", "Meets only some expectations and needs repeated reminders.", "Does not meet participation or cleanup expectations."),
    ("Individual reflection", "Records one specific result, one challenge, and one improvement from the project experience.", "Includes all three parts; one needs more detail.", "Includes one or two clear parts.", "Reflection is missing or unrelated."),
]

def C(name, high, good, developing, beginning):
    return (name, high, good, developing, beginning)

GRADES = {
    7: [
        (1, 1, "Spreadsheet Vocabulary and Examples", [
            C("Required terms", "Includes every assigned spreadsheet term.", "One term is missing or unclear.", "Several terms are missing or unclear.", "Most terms are missing."),
            C("Definitions", "Definitions are accurate and student-friendly.", "Most definitions are accurate.", "Several definitions need correction.", "Definitions are mostly missing or incorrect."),
            C("Examples", "Every term has a correct spreadsheet example.", "One example needs clarification.", "Several examples are weak or incorrect.", "Examples are mostly missing."),
            C("Organization", "The table is complete, consistent, and easy to read.", "One formatting issue remains.", "Several entries are difficult to check.", "The work is incomplete or unreadable."),
        ]),
        (2, 2, "Spreadsheet Formulas and Chart", [
            C("Formulas", "Uses the assigned formulas correctly and shows correct results.", "One formula or result needs correction.", "Several formulas need support.", "Formulas are missing or do not work."),
            C("Data range", "Uses the complete assigned range with no extra or missing values.", "One range reference needs correction.", "Several values are missing or misplaced.", "No usable range is selected."),
            C("Chart", "Creates the correct chart with an accurate title and labels.", "Chart works with one minor label or type issue.", "Chart is present but difficult to interpret.", "Chart is missing or unrelated."),
            C("Interpretation", "Writes a specific conclusion supported by the formula results and chart.", "Conclusion is supported but needs one detail.", "Conclusion is vague or partly supported.", "Conclusion is missing or unsupported."),
        ]),
        (6, 3, "Scratch Subroutine and Loop", [
            C("Decomposition", "Breaks the task into clear parts with a useful plan.", "Plan is usable with one unclear part.", "Plan is incomplete or needs repeated support.", "No usable plan is provided."),
            C("Custom block", "Creates and calls a custom block that performs the intended task.", "Custom block works with one minor issue.", "Custom block is incomplete or called incorrectly.", "No usable custom block is present."),
            C("Loop", "Uses a loop correctly to reduce repeated blocks.", "Loop mostly works with one correction.", "Loop is present but does not control repetition correctly.", "Loop is missing."),
            C("Testing record", "Records at least 3 tests, results, and one improvement.", "Records 2 complete tests and an improvement.", "Testing record is incomplete or vague.", "No checkable testing record is provided."),
        ]),
        (7, 4, "Source Credibility and Image Credit", [
            C("Source checks", "Checks author or organization, date, evidence, and purpose for each source.", "One check is incomplete.", "Several checks are vague or missing.", "Sources are accepted without checks."),
            C("Credibility decisions", "Makes accurate decisions and explains each with source evidence.", "Most decisions are accurate; one explanation needs detail.", "Several decisions lack evidence.", "Decisions are missing or unsupported."),
            C("Image credit", "Creates a complete credit with creator, title, source, and license when available.", "Credit is usable with one missing element.", "Several credit elements are missing.", "Credit is missing or unusable."),
            C("Responsible use", "Uses only approved material and clearly separates fact from opinion.", "One small responsible-use issue is corrected.", "Needs repeated reminders about use or attribution.", "Uses unapproved material or misrepresents a source."),
        ]),
        (9, 5, "Mandrake Detection System Design Plan", [
            C("Goal and user", "States a clear user need and a testable system goal.", "Goal is clear but one detail needs refinement.", "Goal is broad or only partly testable.", "Goal or user need is missing."),
            C("Input-process-output", "Identifies a suitable sensor input, decision process, and warning output.", "All three parts are present; one needs detail.", "One part is missing or connections are unclear.", "The system flow is not usable."),
            C("Flowchart and components", "Provides a logical flowchart and complete component list.", "Both are present with one minor omission.", "One is incomplete or difficult to follow.", "Flowchart and components are mostly missing."),
            C("Test plan", "Defines at least 3 measurable tests including safe conditions and expected results.", "Defines 3 tests; one needs a clearer expected result.", "Tests are vague or incomplete.", "No usable test plan is provided."),
        ]),
    ],
    8: [
        (1, 1, "App Vocabulary and Decomposition", [
            C("Required terms", "Includes every assigned app-design and decomposition term.", "One term is missing or unclear.", "Several terms are missing or unclear.", "Most terms are missing."),
            C("Definitions", "Definitions are accurate and student-friendly.", "Most definitions are accurate.", "Several need correction.", "Definitions are mostly missing or incorrect."),
            C("Examples", "Every term has a correct app or algorithm example.", "One example needs clarification.", "Several examples are weak.", "Examples are mostly missing."),
            C("Task breakdown", "Breaks one app goal into clear, ordered subtasks.", "Breakdown is usable with one unclear step.", "Steps are incomplete or poorly ordered.", "No usable breakdown is provided."),
        ]),
        (2, 2, "App Screen and Event Map", [
            C("Screen map", "Shows every required screen and a clear purpose for each.", "One screen or purpose needs detail.", "Several screens are missing or unclear.", "No usable screen map is provided."),
            C("Navigation", "Connections show how users move between screens without dead ends.", "One connection needs correction.", "Several links are missing or confusing.", "Navigation is not usable."),
            C("Events and actions", "Pairs each user event with the correct app action.", "Most pairs are correct; one needs detail.", "Several event-action pairs are incorrect.", "Events and actions are mostly missing."),
            C("Testing scenarios", "Includes at least 3 user tests with expected results.", "Includes 3 tests; one expected result is vague.", "Tests are incomplete or repetitive.", "No usable tests are provided."),
        ]),
        (7, 3, "Python Arithmetic and Output", [
            C("Variables and values", "Uses meaningful variable names and correct numeric values.", "One name or value needs correction.", "Several variables are unclear or incorrect.", "Variables are missing or unusable."),
            C("Operations", "Uses the required arithmetic operations correctly.", "One operation needs correction.", "Several operations are incorrect.", "No working arithmetic is present."),
            C("Output", "Displays correct, clearly labeled results for every calculation.", "Results are correct with one unclear label.", "Several results or labels are incorrect.", "No usable output is produced."),
            C("Testing", "Tests at least 3 inputs and records expected and actual results.", "Tests 2-3 inputs with one missing detail.", "Testing is limited or unclear.", "No checkable tests are provided."),
        ]),
        (8, 4, "Python Debugging", [
            C("Error identification", "Correctly identifies every assigned syntax or logic error.", "Identifies most errors; one needs correction.", "Several errors are missed.", "Errors are not identified."),
            C("Corrections", "Produces working corrected code without changing the intended task.", "Code works after one minor correction.", "Code only partly works or changes the task.", "No working correction is provided."),
            C("Explanation", "Explains what caused each error and why the fix works.", "Most explanations are accurate.", "Several explanations are vague or incorrect.", "Explanations are missing."),
            C("Test log", "Records inputs, expected output, actual output, and final status for at least 3 tests.", "Records 3 tests with one missing field.", "Records fewer than 3 usable tests.", "No test log is provided."),
        ]),
        (9, 5, "Digital Representation and Binary", [
            C("Binary conversion", "Correctly converts every assigned value between binary and decimal.", "One conversion error remains.", "Several conversions need correction.", "Conversions are mostly missing or incorrect."),
            C("Image representation", "Accurately explains pixels, resolution, and color depth with an example.", "Explanation is accurate with one missing detail.", "Several ideas are confused.", "Explanation is missing or incorrect."),
            C("Sound representation", "Accurately explains samples, sample rate, and bit depth with an example.", "Explanation is accurate with one missing detail.", "Several ideas are confused.", "Explanation is missing or incorrect."),
            C("File-size reasoning", "Compares two representations and correctly predicts the larger file with evidence.", "Prediction is correct but evidence needs detail.", "Reasoning is only partly supported.", "Prediction or evidence is missing."),
        ]),
    ],
    9: [
        (1, 1, "Audiovisual Representation Vocabulary", [
            C("Ten required terms", "Includes pixel, resolution, color depth, RGB, sample, sample rate, sample size, microphone, speaker, and compression.", "Includes 9 of the 10 named terms.", "Includes 6-8 of the named terms.", "Includes 0-5 of the named terms."),
            C("Ten definitions", "Defines all 10 terms accurately in clear student wording using the provided source.", "Defines 9 terms accurately; one needs correction.", "Defines 6-8 terms accurately.", "Defines 0-5 terms accurately."),
            C("Ten examples", "Gives one correct image, sound, device, or file example for all 10 terms.", "Gives 9 correct examples; one needs correction.", "Gives 6-8 correct examples.", "Gives 0-5 correct examples."),
            C("Three-to-five-sentence connection", "Names one practiced setting, describes its change, and correctly explains the file-size result in 3-5 sentences.", "Includes all three ideas; one needs clearer detail.", "Includes only one or two accurate ideas or the length is outside 3-5 sentences.", "The connection is missing or contradicts the taught relationship."),
        ]),
        (2, 2, "Image and Sound Representation Check", [
            C("Vocabulary meanings", "Accurately explains pixel, resolution, color depth, sample rate, and sample size in context.", "Four meanings are accurate; one needs correction.", "Two or three meanings are accurate.", "Zero or one meaning is accurate."),
            C("Image operations", "Shows `10 x 8 = 80 pixels` and `10 x 8 x 8 = 640 bits` with correct units.", "Both results are correct; one operation or unit is incomplete.", "One result is correct or both show partial work.", "Both operations are missing or incorrect."),
            C("Sound and quality reasoning", "Distinguishes samples per second from bits per sample and explains how a higher setting can increase detail and file size.", "The two meanings are correct; one relationship needs detail.", "One meaning or relationship is correct.", "The sound reasoning is missing or incorrect."),
            C("Compression recommendation", "Names one taught change, one file-size benefit, and one possible quality cost.", "Names all three; one needs clearer detail.", "Names only one or two accurate parts.", "The recommendation is missing or not supported by the taught trade-off."),
        ]),
        (6, 3, "Cybersecurity Scenario Analysis", [
            C("Six threat identifications", "Correctly identifies phishing, malware family, DDoS, brute force, social engineering/phishing, and worm/malware in scenarios 1-6.", "Five identifications are correct.", "Three or four identifications are correct.", "Zero to two identifications are correct."),
            C("Six protection choices", "Matches one safe, practical source-page protection to all six scenarios.", "Five protections fit; one needs correction.", "Three or four protections fit.", "Zero to two protections fit or an unsafe action is proposed."),
            C("Three evidence justifications", "For three different scenarios, names the clue, explains why the threat fits, and explains how the protection reduces risk.", "All three are present; one explanation needs detail.", "Only one or two complete justifications are present.", "No complete evidence justification is present."),
            C("School protection rule", "Chooses one taught protection and explains how it helps many school users without unsafe testing or blaming a user.", "The rule is safe and broad; one reason needs detail.", "The rule is partly useful or weakly explained.", "The rule is missing, unsafe, or unrelated."),
        ]),
        (7, 4, "Cybersecurity Risk Map", [
            C("Fixed records R1-R6", "Completes asset and threat or harmful-event analysis for all six fixed records.", "Completes five records; one is missing or unclear.", "Completes three or four records.", "Completes zero to two records."),
            C("Probability and impact", "For all six records, assigns low/high probability and low/high impact separately and cites a supporting fact for each rating.", "All ratings are present; one or two facts need correction.", "Several ratings or supporting facts are missing.", "Ratings are mostly missing or arbitrary."),
            C("Six protection strategies", "Matches one feasible source-page protection to every record and explains how each helps.", "Five protections fit; one needs correction.", "Three or four protections fit.", "Zero to two protections fit or unsafe actions are proposed."),
            C("Priority paragraph", "Names one highest-priority record and defends it with its probability, impact, and at least one record fact.", "Includes all three supports; one needs detail.", "Includes only one or two supports.", "The paragraph is missing or contradicts the completed map."),
        ]),
        (9, 5, "STEM Project Proposal", [
            C("Assigned problem, users, and goal", "Names the assigned option, copies its fixed problem and users, explains the need, and writes a goal with a visible or numeric success condition.", "All parts are present; one needs detail.", "Two parts are missing or too broad.", "No usable assigned problem and goal are present."),
            C("Product, tool, and labeled design", "Names an allowed product and approved tool, lists needed resources and file location, and shows input, process or decision, and output in a labeled design.", "All parts are present; one needs correction.", "Several named parts are incomplete.", "No feasible product and labeled design are present."),
            C("Criteria C1-C3", "Uses all three required option-card inputs and gives an expected result and visible or numeric pass rule for C1, C2, and C3.", "All three criteria are present; one result or pass rule needs detail.", "Only one or two measurable criteria are complete.", "No usable measurable criteria are present."),
            C("Evidence plan and limitation", "Names the v1 file, final file, dated logbook, three test rows, before/after comparison, live demo, and one realistic limitation.", "All items are present; one needs detail.", "Several evidence items or the limitation are missing.", "No checkable evidence plan is present."),
        ]),
    ],
}

def set_cell(cell, text):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = FONT
            run.font.size = Pt(12)
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{key}"), FONT)

def build(grade, folder, week, number, title, criteria, label="Summative"):
    doc = Document(TEMPLATE)
    table = doc.tables[0]
    groups = "A & B & C" if grade == 7 else "A & B"
    letters = "A B C" if grade == 7 else "A B"
    table.cell(0, 0).text = "\n".join([
        "Academia Internacional David", "Robotics and Technology", "3rd Trimester",
        f"{label} #{number}", f"{grade}th {groups}",
        f"Name: ________________________________   Date: __________________   Group: {grade}° {letters}",
        "Teacher: Porfirio Rios                  Score: _____ / 40pts", title,
    ])
    for idx, criterion in enumerate(criteria, start=2):
        for cell, value in zip(table.rows[idx].cells[:5], criterion):
            set_cell(cell, value)
    for cell, value in zip(table.rows[7].cells[:5], READY):
        set_cell(cell, value)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                set_cell(cell, cell.text)
    out = ROOT / f"plans/{grade}th Grade Technology/Assessments/Rubrics/IIIT/{folder}"
    out.mkdir(parents=True, exist_ok=True)
    filename = f"{grade}th grade - IIIT - Week {week} - Rubric for {label} {number} - {title}.docx"
    doc.save(out / filename)
    print(out / filename)

for grade, daily in GRADES.items():
    for spec in daily:
        build(grade, "Daily", *spec)
    build(grade, "Appreciation", 4, 1, "STEAM Preparation and Work Process", STEAM_1, "Appreciation Grade")
    build(grade, "Appreciation", 5, 2, "STEAM Expo Participation and Closure", STEAM_2, "Appreciation Grade")
