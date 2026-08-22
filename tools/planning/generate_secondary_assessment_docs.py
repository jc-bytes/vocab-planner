#!/usr/bin/env python3
"""Generate missing secondary assessment DOCX documents.

Creates appreciation grade forms, exam project packets, and one missing
8th grade rubric from the approved 2026 assessment plan.
"""

from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PLANS_ROOT = ROOT / "plans"
GENERATED_ROOT = PLANS_ROOT / "Shared/Generated Outputs/Assessment Docs 2026"
GENERATED_RUBRICS_ROOT = PLANS_ROOT / "Shared/Generated Outputs/Rubrics 2026"
RUBRIC_TEMPLATE = PLANS_ROOT / "Shared/Templates" / "Xth grade - XT - Week X - Rubric for Summative X.docx"
LOGO_PATH = ROOT / "apps" / "sparks" / "logo.jpeg"
RENDER_DOCX = (
    Path.home()
    / ".codex/plugins/cache/openai-primary-runtime/documents/26.521.10419/skills/documents/render_docx.py"
)

TRIMESTERS = {
    "IT": "1st Trimester",
    "IIT": "2nd Trimester",
    "IIIT": "3rd Trimester",
}

GRADE_FOLDERS = {
    "7th": PLANS_ROOT / "7th Grade Technology",
    "8th": PLANS_ROOT / "8th Grade Technology",
    "9th": PLANS_ROOT / "9th Grade Technology",
}


@dataclass(frozen=True)
class AppreciationDoc:
    grade: str
    trimester: str
    week: int
    number: int
    title: str
    evidence: str
    criteria: tuple[str, ...]
    reflection: str
    special_form: str = ""
    template_criteria: tuple[tuple[str, str, str, str, str], ...] = ()


@dataclass(frozen=True)
class ExamDoc:
    grade: str
    trimester: str
    title: str
    overview: str
    evidence: tuple[str, ...]
    schedule: tuple[str, ...]
    presentation: tuple[str, ...]
    criteria: tuple[tuple[str, str, str, str, str], ...]


@dataclass(frozen=True)
class RubricDoc:
    grade: str
    trimester: str
    week: int
    summative: int
    title: str
    evidence: str
    criteria: tuple[tuple[str, str, str, str, str], ...]


@dataclass(frozen=True)
class OutputRecord:
    kind: str
    title: str
    primary: Path
    mirrors: tuple[Path, ...]
    source: str


APPRECIATION_DOCS: tuple[AppreciationDoc, ...] = (
    AppreciationDoc(
        "7th",
        "IT",
        4,
        1,
        "Robotics Safety and Responsibility",
        "Completed checklist during mBot setup, testing, movement practice, and cleanup.",
        (
            "Prepared materials and followed robot safety routines.",
            "Handled mBot parts, cables, batteries, and work area carefully.",
            "Shared builder, coder, tester, or recorder roles responsibly.",
            "Listened to instructions and used class time productively.",
            "Cleaned up materials and recorded one useful debugging note.",
        ),
        "What safe routine helped your group work better today?",
    ),
    AppreciationDoc(
        "7th",
        "IT",
        6,
        2,
        "Perseverance and Feedback",
        "Self-reflection from mBot route testing and debugging.",
        (
            "Kept trying when the route or robot behavior did not work at first.",
            "Used feedback from the teacher or peers to improve one test.",
            "Stayed organized with route notes, code changes, or screenshots.",
            "Explained one problem and one attempted fix clearly.",
            "Showed respectful teamwork while testing and retesting.",
        ),
        "Which feedback or test helped you improve the robot route?",
    ),
    AppreciationDoc(
        "7th",
        "IIT",
        6,
        1,
        "Countdown Program and Constructive Peer Feedback",
        "Create and submit a Scratch countdown or rhythm-timing program, then evaluate a classmate's program with evidence.",
        (
            "Use of class time",
            "Organization and readiness",
            "Scratch practice evidence",
            "Constructive classmate feedback",
            "Participation and response to feedback",
        ),
        "What value did you adjust, and what specific feedback did you give your classmate?",
        "constructive_peer_feedback",
    ),
    AppreciationDoc(
        "7th",
        "IIT",
        8,
        2,
        "Pair Programming Responsibility",
        "Complete a pair-programming self-reflection, then build a practice dance move with a partner as evidence.",
        (
            "Use of class time",
            "Organization and readiness",
            "Driver/navigator roles",
            "Scratch dance practice evidence",
            "Communication and response to feedback",
        ),
        "Which pair-programming role helped you most, and why?",
        "pair_programming_reflection",
    ),
    AppreciationDoc(
        "7th",
        "IIIT",
        3,
        1,
        "Data Teamwork",
        "Checklist during spreadsheet, formatting, and data interpretation practice.",
        (
            "Handled shared data carefully and avoided changing another person's work without permission.",
            "Communicated respectfully while sorting, filtering, or checking data.",
            "Shared work fairly during spreadsheet practice.",
            "Stayed organized with file names, cells, formulas, or charts.",
            "Used feedback to improve a spreadsheet or interpretation.",
        ),
        "What teamwork habit helped the data task stay accurate?",
    ),
    AppreciationDoc(
        "7th",
        "IIIT",
        9,
        2,
        "Project Readiness and Responsibility",
        "Self-reflection before the Mandrake Detection System exam project.",
        (
            "Prepared project files, equipment notes, or setup materials.",
            "Showed responsibility with safe equipment use.",
            "Identified a realistic first build task for the project.",
            "Responded to feedback or review before project work.",
            "Stayed organized with roles, notes, and presentation outline.",
        ),
        "What is one readiness habit you will use during the exam project?",
    ),
    AppreciationDoc(
        "8th",
        "IT",
        4,
        1,
        "Collaboration During Peer Review",
        "Collaboration checklist during vector graphic peer review and improvement.",
        (
            "Gave respectful, useful feedback about design choices.",
            "Used peer or self-review notes to improve the campaign graphic.",
            "Stayed organized with before and after evidence.",
            "Used class time responsibly during design revision.",
            "Communicated clearly and respected other students' work.",
        ),
        "Which review note created the strongest improvement in your design?",
    ),
    AppreciationDoc(
        "8th",
        "IT",
        8,
        2,
        "Web Unit Self-Reflection",
        "Self-reflection about effort, organization, digital responsibility, feedback, and improvement.",
        (
            "Stayed organized with pages, sources, files, and portfolio evidence.",
            "Used feedback to improve web or research work.",
            "Practiced digital responsibility with sources and credits.",
            "Used class time productively while building portfolio evidence.",
            "Identified a realistic next step before the exam project.",
        ),
        "What web or research habit will help your portfolio project most?",
    ),
    AppreciationDoc(
        "8th",
        "IIT",
        4,
        1,
        "Safe Build Teamwork",
        "Safe material use and teamwork checklist during Arduino build practice.",
        (
            "Handled the Arduino/Freenove kit and components carefully.",
            "Checked power, ground, LED direction, and wiring before testing.",
            "Shared build, code, test, and recorder roles responsibly.",
            "Communicated respectfully during troubleshooting.",
            "Returned parts and left the workstation ready.",
        ),
        "What safe build habit prevented or fixed a problem?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Safe material use",
                "Handles kit, parts, wiring, power and ground, and LED direction safely.",
                "Uses materials safely with one minor reminder.",
                "Needs reminders for handling, wiring, or testing.",
                "Unsafe handling or no safe build evidence.",
            ),
            (
                "Teamwork and roles",
                "Shares build, code, test, and recorder roles responsibly.",
                "Usually shares roles and communicates respectfully.",
                "Roles are uneven or need several reminders.",
                "Does not participate responsibly.",
            ),
            (
                "Build challenge evidence",
                "LED or button build is tested and documented with a diagram or photo.",
                "Build evidence is mostly complete.",
                "Build evidence is partial or hard to check.",
                "Build evidence is missing.",
            ),
            (
                "Organization and perseverance",
                "Keeps parts organized, troubleshoots, and leaves station ready.",
                "Mostly organized and keeps trying.",
                "Organization or perseverance needs support.",
                "Workspace or troubleshooting evidence is poor.",
            ),
        ),
    ),
    AppreciationDoc(
        "8th",
        "IIT",
        8,
        2,
        "Arduino Perseverance",
        "Self-reflection about perseverance, organization, careful testing, feedback, and improvement.",
        (
            "Kept trying during wiring, code, upload, or serial monitor problems.",
            "Recorded test evidence clearly.",
            "Used feedback or a checklist to improve the build.",
            "Stayed organized with diagrams, code, photos, or notes.",
            "Worked safely and respectfully with materials and partners.",
        ),
        "What Arduino problem did you keep working on, and what changed?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Perseverance during testing",
                "Keeps trying through wiring, code, or upload issues and records result.",
                "Usually keeps trying and records result.",
                "Needs support to keep testing or explain.",
                "Stops without useful testing evidence.",
            ),
            (
                "Organization of evidence",
                "Diagrams, code, photos, notes, or tests are organized.",
                "Evidence is mostly organized.",
                "Evidence is partial or hard to check.",
                "Evidence is missing.",
            ),
            (
                "Feedback and improvement",
                "Uses feedback or a checklist to improve build or code.",
                "Uses feedback for a basic improvement.",
                "Feedback or improvement is unclear.",
                "No feedback response or improvement.",
            ),
            (
                "Integrated practice challenge",
                "Builds and tests one challenge with one input and one output.",
                "Challenge mostly works.",
                "Input, output, or test result is incomplete.",
                "Challenge evidence is missing.",
            ),
        ),
    ),
    AppreciationDoc(
        "8th",
        "IIIT",
        8,
        1,
        "Pair Programming and Debugging",
        "Checklist during Python/Scratch/micro:bit prototype work.",
        (
            "Used driver, navigator, tester, or recorder roles responsibly.",
            "Communicated respectfully while debugging.",
            "Participated with effort and perseverance.",
            "Kept testing notes or screenshots organized.",
            "Used feedback to improve one prototype feature.",
        ),
        "Which debugging habit helped your prototype work better?",
    ),
    AppreciationDoc(
        "8th",
        "IIIT",
        9,
        2,
        "Project Readiness",
        "Self-reflection about responsibility, organization, effort, safe device use, feedback, and improvement.",
        (
            "Prepared the final game plan and required evidence.",
            "Used Scratch Link, micro:bit, or simulator safely and responsibly.",
            "Stayed organized with files, screenshots, and test cases.",
            "Responded to feedback before project building.",
            "Identified one realistic improvement goal for the game.",
        ),
        "What must be ready before the Scratch + micro:bit exam project begins?",
    ),
    AppreciationDoc(
        "9th",
        "IT",
        3,
        1,
        "Lab Safety and Responsibility",
        "Checklist about preparation, careful material use, listening, cleanup, and persistence.",
        (
            "Came prepared with required materials and setup.",
            "Handled electronics, wires, boards, and workspace carefully.",
            "Listened to safety and build instructions.",
            "Cleaned up and returned materials correctly.",
            "Persisted through one test or fix and recorded evidence.",
        ),
        "What lab safety habit should you keep using next class?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Preparation and setup",
                "Arrives prepared and sets up materials, board, wires, and workspace correctly.",
                "Usually prepared with one minor setup issue.",
                "Needs reminders for materials or setup.",
                "Not prepared or setup cannot be checked.",
            ),
            (
                "Careful material use",
                "Handles electronics, wires, boards, and workspace safely and carefully.",
                "Uses materials safely with one minor reminder.",
                "Needs reminders for careful handling or safety.",
                "Unsafe material use or no safety evidence.",
            ),
            (
                "Instructions and cleanup",
                "Follows build instructions and cleans up/returns materials correctly.",
                "Usually follows instructions and cleans up.",
                "Needs reminders for instructions or cleanup.",
                "Does not follow instructions or clean up.",
            ),
            (
                "Persistence and evidence",
                "Persists through one test or fix and records useful evidence.",
                "Keeps trying and records basic evidence.",
                "Testing or fix evidence is incomplete.",
                "No useful test or fix evidence.",
            ),
        ),
    ),
    AppreciationDoc(
        "9th",
        "IT",
        7,
        2,
        "Teamwork and Feedback",
        "Checklist about communication, shared responsibility, respectful testing, organization, and feedback.",
        (
            "Communicated respectfully during robotics practice.",
            "Shared team roles and responsibilities fairly.",
            "Tested carefully and recorded useful evidence.",
            "Stayed organized with files, code, diagrams, or notes.",
            "Used feedback to improve a light, sound, or robotics system.",
        ),
        "What feedback changed your team project or test result?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Communication",
                "Communicates respectfully during robotics practice and testing.",
                "Usually communicates respectfully.",
                "Needs reminders for respectful communication.",
                "Communication blocks teamwork.",
            ),
            (
                "Shared responsibility",
                "Shares builder, coder, tester, or recorder roles fairly.",
                "Usually shares roles fairly.",
                "Roles are uneven or unclear.",
                "Does not share team responsibility.",
            ),
            (
                "Testing evidence",
                "Tests carefully and records useful evidence for the light, sound, or robotics system.",
                "Testing evidence is mostly complete.",
                "Testing evidence is partial or unclear.",
                "Testing evidence is missing.",
            ),
            (
                "Feedback and improvement",
                "Uses feedback to improve the system or explain a better test result.",
                "Uses feedback for a basic improvement.",
                "Feedback response is unclear.",
                "No feedback response or improvement.",
            ),
        ),
    ),
    AppreciationDoc(
        "9th",
        "IIT",
        6,
        1,
        "Data Cleaning Collaboration and Care",
        "Checklist about responsibility, communication, shared work, data care, and persistence during data cleaning and change-log work.",
        (
            "Handled data responsibly and avoided careless changes.",
            "Communicated clearly during data cleaning or review.",
            "Shared work fairly with partners or team members.",
            "Stayed organized with dataset, chart, or notes.",
            "Persisted through errors, missing values, or unclear results.",
        ),
        "What data-care choice made the dataset more trustworthy?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Responsible data handling",
                "Handles data carefully and avoids careless or harmful changes.",
                "Usually handles data responsibly.",
                "Needs reminders for data care.",
                "Data handling is careless or unsafe.",
            ),
            (
                "Data-cleaning collaboration",
                "Communicates clearly and shares data-cleaning work fairly.",
                "Usually communicates and shares work fairly.",
                "Shared work or communication is uneven.",
                "Does not collaborate responsibly.",
            ),
            (
                "Data-cleaning evidence",
                "Records at least three cleaned values or clear change-log evidence.",
                "Records most required change evidence.",
                "Change evidence is partial or hard to check.",
                "Change evidence is missing.",
            ),
            (
                "Organization and persistence",
                "Keeps dataset, chart, notes, or model evidence organized while solving errors.",
                "Mostly organized and keeps trying.",
                "Organization or persistence needs support.",
                "Evidence is disorganized or incomplete.",
            ),
        ),
    ),
    AppreciationDoc(
        "9th",
        "IIT",
        9,
        2,
        "Data Project Readiness and Planning",
        "Self-reflection about effort, organization, responsible data use, deadlines, and response to feedback during data project planning.",
        (
            "Used class time well during the data unit.",
            "Kept files, datasets, charts, and notes organized.",
            "Used data responsibly and protected privacy.",
            "Met deadlines or made a clear make-up plan.",
            "Responded to feedback before the data exam project.",
        ),
        "What organization habit will help your data project most?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Data project planning effort",
                "Uses class time well and stays focused during data project planning.",
                "Usually focused and productive.",
                "Needs reminders to use time well.",
                "Does not use class time productively.",
            ),
            (
                "Project evidence organization",
                "Files, datasets, charts, notes, and project plan are organized.",
                "Evidence is mostly organized.",
                "Organization is partial or hard to check.",
                "Organization evidence is missing.",
            ),
            (
                "Responsible data use",
                "Uses data responsibly, protects privacy, and names dataset source.",
                "Mostly responsible data use.",
                "Needs reminders for source or privacy care.",
                "Data use is irresponsible or unclear.",
            ),
            (
                "Deadlines, feedback, and plan",
                "Project plan meets deadlines and uses feedback before the data exam project.",
                "Mostly on time and responds to feedback.",
                "Deadline or feedback response is incomplete.",
                "No project plan, deadline plan, or feedback response.",
            ),
        ),
    ),
    AppreciationDoc(
        "9th",
        "IIIT",
        4,
        1,
        "Collaboration and Digital Responsibility",
        "Checklist about respectful communication, source care, privacy, participation, and safe account behavior.",
        (
            "Communicated respectfully during cybersecurity and STEM work.",
            "Handled sources and online information responsibly.",
            "Protected privacy and used accounts safely.",
            "Participated actively in group or class tasks.",
            "Used feedback to improve a comparison, checklist, or plan.",
        ),
        "Which digital responsibility habit protects people or data?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Respectful communication",
                "Communicates respectfully during cybersecurity and STEM work.",
                "Usually communicates respectfully.",
                "Needs reminders for respectful communication.",
                "Communication is unsafe or disrespectful.",
            ),
            (
                "Source care",
                "Handles sources and online information responsibly and checks credibility.",
                "Mostly responsible with sources.",
                "Source care is partial or unclear.",
                "Source care evidence is missing.",
            ),
            (
                "Privacy and account safety",
                "Protects privacy and uses accounts safely during class tasks.",
                "Mostly protects privacy and accounts.",
                "Needs reminders for privacy or account safety.",
                "Unsafe account or privacy behavior.",
            ),
            (
                "Participation and improvement",
                "Participates actively and uses feedback to improve work.",
                "Usually participates and responds to feedback.",
                "Participation or feedback response is incomplete.",
                "No useful participation or improvement evidence.",
            ),
        ),
    ),
    AppreciationDoc(
        "9th",
        "IIIT",
        9,
        2,
        "STEM Project Habits",
        "Self-reflection about effort, organization, safe material use, perseverance, and response to feedback during STEM planning.",
        (
            "Showed effort during project planning and preparation.",
            "Kept the logbook, files, materials, or task list organized.",
            "Used materials or tools safely.",
            "Persevered through design, testing, or planning difficulty.",
            "Responded to feedback before project building.",
        ),
        "What STEM project habit will you focus on during the build?",
        special_form="official_appreciation_template",
        template_criteria=(
            (
                "Effort and planning",
                "Shows effort and completes task list, team roles, build schedule, and outline.",
                "Planning evidence is mostly complete.",
                "Planning evidence is partial or unclear.",
                "Planning evidence is missing.",
            ),
            (
                "Organization",
                "Keeps logbook, files, materials, or project evidence organized.",
                "Evidence is mostly organized.",
                "Organization is inconsistent.",
                "Evidence is missing or disorganized.",
            ),
            (
                "Safe material/tool use",
                "Uses project materials, tools, or approved software safely.",
                "Mostly uses materials/tools safely.",
                "Needs reminders for safe tool use.",
                "Unsafe material or tool use.",
            ),
            (
                "Perseverance and feedback",
                "Persists through design difficulty and responds to feedback before building.",
                "Usually keeps trying and uses feedback.",
                "Perseverance or feedback response is incomplete.",
                "No useful feedback response or persistence evidence.",
            ),
        ),
    ),
)


EXAM_DOCS: tuple[ExamDoc, ...] = (
    ExamDoc(
        "7th",
        "IT",
        "mBot Maze Navigator",
        "Students program an mBot to navigate a maze or best available maze section using movement, sensors, testing, debugging, and a short explanation.",
        (
            "Maze sketch or route plan with start, finish, checkpoints, and obstacles.",
            "mBlock program evidence or screenshots.",
            "Testing notes showing at least two trials and one improvement.",
            "Final robot demonstration or recorded/photo evidence if live testing is interrupted.",
            "Short reflection about one challenge and one solution.",
        ),
        (
            "Week 10: create project plan and begin build.",
            "Week 11: test, debug, document, and improve.",
            "Week 12: demonstrate, explain, and submit reflection.",
        ),
        (
            "Demonstrate the robot or best available maze section.",
            "Explain how code controls movement and how sensors support decisions.",
            "Name one problem, one fix, and one realistic improvement.",
        ),
        (
            ("Programming and logic", "Uses sequence, loops, conditions, and sensor logic effectively.", "Uses basic sequence, loops, or conditions with minor gaps.", "Program moves the robot but logic is incomplete.", "Program is missing, unclear, or does not run."),
            ("Robot performance", "Robot completes most of the route smoothly and safely.", "Robot completes part of the route with some issues.", "Robot moves but often misses the route or needs help.", "Robot cannot complete a meaningful test."),
            ("Testing and debugging", "Records trials, identifies problems, and improves the solution.", "Records some testing and makes at least one fix.", "Testing notes are limited or fixes are unclear.", "Little or no testing evidence is submitted."),
            ("Project documentation", "Plan, code evidence, and reflection are clear and complete.", "Most required evidence is present.", "Some evidence is missing or hard to check.", "Evidence is mostly missing."),
            ("Presentation and understanding", "Explains code, sensors, challenge, result, and improvement clearly.", "Explains most project parts with minor gaps.", "Gives a short or partly unclear explanation.", "Cannot explain how the project works."),
            ("Responsibility", "Works safely, uses class time well, and meets deadlines.", "Usually prepared and responsible.", "Needs reminders for readiness or deadlines.", "Does not use time or materials responsibly."),
        ),
    ),
    ExamDoc(
        "7th",
        "IIT",
        "Scratch Dance Game",
        "Students create an interactive Scratch dance game where the player uses input at the right time and receives clear feedback.",
        (
            "Game plan with goal, keys, moves, variables, subroutines, and roles.",
            "Scratch project with at least two dance move routines.",
            "Input, conditionals, loops, variables, feedback, score or timer evidence.",
            "Testing checklist with one bug and one improvement.",
            "Final reflection and project link or screenshot evidence.",
        ),
        (
            "Week 10: set up project and first subroutine.",
            "Week 11: complete required features, test, debug, and prepare explanation.",
            "Week 12: demonstrate the game and submit reflection.",
        ),
        (
            "Demonstrate the game from start to finish.",
            "Explain subroutines, input, conditionals, loops, variables, and feedback.",
            "Explain one real-world application or benefit of interactive games.",
        ),
        (
            ("Game design and application", "Goal, audience, and real-world application are clear and specific.", "Goal and application are present with some detail.", "Goal is basic or application is vague.", "Goal/application is missing."),
            ("Subroutines and organization", "Code is well organized with reusable subroutines.", "Code has some organization and subroutines.", "Code organization is limited.", "Code is disorganized or missing."),
            ("Functionality", "Moves, timing, input, score/timer, and feedback work reliably.", "Most features work with minor issues.", "Some features work but gameplay is limited.", "Game cannot be played or checked."),
            ("Programming concepts", "Uses input, conditions, loops, variables, and events correctly.", "Uses most required concepts.", "Uses a few concepts with errors.", "Required concepts are missing."),
            ("Testing and improvement", "Testing evidence leads to a clear improvement.", "Some testing and improvement are shown.", "Testing is limited or improvement is unclear.", "Little or no testing evidence."),
            ("Presentation and responsibility", "Clear demo, explanation, reflection, and timely submission.", "Mostly clear presentation and submission.", "Presentation or reflection is incomplete.", "Presentation or submission is missing."),
        ),
    ),
    ExamDoc(
        "7th",
        "IIIT",
        "Mandrake Detection System",
        "Students build or simulate a micro:bit sensor system that detects when the Mandrake is removed and triggers an output signal.",
        (
            "Project plan with purpose, parts, input/output diagram, threshold rule, code plan, and test plan.",
            "micro:bit or simulator evidence showing sensor input and output response.",
            "Testing table with distance values or sample readings.",
            "Canva or Google Slides support slide if used.",
            "Short reflection about challenge, result, and improvement.",
        ),
        (
            "Week 10: begin project setup and sensor logic.",
            "Week 11: debug reliability and prepare presentation.",
            "Week 12: present final Mandrake Detection System.",
        ),
        (
            "Explain purpose, parts, sensor reading, threshold rule, signal/output, challenge, result, and improvement.",
            "Use a clearly labeled simulation if equipment is unavailable.",
        ),
        (
            ("Project purpose", "Purpose and Mandrake behavior are explained clearly.", "Purpose is explained with some detail.", "Purpose is basic or partly unclear.", "Purpose is missing."),
            ("System parts and input/output", "micro:bit, sensor, threshold, signal, and output are connected accurately.", "Most system parts are explained.", "Some parts are named but not connected.", "System explanation is missing or incorrect."),
            ("Code and logic", "Code plan uses sensor input and threshold logic correctly.", "Code mostly matches the intended behavior.", "Code is incomplete or only partly connected.", "Code evidence is missing."),
            ("Testing and reliability", "Test evidence supports threshold choice and improvement.", "Some test evidence is present.", "Testing is limited or unclear.", "No meaningful test evidence."),
            ("Result and improvement", "Final result is shown with a realistic improvement.", "Result is shown with a simple improvement.", "Result or improvement is vague.", "Result is not shown."),
            ("Presentation and responsibility", "Clear presentation, reflection, preparation, and safe material use.", "Mostly complete presentation and responsible work.", "Presentation or readiness is incomplete.", "Presentation/submission is missing."),
        ),
    ),
    ExamDoc(
        "8th",
        "IT",
        "Google Sites Environmental Campaign Portfolio",
        "Students build a Google Sites portfolio that presents an environmental campaign with original vector graphics, required media, source credits, testing, and reflection.",
        (
            "Portfolio map with at least five pages.",
            "At least three original vector graphics.",
            "Required environmental video location, storyboard, or evidence.",
            "Two source credits and one call to action.",
            "Final site link, testing notes, reflection, and presentation notes.",
        ),
        (
            "Week 10: homepage, campaign purpose, and media integration.",
            "Week 11: testing, feedback, revisions, credits, and reflection.",
            "Week 12: portfolio presentation and final submission.",
        ),
        (
            "Show the site navigation and required pages.",
            "Explain original graphics, media choices, source credits, one design choice, and one improvement.",
        ),
        (
            ("Campaign purpose and content", "Campaign message, audience, pages, and call to action are clear.", "Most content is clear and connected.", "Content is basic or partly disconnected.", "Campaign purpose is missing."),
            ("Vector graphics", "At least three original graphics are complete and well integrated.", "Original graphics are present with minor gaps.", "Graphics are limited or not clearly original.", "Required graphics are missing."),
            ("Website structure and design", "Navigation, layout, readability, and consistency support the audience.", "Site structure mostly works.", "Navigation or design is confusing in places.", "Site is hard to navigate or incomplete."),
            ("Media, credits, and accessibility", "Video/media, captions, alt text, and credits are complete.", "Most media and source requirements are met.", "Some source or accessibility evidence is missing.", "Media/credit evidence is mostly missing."),
            ("Testing and reflection", "Peer testing leads to at least three clear fixes and a thoughtful reflection.", "Testing and reflection are present.", "Testing or reflection is limited.", "Little or no revision evidence."),
            ("Presentation and responsibility", "Clear presentation, timely submission, and organized files.", "Mostly complete presentation and submission.", "Presentation/submission is incomplete.", "Final link or presentation is missing."),
        ),
    ),
    ExamDoc(
        "8th",
        "IIT",
        "Arduino Freenove Prototype",
        "Students design, build, test, debug, and demonstrate an Arduino/Freenove prototype where one input controls at least one output.",
        (
            "Arduino project design plan with purpose, role, input, output, diagram, behavior, tests, and materials.",
            "Prototype circuit evidence and starter/final sketch evidence.",
            "Testing log with at least three conditions.",
            "Debugging note with problem, likely cause, and solution.",
            "Final diagram, code logic explanation, demonstration, and reflection.",
        ),
        (
            "Week 10: build first circuit version and program input/output behavior.",
            "Week 11: debug, improve, document, and collect final testing evidence.",
            "Week 12: demonstrate prototype and submit evidence.",
        ),
        (
            "Demonstrate the prototype or provide recorded/photo evidence.",
            "Identify the input, process, output, one code decision, one test result, and one improvement.",
        ),
        (
            ("Design plan", "Purpose, input, output, diagram, behavior, tests, and materials are complete.", "Most plan parts are complete.", "Plan is partial or unclear.", "Plan is missing."),
            ("Circuit and build quality", "Circuit is safe, labeled, and works reliably.", "Circuit works with minor support.", "Circuit partly works or labels are incomplete.", "Circuit cannot be checked."),
            ("Code and input/output logic", "Code clearly connects input to output with correct logic.", "Code mostly controls the output.", "Code is incomplete or only partly connected.", "Code evidence is missing."),
            ("Testing and debugging", "Three conditions, issue evidence, fix, and improvement are clear.", "Some testing and debugging are shown.", "Testing is limited or unclear.", "No meaningful testing evidence."),
            ("Documentation and explanation", "Diagram, code explanation, photos/screenshots, and reflection are complete.", "Most documentation is present.", "Documentation has gaps.", "Documentation is mostly missing."),
            ("Demo and responsibility", "Clear demo, safe material use, cleanup, and timely submission.", "Mostly responsible work and demo.", "Demo or responsibility evidence is incomplete.", "Demo/submission is missing."),
        ),
    ),
    ExamDoc(
        "8th",
        "IIIT",
        "Scratch + micro:bit Game",
        "Students create a Scratch game controlled by micro:bit input or a clearly labeled simulator setup, with game logic, testing, debugging, and presentation evidence.",
        (
            "Game plan with title, role, goal, controls, scoring or win condition, sprites, and three test cases.",
            "Scratch game shell and micro:bit input mapping evidence.",
            "Game loop, sensor-controlled movement, scoring/success condition, and challenge or obstacle.",
            "Peer test notes and at least one fix.",
            "Screenshots or live demo evidence, speaking notes, and reflection.",
        ),
        (
            "Week 10: build game shell and sensor-controlled movement.",
            "Week 11: user testing, debugging, documentation, and presentation prep.",
            "Week 12: present game and submit final evidence.",
        ),
        (
            "Demonstrate controls, game logic, challenge, result, and solution.",
            "Use Scratch Link, a physical micro:bit, or a clearly labeled simulator option.",
        ),
        (
            ("Game design", "Goal, controls, scoring/win condition, sprites, and challenge are clear.", "Most design requirements are present.", "Design is partial or unclear.", "Design evidence is missing."),
            ("micro:bit input", "micro:bit or simulator input controls the game clearly.", "Input works with minor issues.", "Input is attempted but unreliable.", "Input evidence is missing."),
            ("Scratch code and game logic", "Events, movement, loop, condition, and feedback work together.", "Most code logic works.", "Some code works but game logic is limited.", "Game cannot be checked."),
            ("Testing and debugging", "Peer test notes and fixes clearly improve the game.", "Some testing and at least one fix are shown.", "Testing is limited or unclear.", "No meaningful testing evidence."),
            ("Documentation and reflection", "Screenshots, notes, and reflection explain choices and improvement.", "Most documentation is present.", "Documentation has gaps.", "Documentation is mostly missing."),
            ("Presentation and responsibility", "Clear demo, explanation, timely submission, and respectful audience behavior.", "Mostly complete presentation and submission.", "Presentation is incomplete.", "Presentation/submission is missing."),
        ),
    ),
    ExamDoc(
        "9th",
        "IT",
        "Robotics Project",
        "Students build, test, debug, and demonstrate an Arduino/Freenove robotics prototype that uses inputs, outputs, state logic, documentation, and reflection.",
        (
            "Robotics design plan with goal, inputs/controls, outputs, state chart, components, and tests.",
            "Prototype progress, code/screenshots, project log, test results, and fixes.",
            "Final demo evidence and reflection.",
        ),
        (
            "Weeks 10-11: build, test, debug, improve, and document.",
            "Week 12: demonstrate prototype and submit final log/reflection.",
            "Week 13: make-up or admin buffer if needed.",
        ),
        (
            "Demonstrate goal, input, output, state logic, testing evidence, and improvement.",
            "Use recorded demo or photo sequence if live presentation is interrupted.",
        ),
        (
            ("Design and state logic", "Goal, inputs, outputs, state chart, components, and tests are complete.", "Most design elements are complete.", "Design is partial or unclear.", "Design evidence is missing."),
            ("Prototype and code", "Prototype and code work together reliably.", "Prototype mostly works.", "Prototype partly works or needs support.", "Prototype cannot be checked."),
            ("Testing and debugging", "Testing evidence, issues, fixes, and improvement are clear.", "Some testing and fixes are shown.", "Testing is limited.", "No meaningful testing evidence."),
            ("Project log and documentation", "Log, screenshots/code, photos, and reflection are complete.", "Most documentation is present.", "Documentation has gaps.", "Documentation is mostly missing."),
            ("Demo and explanation", "Explains goal, logic, evidence, and improvement clearly.", "Explains most project parts.", "Explanation is short or unclear.", "Demo/explanation is missing."),
            ("Responsibility", "Safe, organized, prepared, and timely throughout project work.", "Usually responsible and prepared.", "Needs reminders for organization or deadlines.", "Does not use time/materials responsibly."),
        ),
    ),
    ExamDoc(
        "9th",
        "IIT",
        "Data Project",
        "Students clean or prepare a dataset, build a data product or report, test clarity, revise conclusions, and present evidence-based findings.",
        (
            "Approved project question, dataset source, chart plan, product format, task list, and presentation outline.",
            "Clean dataset, chart/visualization, data product/report/dashboard, and testing feedback.",
            "Conclusion, limitation, recommendation, final reflection, and presentation evidence.",
        ),
        (
            "Weeks 10-11: clean data, build chart/product, test clarity, revise evidence, and prepare reflection.",
            "Week 12: present data question, method, chart, conclusion, limitation, recommendation, and product.",
            "Week 13: make-up or admin buffer if needed.",
        ),
        (
            "Present the data question, method, chart, conclusion, limitation, recommendation, final product, and reflection.",
            "Use backup dataset if approved data is unusable.",
        ),
        (
            ("Question and method", "Question, dataset source, method, and privacy/data-care choices are clear.", "Most method details are clear.", "Method is partial or unclear.", "Question/method is missing."),
            ("Data product and chart", "Clean dataset, chart, and product are accurate and readable.", "Product mostly communicates the data.", "Product is incomplete or hard to read.", "Product cannot be checked."),
            ("Interpretation", "Conclusion, limitation, recommendation, and evidence statements match the data.", "Most interpretation matches the data.", "Interpretation is basic or partly unsupported.", "Interpretation is missing or unsupported."),
            ("Testing and revision", "Feedback leads to clear fixes in labels, calculations, or clarity.", "Some testing and revisions are shown.", "Testing is limited.", "No meaningful testing evidence."),
            ("Presentation and reflection", "Presentation explains the data story clearly and includes thoughtful reflection.", "Presentation covers most required parts.", "Presentation is short or incomplete.", "Presentation/reflection is missing."),
            ("Responsibility", "Responsible data use, organization, deadlines, and effort are evident.", "Usually organized and responsible.", "Needs reminders for organization or deadlines.", "Work habits limit project evidence."),
        ),
    ),
    ExamDoc(
        "9th",
        "IIIT",
        "STEM Project",
        "Students build a STEM prototype or digital product, test it against criteria, improve it, document evidence, and demonstrate the final result.",
        (
            "Final plan with task list, team roles, build schedule, testing plan, and presentation outline.",
            "Prototype or digital product progress, logbook, testing evidence, screenshots/photos, and improvement.",
            "Final product evidence, reflection, and presentation notes.",
        ),
        (
            "Weeks 10-11: build prototype, test features, identify weakness, improve, and update logbook.",
            "Week 12: demonstrate STEM product and submit final evidence.",
            "Week 13: make-up, final reflection corrections, or portfolio cleanup if needed.",
        ),
        (
            "Demonstrate the STEM product or prototype.",
            "Explain problem, users, tools, design choices, testing evidence, improvement, limitation, and reflection.",
        ),
        (
            ("Problem and design", "Problem, users, goal, tools, sketch/model, materials, and test criteria are clear.", "Most design elements are clear.", "Design is partial or broad.", "Design evidence is missing."),
            ("Prototype or product", "Prototype/product addresses the problem and shows the planned feature clearly.", "Prototype/product mostly works.", "Prototype/product is partial.", "Prototype/product cannot be checked."),
            ("Testing evidence", "At least three trials or clear evidence compare results to criteria.", "Some testing evidence is present.", "Testing is limited or unclear.", "No meaningful testing evidence."),
            ("Improvement and iteration", "Improvement responds to evidence and is explained clearly.", "Improvement is present with some explanation.", "Improvement is vague or not evidence-based.", "No improvement evidence."),
            ("Logbook and reflection", "Logbook, screenshots/photos, notes, and reflection are complete.", "Most documentation is present.", "Documentation has gaps.", "Documentation is mostly missing."),
            ("Presentation and responsibility", "Clear demo, explanation, safe material use, and timely submission.", "Mostly complete presentation and responsible work.", "Presentation or responsibility evidence is incomplete.", "Presentation/submission is missing."),
        ),
    ),
)


RUBRIC_DOC = RubricDoc(
    "8th",
    "IIT",
    1,
    1,
    "Arduino Basics Vocabulary Table",
    "Vocabulary table with definitions, illustrations, and example sentences for Arduino, Freenove, microcontroller, circuit, breadboard, pin, power, ground, resistor, and LED.",
    (
        (
            "Vocabulary completion and definitions",
            "All 10 words have accurate, student-friendly definitions.",
            "Most definitions are accurate.",
            "Some definitions are missing or unclear.",
            "Definitions are mostly missing or incorrect.",
        ),
        (
            "Illustrations or diagrams",
            "Each word has a relevant illustration or diagram.",
            "Most illustrations match the words.",
            "Some illustrations are missing or unclear.",
            "Illustrations are mostly missing or unrelated.",
        ),
        (
            "Example sentences",
            "Sentences use each word in a clear Arduino or Freenove context.",
            "Most sentences use the words correctly.",
            "Some sentences are incomplete or vague.",
            "Sentences are mostly missing or incorrect.",
        ),
        (
            "Hardware accuracy",
            "Board, circuit, breadboard, pin, power, ground, resistor, and LED ideas are correct.",
            "Most hardware ideas are correct.",
            "Some hardware ideas are confused.",
            "Hardware ideas are mostly incorrect.",
        ),
    ),
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate secondary assessment DOCX documents.")
    parser.add_argument("--dry-run", action="store_true", help="Print planned outputs without writing files.")
    parser.add_argument("--render", action="store_true", help="Render generated primary DOCX files for visual QA.")
    parser.add_argument("--render-docx", default=str(RENDER_DOCX), help="Path to render_docx.py.")
    parser.add_argument("--force", action="store_true", help="Overwrite existing generated outputs.")
    return parser.parse_args()


def slug_filename(value: str) -> str:
    value = value.replace("/", " ")
    value = re.sub(r"[^A-Za-z0-9+:. -]+", "", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def grade_number(grade: str) -> str:
    return grade[:-2]


def grade_folder(grade: str) -> Path:
    return GRADE_FOLDERS[grade]


def trimester_label(code: str) -> str:
    return TRIMESTERS[code]


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 12) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_lines(cell, lines: Iterable[str], bold: bool = False, align: int | None = None) -> None:
    cell.text = ""
    paragraphs = cell.paragraphs
    for index, line in enumerate(lines):
        paragraph = paragraphs[0] if index == 0 else cell.add_paragraph()
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(12)
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_template_rubric_row(row, values: tuple[str, str, str, str, str]) -> None:
    # The official template merges duplicate 7-point and 5-point cells.
    for cell_index, value in zip((0, 1, 2, 4, 6), values):
        set_cell_text(row.cells[cell_index], value)


def set_template_punctuality_row(row) -> None:
    set_cell_lines(row.cells[0], ("Punctuality,", "Readiness &", "Respect"), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], "On time, prepared, and has needed resources.")
    set_cell_text(row.cells[3], "Submitted, but late/incomplete or missing a needed resource.")
    set_cell_text(row.cells[5], "Not submitted on time or not prepared.")


def enforce_arial_12(doc: Document) -> None:
    blocks = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(cell.paragraphs)
    for paragraph in blocks:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            if not run.text:
                continue
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run_properties = run._element.get_or_add_rPr()
            run_fonts = run_properties.find(qn("w:rFonts"))
            if run_fonts is None:
                run_fonts = OxmlElement("w:rFonts")
                run_properties.append(run_fonts)
            for key in ("w:ascii", "w:hAnsi", "w:cs"):
                run_fonts.set(qn(key), "Arial")
            size_cs = run_properties.find(qn("w:szCs"))
            if size_cs is None:
                size_cs = OxmlElement("w:szCs")
                run_properties.append(size_cs)
            size_cs.set(qn("w:val"), "24")


def format_document(doc: Document, *, landscape: bool = False) -> None:
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, size: int = 14) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(0)


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        run = paragraph.add_run("- " + item)
        run.font.name = "Arial"
        run.font.size = Pt(12)


def add_header_block(doc: Document, title: str, grade: str, trimester: str, subtitle: str, score: str) -> None:
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(0.42))
    add_heading(doc, "Academia Internacional David", 13)
    add_heading(doc, "Robotics and Technology", 12)
    add_heading(doc, title, 14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{grade} Grade - {trimester_label(trimester)} - {subtitle}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(0)

    info = doc.add_table(rows=2, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = True
    cells = info.rows[0].cells + info.rows[1].cells
    values = [
        "Name: ________________________________",
        "Date: ______________________",
        f"Teacher: Porfirio Rios     Group: {grade_number(grade)}th ______",
        f"Score: {score}",
    ]
    for cell, value in zip(cells, values):
        set_cell_text(cell, value, bold=False)


def add_appreciation_table(doc: Document, criteria: tuple[str, ...]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Criterion", "4 pts", "3 pts", "2 pts", "1 pt"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
    for criterion in criteria:
        row = table.add_row()
        values = [
            criterion,
            "Consistently shown",
            "Usually shown",
            "Sometimes shown",
            "Needs support",
        ]
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)


def add_constructive_feedback_appreciation_rubric(doc: Document) -> None:
    add_section_heading(doc, "Rubric")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Criteria", "4 pts", "3 pts", "2 pts", "1 pt"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")

    rows = [
        (
            "Use of class time",
            "Starts promptly, stays focused, and uses coding and feedback time responsibly.",
            "Usually focused and uses time well, with only one small reminder.",
            "Needs several reminders; some class time is not used for the assigned work.",
            "Often off task or does not complete the assigned work during class.",
        ),
        (
            "Organization and readiness",
            "Materials, login, Scratch file, and evidence are ready, saved, and easy to check.",
            "Mostly prepared and organized, with one minor file, material, or submission issue.",
            "Organization issues make evidence harder to check or require teacher support.",
            "Missing readiness, organization, or submission evidence.",
        ),
        (
            "Scratch practice evidence",
            "Program evidence shows a loop, variable, timing test, and one adjusted value.",
            "Program shows most required parts, with one minor missing or unclear detail.",
            "Program is started, but loop, variable, test, or adjustment evidence is incomplete.",
            "Program evidence is missing or too unclear to support the grade.",
        ),
        (
            "Constructive classmate feedback",
            "Feedback names an observed strength and a specific suggestion in a respectful tone.",
            "Feedback is respectful and useful but needs more specific evidence or explanation.",
            "Feedback is mostly general praise or unclear advice with limited evidence.",
            "Feedback is missing, disrespectful, or does not help the classmate improve.",
        ),
        (
            "Participation and response to feedback",
            "Actively evaluates a classmate and records the response, change, or next step.",
            "Completes the evaluation and records a basic response, change, or next step.",
            "Evaluation or response is incomplete and needs teacher support.",
            "Does not complete the peer-evaluation process.",
        ),
    ]
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)


def add_constructive_feedback_questions(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Constructive Peer Feedback Questions", 14)
    p = doc.add_paragraph(
        "Constructive criticism is specific, respectful, and useful. Use evidence from the program, then give one clear suggestion."
    )
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    rows = [
        ("Reviewer", "Name: ________________________________"),
        ("Classmate evaluated", "Name: ________________________________"),
        ("My program evidence", "Link/screenshot: ____________________  Value adjusted: ____________________"),
        ("Code evidence observed", "[ ] loop  [ ] variable  [ ] timing test  [ ] adjusted value"),
        ("Specific strength", "I noticed that ______________________________________________________________"),
        ("Helpful suggestion", "I suggest __________________________________________________ because ______________________________"),
        ("Constructive question", "Have you tried ______________________________________________________________?"),
        ("Classmate response/change", "After the feedback, my classmate ____________________________________________"),
    ]
    for row, (label, prompt) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], "D9EAF7")
        set_cell_text(row.cells[1], prompt)


def add_pair_programming_appreciation_rubric(doc: Document) -> None:
    add_section_heading(doc, "Rubric")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Criteria", "4 pts", "3 pts", "2 pts", "1 pt"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")

    rows = [
        (
            "Use of class time",
            "Starts promptly, stays focused, and uses coding and reflection time responsibly.",
            "Usually focused and uses time well, with only one small reminder.",
            "Needs several reminders; some class time is not used for the assigned work.",
            "Often off task or does not complete the assigned work during class.",
        ),
        (
            "Organization and readiness",
            "Login, Scratch file, screenshot, and reflection evidence are ready and easy to check.",
            "Mostly prepared and organized, with one minor file, material, or submission issue.",
            "Organization issues make evidence harder to check or require teacher support.",
            "Missing readiness, organization, or submission evidence.",
        ),
        (
            "Driver/navigator roles",
            "Uses driver and navigator roles fairly and explains personal contribution clearly.",
            "Uses roles mostly fairly, with only a small imbalance or unclear detail.",
            "Roles are attempted, but one partner does most of the work or the evidence is unclear.",
            "Roles are not used responsibly or cannot be explained.",
        ),
        (
            "Scratch dance practice evidence",
            "Feature shows one key press, costume change, sound, movement, and success/failure response.",
            "Feature shows most required parts, with one minor missing or unclear detail.",
            "Feature is started, but several required parts are missing or hard to check.",
            "Practice feature is missing or too incomplete to support the grade.",
        ),
        (
            "Communication and response to feedback",
            "Communicates respectfully, listens, and records feedback or a next improvement step.",
            "Communicates respectfully and records a basic feedback response or next step.",
            "Communication or feedback response is incomplete and needs teacher support.",
            "Does not show respectful communication or a feedback response.",
        ),
    ]
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)


def add_pair_programming_reflection_questions(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Pair-Programming Self-Reflection Questions", 14)
    p = doc.add_paragraph(
        "Use evidence from the Scratch practice feature and your partner work. Your answers should support the appreciation grade."
    )
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    rows = [
        ("Student", "Name: ________________________________"),
        ("Partner", "Name: ________________________________"),
        ("Roles used", "I was: [ ] driver  [ ] navigator  [ ] both. Evidence: ______________________________"),
        ("Practice feature evidence", "[ ] key press  [ ] costume change  [ ] sound  [ ] movement  [ ] success/failure response"),
        ("My contribution", "I helped by ______________________________________________________________"),
        ("Partner contribution", "My partner helped by ______________________________________________________"),
        ("Communication", "One respectful communication habit we used was ______________________________"),
        ("Organization", "Our file/screenshot/reflection evidence is saved here: ________________________"),
        ("Feedback and next step", "Feedback we used or will use next: _________________________________________"),
    ]
    for row, (label, prompt) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], "D9EAF7")
        set_cell_text(row.cells[1], prompt)


def save_doc(doc: Document, path: Path) -> None:
    ensure_parent(path)
    doc.save(str(path))


def create_official_appreciation_template_doc(spec: AppreciationDoc, path: Path) -> None:
    if len(spec.template_criteria) != 4:
        raise ValueError(f"Official appreciation template requires exactly 4 criteria: {spec.title}")
    doc = Document(str(RUBRIC_TEMPLATE))
    table = doc.tables[0]
    set_cell_lines(
        table.rows[0].cells[0],
        (
            "Academia Internacional David",
            "Robotics and Technology",
            f"{trimester_label(spec.trimester)}",
            f"Appreciation Summative # {spec.number}",
            f"{spec.grade} A & B",
            f"Name: ________________________________      Date: __________________",
            f"Group: {grade_number(spec.grade)}° A B",
            "Teacher: Porfirio Rios                               Score: _____ / 40pts",
            spec.title,
        ),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for row_index, criterion in enumerate(spec.template_criteria, start=2):
        set_template_rubric_row(table.rows[row_index], criterion)
    set_template_punctuality_row(table.rows[7])
    enforce_arial_12(doc)
    save_doc(doc, path)


def create_appreciation_doc(spec: AppreciationDoc, path: Path) -> None:
    if spec.special_form == "official_appreciation_template":
        create_official_appreciation_template_doc(spec, path)
        return

    doc = Document()
    format_document(doc, landscape=True)
    add_header_block(
        doc,
        spec.title,
        spec.grade,
        spec.trimester,
        f"Week {spec.week} - Appreciation Grade #{spec.number}",
        "____ / 20 pts",
    )
    add_section_heading(doc, "Evidence Reminder")
    p = doc.add_paragraph(spec.evidence)
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)
    if spec.special_form == "constructive_peer_feedback":
        add_constructive_feedback_appreciation_rubric(doc)
        add_constructive_feedback_questions(doc)
    elif spec.special_form == "pair_programming_reflection":
        add_pair_programming_appreciation_rubric(doc)
        add_pair_programming_reflection_questions(doc)
    else:
        add_section_heading(doc, "Checklist")
        add_appreciation_table(doc, spec.criteria)
        add_section_heading(doc, "Short Reflection")
        p = doc.add_paragraph(spec.reflection)
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(12)
        p = doc.add_paragraph("Response: __________________________________________________________________________________")
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(12)
    save_doc(doc, path)


def add_exam_rubric(doc: Document, criteria: tuple[tuple[str, str, str, str, str], ...]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Criterion", "15-13 pts", "12-9 pts", "8-5 pts", "4-0 pts"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
    for criterion, excellent, good, developing, beginning in criteria:
        row = table.add_row()
        values = [criterion, excellent, good, developing, beginning]
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)


def create_exam_doc(spec: ExamDoc, path: Path) -> None:
    doc = Document()
    format_document(doc, landscape=True)
    add_header_block(
        doc,
        spec.title,
        spec.grade,
        spec.trimester,
        "Final Project Packet and Rubric",
        "____ / 90 pts",
    )
    add_section_heading(doc, "Project Overview")
    p = doc.add_paragraph(spec.overview)
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)
    add_section_heading(doc, "Required Evidence")
    add_bullets(doc, spec.evidence)
    add_section_heading(doc, "Schedule and Evidence Expectations")
    add_bullets(doc, spec.schedule)
    add_section_heading(doc, "Presentation or Demonstration Requirements")
    add_bullets(doc, spec.presentation)
    add_section_heading(doc, "Rubric")
    add_exam_rubric(doc, spec.criteria)
    save_doc(doc, path)


def create_rubric_doc(spec: RubricDoc, path: Path) -> None:
    doc = Document(str(RUBRIC_TEMPLATE))
    table = doc.tables[0]
    set_cell_lines(
        table.rows[0].cells[0],
        (
            "Academia Internacional David",
            "Robotics and Technology",
            f"{trimester_label(spec.trimester)}",
            f"Summative # {spec.summative}",
            f"{spec.grade} A & B",
            f"Name: ________________________________      Date: __________________",
            f"Group: {grade_number(spec.grade)}° A B",
            "Teacher: Porfirio Rios                               Score: _____ / 40pts",
            spec.title,
        ),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for row_index, criterion in enumerate(spec.criteria, start=2):
        set_template_rubric_row(table.rows[row_index], criterion)
    set_template_punctuality_row(table.rows[7])
    enforce_arial_12(doc)
    save_doc(doc, path)


def app_primary_path(spec: AppreciationDoc) -> Path:
    name = f"{spec.grade} grade - {spec.trimester} - Week {spec.week} - Appreciation Grade {spec.number} - {slug_filename(spec.title)}.docx"
    return grade_folder(spec.grade) / "Assessments" / "Appreciation Grades" / trimester_label(spec.trimester) / name


def exam_primary_path(spec: ExamDoc) -> Path:
    name = f"{spec.grade} grade - {spec.trimester} - Final Project Packet and Rubric - {slug_filename(spec.title)}.docx"
    return grade_folder(spec.grade) / "Assessments" / "Exam Projects" / trimester_label(spec.trimester) / name


def rubric_primary_path(spec: RubricDoc) -> Path:
    name = f"{spec.grade} grade - {spec.trimester} - Week {spec.week} - Rubric for Summative {spec.summative}.docx"
    return grade_folder(spec.grade) / "Assessments" / "Rubrics" / name


def assessment_mirror(primary: Path) -> Path:
    relative = primary.relative_to(PLANS_ROOT)
    return GENERATED_ROOT / relative


def generated_rubric_mirror(spec: RubricDoc) -> Path:
    return (
        GENERATED_RUBRICS_ROOT
        / f"{spec.grade} Grade Technology"
        / trimester_label(spec.trimester)
        / rubric_primary_path(spec).name
    )


def planned_outputs() -> list[OutputRecord]:
    records: list[OutputRecord] = []
    for spec in APPRECIATION_DOCS:
        primary = app_primary_path(spec)
        records.append(
            OutputRecord(
                "appreciation",
                spec.title,
                primary,
                (assessment_mirror(primary),),
                f"{spec.grade} {spec.trimester} Week {spec.week} Appreciation Grade {spec.number}",
            )
        )
    for spec in EXAM_DOCS:
        primary = exam_primary_path(spec)
        records.append(
            OutputRecord(
                "exam project",
                spec.title,
                primary,
                (assessment_mirror(primary),),
                f"{spec.grade} {spec.trimester} exam project plan",
            )
        )
    primary = rubric_primary_path(RUBRIC_DOC)
    records.append(
        OutputRecord(
            "rubric",
            RUBRIC_DOC.title,
            primary,
            (assessment_mirror(primary), generated_rubric_mirror(RUBRIC_DOC)),
            "8th IIT assessment map",
        )
    )
    return records


def copy_mirrors(primary: Path, mirrors: tuple[Path, ...]) -> None:
    for mirror in mirrors:
        ensure_parent(mirror)
        shutil.copy2(primary, mirror)


def generate(records: list[OutputRecord], *, force: bool) -> None:
    for spec in APPRECIATION_DOCS:
        path = app_primary_path(spec)
        if path.exists() and not force:
            raise FileExistsError(f"Refusing to overwrite existing file without --force: {path}")
        create_appreciation_doc(spec, path)
        copy_mirrors(path, (assessment_mirror(path),))

    for spec in EXAM_DOCS:
        path = exam_primary_path(spec)
        if path.exists() and not force:
            raise FileExistsError(f"Refusing to overwrite existing file without --force: {path}")
        create_exam_doc(spec, path)
        copy_mirrors(path, (assessment_mirror(path),))

    path = rubric_primary_path(RUBRIC_DOC)
    if path.exists() and not force:
        raise FileExistsError(f"Refusing to overwrite existing file without --force: {path}")
    create_rubric_doc(RUBRIC_DOC, path)
    copy_mirrors(path, (assessment_mirror(path), generated_rubric_mirror(RUBRIC_DOC)))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def open_docx_check(path: Path) -> str:
    try:
        Document(str(path))
    except Exception as exc:  # pragma: no cover - reported to user
        return f"FAILED: {exc}"
    return "OK"


def visible_text_font_check(path: Path) -> str:
    doc = Document(str(path))
    issues: list[str] = []
    blocks = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(cell.paragraphs)
    for paragraph in blocks:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            name = run.font.name or paragraph.style.font.name or doc.styles["Normal"].font.name
            size = run.font.size or paragraph.style.font.size or doc.styles["Normal"].font.size
            if name != "Arial":
                issues.append(f"font={name!r}")
            if size is not None and round(size.pt, 2) != 12 and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                # Centered title/header text uses 13-14 pt by design.
                issues.append(f"size={size.pt:g}")
    return "OK" if not issues else "WARN: " + ", ".join(sorted(set(issues))[:6])


def render_output_dir(primary: Path) -> Path:
    relative = primary.relative_to(PLANS_ROOT)
    safe = "__".join(part.replace(" ", "_").replace("/", "_") for part in relative.parts)
    safe = safe[:-5] if safe.lower().endswith(".docx") else safe
    return GENERATED_ROOT / "QA" / "rendered" / safe


def render_docx(path: Path, render_script: Path) -> str:
    output_dir = render_output_dir(path)
    output_dir.mkdir(parents=True, exist_ok=True)
    for old_page in output_dir.glob("page-*.png"):
        old_page.unlink()
    cmd = [sys.executable, str(render_script), str(path), "--output_dir", str(output_dir)]
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False)
    if proc.returncode != 0:
        return "FAILED: " + (proc.stderr.strip() or proc.stdout.strip())
    pages = sorted(output_dir.glob("page-*.png"))
    return f"OK ({len(pages)} page{'s' if len(pages) != 1 else ''})"


def write_report(records: list[OutputRecord], render_results: dict[Path, str] | None = None) -> Path:
    render_results = render_results or {}
    GENERATED_ROOT.mkdir(parents=True, exist_ok=True)
    report = GENERATED_ROOT / "generation-report.md"
    primary_count = len(records)
    mirror_count = sum(len(record.mirrors) for record in records)
    lines = [
        "# Generated Secondary Assessment Documents Report",
        "",
        f"Primary DOCX files: {primary_count}",
        f"Mirror DOCX copies: {mirror_count}",
        "Optional quiz-style checks: excluded",
        "",
        "## Outputs",
        "",
        "| Kind | Title | Primary File | Mirrors | Open Check | Font Check | Render QA | Mirror Hash Check |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for record in records:
        open_check = open_docx_check(record.primary) if record.primary.exists() else "PENDING"
        font_check = visible_text_font_check(record.primary) if record.primary.exists() else "PENDING"
        render_check = render_results.get(record.primary, "Not run")
        primary_hash = sha256(record.primary) if record.primary.exists() else ""
        mirror_hashes = []
        for mirror in record.mirrors:
            if not mirror.exists():
                mirror_hashes.append(f"{mirror.name}: missing")
            elif primary_hash and sha256(mirror) == primary_hash:
                mirror_hashes.append(f"{mirror.name}: OK")
            else:
                mirror_hashes.append(f"{mirror.name}: mismatch")
        mirror_text = "<br>".join(str(m.relative_to(ROOT)) for m in record.mirrors)
        lines.append(
            "| {kind} | {title} | `{primary}` | {mirrors} | {open_check} | {font_check} | {render_check} | {mirror_hash} |".format(
                kind=record.kind,
                title=record.title,
                primary=record.primary.relative_to(ROOT),
                mirrors=mirror_text,
                open_check=open_check,
                font_check=font_check,
                render_check=render_check.replace("|", "/"),
                mirror_hash="<br>".join(mirror_hashes),
            )
        )

    lines.extend(
        [
            "",
            "## Sources",
            "",
            "- 7th grade monthly planning drafts under `plans/7th Grade Technology/Planning/Drafts`.",
            "- 8th grade final assessment map: `plans/8th Grade Technology/Planning/Reviews/final-assessment-map-2026.md`.",
            "- 9th grade final assessment map: `plans/9th Grade Technology/Planning/Reviews/Final Assessment Map - 9th Grade Technology 2026.md`.",
            "",
            "## Assumptions",
            "",
            "- 6th grade counted summatives were already covered and were not changed.",
            "- Optional 8th grade quiz-style checks were not generated.",
            "- Appreciation grades normally use a compact 20-point checklist format; 8th grade IIT and all 9th grade appreciation summatives use the official 40-point rubric template.",
            "- Exam projects use a 90-point project packet and rubric format.",
        ]
    )
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def dry_run(records: list[OutputRecord]) -> None:
    print(f"Primary DOCX planned: {len(records)}")
    print(f"Mirror DOCX planned: {sum(len(record.mirrors) for record in records)}")
    by_kind: dict[str, int] = {}
    for record in records:
        by_kind[record.kind] = by_kind.get(record.kind, 0) + 1
    for kind, count in sorted(by_kind.items()):
        print(f"{kind}: {count}")
    for record in records:
        print(f"- {record.kind}: {record.primary}")


def main() -> int:
    args = parse_args()
    records = planned_outputs()
    if args.dry_run:
        dry_run(records)
        return 0

    generate(records, force=args.force)

    render_results: dict[Path, str] = {}
    if args.render:
        render_script = Path(args.render_docx).expanduser().resolve()
        if not render_script.exists():
            raise SystemExit(f"Render script not found: {render_script}")
        for record in records:
            render_results[record.primary] = render_docx(record.primary, render_script)

    report = write_report(records, render_results)
    print(f"Primary DOCX files: {len(records)}")
    print(f"Mirror DOCX copies: {sum(len(record.mirrors) for record in records)}")
    print(f"Report: {report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
