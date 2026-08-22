from pathlib import Path
import re

from docx import Document

from update_grade6_updated_docx_vocab_in_place import DEFINITIONS, DOCX_DIR, PRACTICE


def definition_lines(words):
    return "\n".join(f"{word}: {DEFINITIONS[word]}" for word in words)


def split_words(words):
    midpoint = (len(words) + 1) // 2
    return words[:midpoint], words[midpoint:]


def normalize_label_sentence(sentence, part_number):
    sentence = sentence.strip()
    sentence = re.sub(r"\s*\(part [12]\)", "", sentence)
    sentence = sentence.rstrip(".")
    return f"{sentence} (part {part_number})."


def update_text_for_label(text, label, part_words, part_number):
    all_words = PRACTICE[label]
    full_defs = definition_lines(all_words)
    if full_defs not in text:
        return text, False

    label_match = re.search(rf"([^\n]*{re.escape(label)} practice vocabulary[^\n]*\.)", text, flags=re.I)
    if not label_match:
        return text, False

    updated_sentence = normalize_label_sentence(label_match.group(1), part_number)
    text = text[:label_match.start(1)] + updated_sentence + text[label_match.end(1):]
    text = text.replace(full_defs, definition_lines(part_words), 1)
    return text, True


def update_paragraph(paragraph, label, part_words, part_number):
    text, changed = update_text_for_label(paragraph.text, label, part_words, part_number)
    if not changed:
        return 0

    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return 1


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def main():
    total = 0
    for path in sorted(DOCX_DIR.glob("**/6° Technology - *.docx")):
        if path.name.startswith("~$"):
            continue

        doc = Document(path)
        paragraphs = list(iter_paragraphs(doc))
        doc_count = 0

        for label, words in PRACTICE.items():
            matching = [
                paragraph for paragraph in paragraphs
                if f"{label} practice vocabulary" in paragraph.text
                and definition_lines(words) in paragraph.text
            ]
            if len(matching) < 2:
                continue

            first_half, second_half = split_words(words)
            doc_count += update_paragraph(matching[0], label, first_half, 1)
            doc_count += update_paragraph(matching[1], label, second_half, 2)

        if doc_count:
            doc.save(path)
        total += doc_count
        print(f"{path.name}: {doc_count} paragraphs updated")

    print(f"total paragraphs updated: {total}")


if __name__ == "__main__":
    main()
