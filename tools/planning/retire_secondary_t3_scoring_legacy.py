#!/usr/bin/env python3
"""Move superseded T3 rubrics aside, correct check filenames, and remove phone evidence."""

from pathlib import Path
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]

CURRENT_MARKERS = (" - Spreadsheet", " - Scratch", " - Source", " - Mandrake", " - App ",
                   " - Python", " - Digital", " - Audiovisual", " - Image", " - Cybersecurity",
                   " - STEM", "STEAM Preparation", "STEAM Expo")

RENAMES = {
    "7th grade - IIIT - Week 5 - Summative 4 - Source Credibility and Image Credit Check.docx":
        "7th grade - IIIT - Week 7 - Summative 4 - Source Credibility and Image Credit Check.docx",
    "7th grade - IIIT - Week 5 - Summative 4 - Source Credibility and Image Credit Check - Teacher Key.docx":
        "7th grade - IIIT - Week 7 - Summative 4 - Source Credibility and Image Credit Check - Teacher Key.docx",
    "8th grade - IIIT - Week 7 - Summative 4 - Python Selection and Loop Debugging Check.docx":
        "8th grade - IIIT - Week 8 - Summative 4 - Python Selection and Loop Debugging Check.docx",
    "8th grade - IIIT - Week 7 - Summative 4 - Python Selection and Loop Debugging Check - Teacher Key.docx":
        "8th grade - IIIT - Week 8 - Summative 4 - Python Selection and Loop Debugging Check - Teacher Key.docx",
    "8th grade - IIIT - Week 8 - Summative 5 - Representation and Binary Check.docx":
        "8th grade - IIIT - Week 9 - Summative 5 - Representation and Binary Check.docx",
    "8th grade - IIIT - Week 8 - Summative 5 - Representation and Binary Check - Teacher Key.docx":
        "8th grade - IIIT - Week 9 - Summative 5 - Representation and Binary Check - Teacher Key.docx",
    "9th grade - IIIT - Week 5 - Summative 3 - Cyber Security Scenario Quiz.docx":
        "9th grade - IIIT - Week 6 - Summative 3 - Cyber Security Scenario Quiz.docx",
    "9th grade - IIIT - Week 5 - Summative 3 - Cyber Security Scenario Quiz - Teacher Key.docx":
        "9th grade - IIIT - Week 6 - Summative 3 - Cyber Security Scenario Quiz - Teacher Key.docx",
}

REPLACEMENTS = {
    "screenshots/photos": "saved project files and typed test records",
    "screenshots and photos": "saved project files and typed test records",
    "screenshots or photos": "saved project files or typed test records",
    "screenshots/photos,": "saved project files and typed test records,",
    "screenshot/photo": "saved-file evidence",
    "screenshots": "saved-file evidence",
    "photos": "teacher-observed evidence",
    "take pictures": "save a typed test record",
}

def replace_in_doc(path):
    doc = Document(path)
    changed = False
    containers = [doc.paragraphs]
    for table in doc.tables:
        containers.extend(cell.paragraphs for row in table.rows for cell in row.cells)
    for paragraphs in containers:
        for paragraph in paragraphs:
            old = paragraph.text
            new = old
            for source, target in REPLACEMENTS.items():
                new = new.replace(source, target).replace(source.title(), target.capitalize())
            if new != old:
                paragraph.text = new
                changed = True
    if changed:
        print(f"updated {path}")
    if doc.tables:
        rubric = doc.tables[-1]
        header_props = rubric.rows[0]._tr.get_or_add_trPr()
        if header_props.find(qn("w:tblHeader")) is None:
            header_props.append(OxmlElement("w:tblHeader"))
        for row in rubric.rows:
            props = row._tr.get_or_add_trPr()
            if props.find(qn("w:cantSplit")) is None:
                props.append(OxmlElement("w:cantSplit"))
        changed = True
    if changed:
        doc.save(path)

for grade in (7, 8, 9):
    base = ROOT / f"plans/{grade}th Grade Technology/Assessments/Rubrics/IIIT"
    for sub in ("Daily", "Appreciation"):
        legacy = base / sub / "Legacy - do not use"
        legacy.mkdir(parents=True, exist_ok=True)
        for path in list((base / sub).glob("*.docx")):
            if "Rubric for" in path.name and not any(marker in path.name for marker in CURRENT_MARKERS):
                target = legacy / path.name
                if target.exists():
                    target.unlink()
                shutil.move(path, target)
                print(f"retired {path}")
        if sub == "Appreciation":
            for path in list((base / sub).glob("*.docx")):
                if "STEAM Preparation" not in path.name and "STEAM Expo" not in path.name:
                    target = legacy / path.name
                    if target.exists():
                        target.unlink()
                    shutil.move(path, target)
                    print(f"retired {path}")
    for old, new in RENAMES.items():
        source = base / "Daily" / old
        if source.exists():
            target = source.with_name(new)
            if target.exists():
                target.unlink()
            source.rename(target)
            print(f"renamed {source.name} -> {target.name}")
    for exam in (base / "Exam Projects").glob("*.docx"):
        replace_in_doc(exam)
