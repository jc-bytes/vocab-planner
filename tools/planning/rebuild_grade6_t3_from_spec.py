#!/usr/bin/env python3
"""Refresh the five active Grade 6 T3 DOCX rubrics from the scoring spec."""

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SPEC_PATH = ROOT / "plans/6th Grade Technology/Materials/Lesson Packages/T3 2026/Teacher Materials/Grade 6 T3 Scoring Specification.json"
CANON = ROOT / "plans/6th Grade Technology/Assessments/Rubrics/IIIT"
MIRROR = ROOT / "plans/Shared/Generated Outputs/Rubrics 2026/6th Grade Technology/3rd Trimester"


def style_run(run):
    run.font.name = "Arial"
    run.font.size = Pt(12)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), "Arial")


def set_cell(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            style_run(run)


def compact_cell(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), "30")
        node.set(qn("w:type"), "dxa")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.95


def main():
    data = json.loads(SPEC_PATH.read_text(encoding="utf-8"))
    MIRROR.mkdir(parents=True, exist_ok=True)
    for rubric in data["rubrics"]:
        path = CANON / rubric["docx"]
        doc = Document(path)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11), Inches(8.5)
        section.top_margin = section.bottom_margin = Inches(0.4)
        section.left_margin = section.right_margin = Inches(0.4)
        table = doc.tables[0]
        if len(table.rows) != 6:
            raise ValueError(f"Expected six rubric rows in {path}")
        for row, criterion in zip(table.rows[1:], rubric["criteria"]):
            set_cell(row.cells[0], f"{criterion['name']}\n{criterion['description']}")
            for cell, (points, description) in zip(row.cells[1:5], criterion["levels"]):
                set_cell(cell, f"{points} points\n{description}")
        for row in table.rows:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                compact_cell(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        style_run(run)
        for paragraph in list(doc.paragraphs):
            if paragraph.text.startswith(("Teacher scoring note:", "Comments:")):
                paragraph._element.getparent().remove(paragraph._element)
                continue
            for run in paragraph.runs:
                style_run(run)
        doc.save(path)
        shutil.copyfile(path, MIRROR / path.name)
        print(path)


if __name__ == "__main__":
    main()
