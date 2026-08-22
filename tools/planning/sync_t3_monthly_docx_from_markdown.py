#!/usr/bin/env python3
"""Synchronize all Grade 6-9 T3 official monthly DOCX files from Markdown."""

from __future__ import annotations

import re
import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
FONT = "Arial"
SIZE = Pt(12)


def section(text: str, heading: str, next_heading: str) -> str:
    match = re.search(
        rf"^## {re.escape(heading)}\n\n(.*?)\n\n## {re.escape(next_heading)}$",
        text,
        flags=re.M | re.S,
    )
    if not match:
        raise ValueError(f"Could not find section {heading}")
    return match.group(1).strip()


def list_items(block: str) -> list[str]:
    return [line[2:].strip().removeprefix("• ").strip() for line in block.splitlines() if line.startswith("- ")]


def monthly_rows(text: str) -> list[list[str]]:
    block = text.split("## Monthly Plan", 1)[1]
    rows: list[list[str]] = []
    for line in block.splitlines():
        if not re.match(r"^\| Week \d+ \| (45|90) minutes \|", line):
            continue
        cells = [cell.strip().replace("<br>", "\n") for cell in line.strip("|").split("|")]
        if len(cells) != 7:
            raise ValueError(f"Unexpected monthly row: {line}")
        if cells[2] == "No regular class scheduled in this monthly plan.":
            continue
        rows.append(cells)
    return rows


def set_font(run, *, bold: bool = False, highlight: bool = False) -> None:
    run.font.name = FONT
    run.font.size = SIZE
    run.bold = bold
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for name in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{name}"), FONT)


def clear_cell(cell) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def paragraph(cell):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_label_value(cell, label: str, value: str, *, assessment_aware: bool = False) -> None:
    p = paragraph(cell)
    set_font(p.add_run(label), bold=True)
    p.add_run().add_break()
    if not assessment_aware:
        set_font(p.add_run(value))
        return
    parts = re.split(r"(?<=[.!?])\s+", value)
    for index, part in enumerate(parts):
        if index:
            set_font(p.add_run(" "))
        highlight = any(
            marker in part
            for marker in (
                "Summative Activity #",
                "Daily Grade #",
                "Appreciation Grade #",
                "Begin Exam Project:",
                "Continue Exam Project:",
                "Complete Exam Project Presentation:",
            )
        )
        set_font(p.add_run(part), highlight=highlight)


def set_summary(cell, label: str, values: list[str]) -> None:
    clear_cell(cell)
    p = paragraph(cell)
    set_font(p.add_run(f"{label}:"), bold=True)
    for value in values:
        p.add_run().add_break()
        set_font(p.add_run(value))


def set_metadata(cell, label: str, value: str, *, add_left_space: bool = False) -> None:
    clear_cell(cell)
    p = paragraph(cell)
    if add_left_space:
        set_font(p.add_run("\u00a0\u00a0"))
    set_font(p.add_run(f"{label}:"), bold=True)
    set_font(p.add_run(f" {value}"))


def set_lesson(cell, data: list[str]) -> None:
    _week, _duration, topic, objective, pre, during, post = data
    clear_cell(cell)
    add_label_value(cell, "Topic:", topic)
    add_label_value(cell, "Class Objective:", objective)
    add_label_value(cell, "Pre-activities:", pre)
    add_label_value(cell, "While activities:", during, assessment_aware=True)
    add_label_value(cell, "Post activities:", post)


def set_resources(cell, grade: int, topic: str) -> None:
    clear_cell(cell)
    if grade == 6:
        value = "Pen, notebook, assigned computer and login; phones prohibited"
    else:
        value = "Pen, pencil, technology notebook, student's computer, email/password, phones prohibited"
    if "STEAM" in topic:
        value += "; project brief, group log, test table, observation checklist, project materials"
    elif "micro:bit" in topic or "Mandrake" in topic or "Counter" in topic or "Sensor" in topic or "Night-light" in topic or "night light" in topic:
        value += "; MOD-MICROBIT-SENSING-01 v0.1.0, fallback, MakeCode, micro:bit/simulator, USB cable"
    elif grade == 6 and ("Tinkercad" in topic or "3D" in topic):
        value += "; MOD-3D-MODELLING-01 v0.1.0, fallback, Tinkercad"
    elif "Scratch" in topic:
        value += ", Scratch, starter file"
    elif "Python" in topic:
        value += ", Python environment, starter file"
    elif "Data" in topic or "data" in topic or "spreadsheet" in topic.lower() or "chart" in topic.lower():
        value += "; MOD-DATA-SPREADSHEETS-01 v0.5.0, fallback, spreadsheet starter"
    p = paragraph(cell)
    set_font(p.add_run("Resources:"), bold=True)
    p.add_run().add_break()
    set_font(p.add_run(value))


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def update(md: Path, docx: Path, grade: int) -> None:
    text = md.read_text(encoding="utf-8")
    rows = monthly_rows(text)
    doc = Document(docx)
    if len(doc.tables) < 2:
        raise ValueError(f"Expected two tables: {docx}")

    overview = doc.tables[0]
    metadata = (
        ("Subject", "Technology"),
        ("Teacher", "Porfirio Rios"),
        ("Grade", f"{grade}th"),
        ("Trimester", "3rd"),
        ("Year", "2026"),
        ("Weekly Hours", "3"),
    )
    for column, (label, value) in enumerate(metadata):
        set_metadata(overview.cell(0, column), label, value, add_left_space=column > 0)
    set_summary(overview.cell(1, 0), "Competences", list_items(section(text, "Competences", "Learning Objectives")))
    set_summary(overview.cell(2, 0), "Learning Objectives", list_items(section(text, "Learning Objectives", "Learning Outcomes")))
    set_summary(overview.cell(3, 0), "Learning Outcomes", list_items(section(text, "Learning Outcomes", "Evaluation")))

    lesson_table = doc.tables[1]
    expected_weeks = len(rows) // 2
    required_rows = expected_weeks * 3
    while len(lesson_table.rows) < required_rows:
        source_rows = list(lesson_table.rows[-3:])
        if len(source_rows) != 3:
            raise ValueError(f"Lesson table too short for {len(rows)} rows: {docx}")
        for source_row in source_rows:
            lesson_table._tbl.append(deepcopy(source_row._tr))
    for index, data in enumerate(rows):
        content_row = 1 + (index // 2) * 3
        resource_row = content_row + 1
        column = index % 2
        set_lesson(lesson_table.cell(content_row, column), data)
        if resource_row < len(lesson_table.rows):
            keep_row_together(lesson_table.rows[resource_row])
            set_resources(lesson_table.cell(resource_row, column), grade, data[2])

    while len(lesson_table.rows) > required_rows:
        lesson_table._tbl.remove(lesson_table.rows[-1]._tr)

    # Remove the template's trailing empty paragraph. In a tightly filled plan,
    # LibreOffice otherwise pushes it to a new page and repeats the final green
    # week heading there by itself.
    trailing = doc.paragraphs[-1]
    if not trailing.text:
        trailing._p.getparent().remove(trailing._p)

    doc.save(docx)
    print(f"updated {docx.relative_to(ROOT)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--grade", type=int, choices=(6, 7, 8, 9))
    args = parser.parse_args()
    grades = (args.grade,) if args.grade else (6, 7, 8, 9)
    for grade in grades:
        root = ROOT / "plans" / f"{grade}th Grade Technology" / "Planning"
        draft_root = root / "Drafts" / "3rd Trimester"
        word_root = root / "Monthly" / "3rd Trimester"
        for month in ("September", "October", "November", "December"):
            update(
                draft_root / f"{grade}° Technology - {month}.md",
                word_root / f"{grade}° Technology - {month}.docx",
                grade,
            )


if __name__ == "__main__":
    main()
