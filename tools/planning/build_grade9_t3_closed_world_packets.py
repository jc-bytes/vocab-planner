#!/usr/bin/env python3
"""Build the closed-world Grade 9 T3 assessment packets and student checks."""

from pathlib import Path
import shutil
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
RUBRIC_ROOT = ROOT / "plans/9th Grade Technology/Assessments/Rubrics/IIIT"
PACKAGE = ROOT / "plans/9th Grade Technology/Materials/Lesson Packages/T3 2026"
MIRROR = ROOT / "plans/Shared/Generated Outputs/Rubrics 2026/9th Grade Technology/3rd Trimester"
BLUE = "174E6A"
LIGHT = "DCECF3"
YELLOW = "FFF0B3"
FONT = "Arial"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Normal"].paragraph_format.line_spacing = 1.0
    for name, size, color in (("Title", 22, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 12, BLUE)):
        styles[name].font.name = FONT
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
    doc.core_properties.title = title
    doc.core_properties.subject = "Grade 9 Technology third-trimester assessment packet"


def add_header(doc, label, title, points):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ACADEMIA INTERNACIONAL DAVID\nROBOTICS AND TECHNOLOGY")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Grade 9 · Third Trimester · {label} · {points} points").bold = True
    info = doc.add_table(rows=2, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    info.cell(0, 0).text = "Student: __________________________________"
    info.cell(0, 1).text = "Date: __________________"
    info.cell(1, 0).text = "Group: 9° A / B"
    info.cell(1, 1).text = f"Score: ______ / {points}"
    for row in info.rows:
        for cell in row.cells:
            set_margins(cell)


def add_bullets(doc, heading, items, numbered=False):
    doc.add_heading(heading, level=1)
    for index, item in enumerate(items, 1):
        marker = f"{index}." if numbered else "•"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(f"{marker}  {item}")


def add_checklist(doc, items):
    doc.add_heading("Finish and submission checklist", level=1)
    for item in items:
        doc.add_paragraph(f"☐ {item}")


def add_rubric(doc, criteria, points=40):
    doc.add_page_break()
    doc.add_heading("Scoring rubric", level=1)
    headers = ["Criterion", "Complete evidence", "Minor gap", "Partial evidence", "Not demonstrated"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        shade(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    if points == 40:
        scored = [(criterion, [9, 7, 5, 2]) for criterion in criteria]
        responsibility = (
            "Punctuality, responsibility, and readiness",
            "Device and login ready; on time; responsible use; deadline met.",
            "One issue corrected after one reminder; work remains checkable.",
            "Device or login missing, late or unprepared, or repeated reminders; some fallback work.",
            "No device or login and no fallback, no checkable work, or responsible-use rules refused.",
        )
        scored.append((responsibility, [4, 3, 2, 0]))
    else:
        maxima = [14, 14, 14, 13, 13, 13]
        scored = [
            (criterion, [maximum, 10 if maximum == 14 else 9, 5, 0])
            for criterion, maximum in zip(criteria, maxima)
        ]
        responsibility = (
            "Punctuality, responsibility, and readiness",
            "Device, login, and materials are ready at every checkpoint; starts and milestones are on time; files are organized; final evidence is submitted.",
            "One readiness, punctuality, or milestone issue is corrected after one reminder; final evidence remains checkable.",
            "Readiness or milestone issues repeat, or one session lacks the device or login; some fallback project work is completed.",
            "Readiness is not shown across checkpoints and the approved fallback or final evidence is not completed.",
        )
        scored.append((responsibility, [9, 7, 4, 0]))
    for criterion, level_points in scored:
        cells = table.add_row().cells
        cells[0].text = criterion[0]
        for cell, value, score in zip(cells[1:], criterion[1:], level_points):
            cell.text = f"{score} points\n{value}"
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_margins(cell)
        shade(cells[0], LIGHT)
        cells[0].paragraphs[0].runs[0].bold = True
    p = doc.add_paragraph()
    p.add_run("Readiness boundary: ").bold = True
    p.add_run(("4 of 40 points" if points == 40 else "9 of 90 points") + " may reflect a forgotten device or login, lateness, or unpreparedness. A documented school network, platform, teacher-provided account, or hardware failure outside the student's control does not reduce these points when the assigned fallback is followed. Academic criteria score the work itself and do not repeat this deduction.")
    doc.add_paragraph("Teacher comments: ______________________________________________________________\n________________________________________________________________________________")


DAILY = [
    {
        "week": 1,
        "number": 1,
        "title": "Audiovisual Representation Vocabulary",
        "overview": "Create a ten-row vocabulary table and one short connection using only the provided representation source.",
        "materials": [
            "01 Representation Vocabulary Source.md",
            "MOD-DIGITAL-REPRESENTATION-01: IMAGE-LEARN-01 and SOUND-LEARN-01",
            "Teacher-provided blank table or approved paper copy",
        ],
        "directions": [
            "Write these ten terms as separate rows: pixel, resolution, color depth, RGB, sample, sample rate, sample size, microphone, speaker, and compression.",
            "For each term, write one accurate definition in your own clear words and one correct image, sound, device, or file example.",
            "Write 3–5 sentences about one practiced relationship: resolution or color depth and image file size, or sample rate or sample size and sound file size. Name the setting, describe the change, and explain the file-size result.",
            "Do not add internet research or a term that is not on the list.",
        ],
        "submit": "Submit one .docx, .pdf, .xlsx, teacher-provided Google file, or approved printed table named 9_-LastName-DG1-Vocabulary.",
        "check": ["All ten terms appear.", "All ten definitions and examples are complete.", "The connection is 3–5 sentences and contains the setting, change, and file-size result.", "The product and filename match the Submit line."],
        "criteria": [
            ("Ten required terms", "Includes all 10 named terms.", "Includes 9 terms.", "Includes 6–8 terms.", "Includes 0–5 terms."),
            ("Ten definitions", "All 10 are accurate, clear, and based on the source.", "Nine are accurate; one needs correction.", "Six to eight are accurate.", "Zero to five are accurate."),
            ("Ten examples", "All 10 examples correctly match their terms.", "Nine examples are correct.", "Six to eight examples are correct.", "Zero to five examples are correct."),
            ("Three-to-five-sentence connection", "Names the setting, change, and correct file-size result in 3–5 sentences.", "Includes all three ideas; one needs detail.", "Includes only one or two accurate ideas or misses the length.", "Missing or contradicts the taught relationship."),
        ],
    },
    {
        "week": 2,
        "number": 2,
        "title": "Image and Sound Representation Check",
        "overview": "Complete the named check using the same vocabulary, image operations, sound meanings, and trade-off practiced in the module.",
        "materials": ["02 Image and Sound Reference.md", "Image and Sound Representation Check.docx", "MOD-DIGITAL-REPRESENTATION-01 practice history"],
        "directions": [
            "Explain pixel, resolution, color depth, sample rate, and sample size in context.",
            "Show 10 × 8 = 80 pixels and 10 × 8 × 8 = 640 bits with units.",
            "Explain what happens to file size when color depth increases.",
            "Distinguish samples per second from bits per sample and explain one quality/file-size relationship.",
            "Recommend one taught way to reduce an audio file and name one benefit and one possible quality cost. No other formula is assessed.",
        ],
        "submit": "Return every numbered response in the provided DOCX or paper check. A device or login failure does not change the task.",
        "check": ["Five vocabulary meanings are complete.", "Both image operations and units are shown.", "Sound-rate and sound-size meanings are not reversed.", "The recommendation contains a change, benefit, and cost."],
        "criteria": [
            ("Vocabulary meanings", "All five meanings are accurate in context.", "Four are accurate.", "Two or three are accurate.", "Zero or one is accurate."),
            ("Image operations", "Both operations, results, and units are correct.", "Both results are correct; one operation or unit is incomplete.", "One result is correct or both show partial work.", "Both are missing or incorrect."),
            ("Sound and quality reasoning", "Correctly distinguishes rate from size and explains detail/file-size change.", "Meanings are correct; one relationship needs detail.", "One meaning or relationship is correct.", "Missing or incorrect."),
            ("Compression recommendation", "Names one taught change, one size benefit, and one quality cost.", "All three are present; one needs detail.", "Only one or two accurate parts are present.", "Missing or unsupported."),
        ],
    },
    {
        "week": 6,
        "number": 3,
        "title": "Cybersecurity Scenario Analysis",
        "overview": "Analyze six fixed written scenarios. Do not test an attack or use real credentials.",
        "materials": ["03 Cybersecurity Source Page.md", "Cyber Security Scenario Quiz.docx", "MOD-CYBERSECURITY-RISK-01 threat and protection practice"],
        "directions": [
            "For scenarios 1–6, name one source-page threat and one matching safe protection.",
            "Choose three different scenario numbers. For each, name the clue, explain why the threat fits, and explain how the protection reduces risk.",
            "Choose one source-page protection as a school rule and explain why it helps many users.",
            "Do not open links or files, contact suspicious senders, enter credentials, or blame the person who received the message.",
        ],
        "submit": "Return every numbered response in the provided DOCX or paper quiz.",
        "check": ["Six threats and six protections are present.", "Three different scenarios are justified from their clues.", "The school rule is safe, practical, and broadly useful.", "Every response is complete."],
        "criteria": [
            ("Six threat identifications", "All six fit the fixed scenarios.", "Five fit.", "Three or four fit.", "Zero to two fit."),
            ("Six protection choices", "All six are safe and match the threats.", "Five fit.", "Three or four fit.", "Zero to two fit or an unsafe action appears."),
            ("Three evidence justifications", "All three name the clue, explain the threat, and explain protection effect.", "All three exist; one explanation needs detail.", "Only one or two are complete.", "None is complete."),
            ("School protection rule", "Uses one taught protection and explains broad school benefit safely.", "Safe and broad; one reason needs detail.", "Partly useful or weakly explained.", "Missing, unsafe, or unrelated."),
        ],
    },
    {
        "week": 7,
        "number": 4,
        "title": "Cybersecurity Risk Map",
        "overview": "Analyze only fixed records R1–R6 and choose one highest-priority record.",
        "materials": ["03 Cybersecurity Source Page.md", "04 Fixed Six-Risk Record.md", "Teacher-provided risk-map template"],
        "directions": [
            "For R1–R6, identify the asset and threat or harmful event.",
            "Rate probability and impact separately as low or high. Cite one record fact for every rating.",
            "Choose one feasible source-page protection for every record and explain how it helps.",
            "Write one paragraph naming the highest-priority record and defend it using probability, impact, and at least one fact.",
        ],
        "submit": "Submit 9_-LastName-DG4-Risk-Map as .docx, .pdf, .xlsx, teacher-provided Google file, or the complete paper fallback.",
        "check": ["R1–R6 are all present.", "Every record has separate probability and impact ratings with facts.", "Every protection is feasible and explained.", "The priority paragraph agrees with the map."],
        "criteria": [
            ("Fixed records R1–R6", "All six have asset and threat/event analysis.", "Five are complete.", "Three or four are complete.", "Zero to two are complete."),
            ("Probability and impact", "All ratings are separate, consistent, and supported by record facts.", "All ratings exist; one or two facts need correction.", "Several ratings or facts are missing.", "Mostly missing or arbitrary."),
            ("Six protections", "Every record has a feasible taught protection and effect.", "Five fit.", "Three or four fit.", "Zero to two fit or an unsafe action appears."),
            ("Priority paragraph", "Uses record ID, probability, impact, and a fact to defend priority.", "All supports appear; one needs detail.", "Only one or two supports appear.", "Missing or contradicts the map."),
        ],
    },
    {
        "week": 9,
        "number": 5,
        "title": "STEM Project Proposal",
        "overview": "Complete the provided nine-section template for the option card assigned by the teacher.",
        "materials": ["05 STEM Proposal Option Cards.md", "06 STEM Proposal Template.md", "Named module sections on the assigned card"],
        "directions": [
            "Write the assigned option, fixed problem, users, need explanation, and a testable goal.",
            "Name an allowed product and approved tool, list only needed resources, give the working-file location, and add a labeled input-process/output design.",
            "Complete C1–C3 using the three required option-card inputs. Give an expected result and visible or numeric pass rule for each.",
            "Name the v1 file, final file, dated logbook, three typed test rows, before/after comparison, live demonstration, and one realistic limitation.",
            "Do not choose a different topic or search the internet unless the teacher writes an approved change.",
        ],
        "submit": "Submit 9_-LastName-DG5-Proposal as .docx, .pdf, the teacher-provided Google file, or the complete printed template.",
        "check": ["All nine template sections are complete.", "The assigned option and fixed inputs are unchanged unless approval is written.", "C1–C3 have measurable pass rules.", "Every named evidence item and one limitation appear."],
        "criteria": [
            ("Assigned problem, users, and goal", "Option, fixed problem/users, need, and measurable goal are complete.", "All are present; one needs detail.", "Two are missing or broad.", "No usable assigned problem and goal."),
            ("Product, tool, and labeled design", "Allowed product/tool, resources, location, and input-process/output design are complete.", "All are present; one needs correction.", "Several parts are incomplete.", "No feasible product and design."),
            ("Criteria C1–C3", "All fixed inputs, expected results, and measurable pass rules are complete.", "All three exist; one result/rule needs detail.", "Only one or two are complete.", "No usable criteria."),
            ("Evidence plan and limitation", "All seven named evidence items and one realistic limitation appear.", "All appear; one needs detail.", "Several are missing.", "No checkable evidence plan."),
        ],
    },
]


APPRECIATION = [
    {
        "week": 4, "number": 1, "title": "STEAM Preparation and Work Process",
        "overview": "Earn an individual Technology process grade while continuing the one project assigned by Arts or Science.",
        "materials": ["Shared Assigned Project Brief and Group Plan", "Shared Student Individual Work Log", "Shared Group Progress Log", "Shared Test and Improvement Record"],
        "directions": ["Use the assigned Grade 9 project: 9A Arts or 9B Science. Do not start a second Technology project.", "In both STEAM weeks, state your assigned role, locate the current group product, complete the named task, and leave checkable evidence.", "Keep at least three dated individual entries. Each entry names the task, result, and next step.", "Use files, materials, time, and communication safely and respectfully. The teacher records individual evidence; the group product alone does not determine the score."],
        "submit": "Submit the individual work log or approved paper fallback. The teacher also uses live observation and the shared product/log.",
        "check": ["The correct 9A Arts or 9B Science project is named.", "The assigned role and contribution can be checked.", "At least three dated task-result-next-step entries exist.", "Files or materials are handed off safely."],
        "criteria": [
            ("Preparation", "In both weeks, uses the brief, states project/role, locates the product, and begins the named task.", "All four actions occur with one reminder.", "Only two or three occur or repeated reminders are needed.", "Zero or one occurs; no preparation evidence."),
            ("Participation and role", "Completes the assigned role and useful work at every checkpoint.", "One checkpoint or detail is incomplete.", "Only part of the role is completed.", "No usable individual contribution evidence."),
            ("Responsible work", "Uses files, materials, time, and communication safely and completes the handoff.", "One corrected reminder is needed.", "Repeated reminders are needed.", "Unsafe use or blocked group progress."),
            ("Individual work log", "At least three dated entries each name task, result, and next step.", "Two are complete and one is partial.", "One is complete or entries are vague.", "No checkable individual log."),
        ],
    },
    {
        "week": 5, "number": 2, "title": "STEAM Expo Participation and Closure",
        "overview": "Complete an individual Technology expo role and reflection for the same assigned Arts or Science project.",
        "materials": ["Shared Expo Explanation and Reflection", "Shared Teacher Individual Observation Sheet", "Current group product and progress log", "Shared STEAM Make-Up Path"],
        "directions": ["Complete the assigned introduction, explanation, demonstration, question, or materials role.", "Explain the problem, product, process, one test result, and one improvement.", "Listen respectfully, follow expo directions, and complete the assigned cleanup or materials-return task.", "Complete the individual reflection with one result, one challenge, and one next improvement."],
        "submit": "Submit the individual reflection; complete the role live or through the shared make-up path.",
        "check": ["The assigned role is complete.", "All five explanation points appear.", "Cleanup or materials return is complete.", "The reflection names a result, challenge, and improvement."],
        "criteria": [
            ("Assigned expo role", "Completes the assigned role without missing parts.", "Completes it with one prompt or small missing part.", "Completes only part or needs several prompts.", "Does not complete the role or approved make-up."),
            ("Five-point explanation", "Clearly explains problem, product, process, test result, and improvement.", "Explains four points.", "Explains two or three points.", "Explains zero or one point."),
            ("Respect and cleanup", "Listens, follows directions, and completes the assigned cleanup/return task.", "One corrected reminder.", "Repeated reminders or one part incomplete.", "Participation or cleanup evidence is absent."),
            ("Individual reflection", "Names one specific result, challenge, and improvement.", "All three appear; one needs detail.", "Only one or two are clear.", "Missing or unrelated."),
        ],
    },
]


EXAM_CRITERIA = [
    ("Approved problem and design", "Assigned option, problem, users, goal, approved tool, resources, labeled design, and C1–C3 are complete and consistent.", "Most elements are complete; one or two need correction.", "Several elements are missing or broad.", "No usable approved plan."),
    ("Prototype or product", "The checkable product performs the card's required action for all three fixed inputs.", "The main action works for two inputs or has one minor issue.", "Only one input works or the product is partial.", "No checkable product or paper equivalent."),
    ("Controlled testing", "Records at least three trials using C1–C3, expected and actual results, pass status, and evidence location.", "Three trials exist; one or two fields need detail.", "Fewer than three complete trials or inconsistent conditions.", "No meaningful test record."),
    ("Evidence-based improvement", "Names a weakness proved by results, makes one focused change, repeats an equivalent test, and compares before/after results.", "All steps appear; one link needs detail.", "A change exists but evidence or comparison is weak.", "No evidence-based improvement."),
    ("Logbook and files", "v1, final, records file, dated log entries, decisions, feedback, and next actions are complete and accessible.", "Most items are complete; one or two need correction.", "Documentation has several gaps.", "Documentation is mostly missing."),
    ("Demo and reflection", "In 1–2 minutes shows the product and explains problem, users, tool, choice, three results, improvement, limitation, and learning; reflection answers all seven prompts.", "Most explanation and reflection parts are clear.", "Several required parts are missing or vague.", "Demo/reflection evidence is missing."),
]


def build_packet(spec, label, out, points=40):
    doc = Document()
    configure(doc, spec["title"])
    add_header(doc, label, spec["title"], points)
    doc.add_heading("What you will make", level=1)
    doc.add_paragraph(spec["overview"])
    add_bullets(doc, "Provided materials", spec["materials"])
    add_bullets(doc, "Exact directions", spec["directions"], numbered=True)
    doc.add_heading("What to submit", level=1)
    doc.add_paragraph(spec["submit"])
    add_checklist(doc, spec["check"])
    doc.add_heading(f"Responsibility points: {4 if points == 40 else 9} of {points}", level=1)
    if points == 40:
        doc.add_paragraph("Bring your assigned charged computer, have your school login ready, begin on time, use files and materials responsibly, and submit or return the required work by the announced deadline. If you forget the computer or login, arrive late or unprepared, or need repeated responsibility reminders, points may be deducted only from these 4 points. Academic criteria are scored from the work itself.")
    else:
        doc.add_paragraph("At every announced project checkpoint, bring your assigned charged computer, required materials, and school login; begin on time; keep files and materials organized; meet the announced milestones; and submit the final evidence by the deadline. If you forget the computer or login, arrive late or unprepared, or miss responsibility checkpoints, points may be deducted only from these 9 points. Academic project criteria are scored from the project evidence itself.")
    doc.add_heading("Make-up and technology-failure path", level=1)
    doc.add_paragraph("Complete the same content with the approved paper copy or scheduled teacher check. A documented school network, platform, teacher-provided account, or hardware failure outside the student's control does not reduce readiness points when the assigned fallback is followed. A forgotten personal device, uncharged device, or unprepared login may reduce only the readiness criterion.")
    add_rubric(doc, spec["criteria"], points)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT
            if run.font.size is None or run.font.size.pt < 12:
                run.font.size = Pt(12)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT
                        run.font.size = Pt(12)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    if RUBRIC_ROOT in out.parents:
        mirror = MIRROR / out.name
        mirror.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(out, mirror)


def build_exam():
    spec = {
        "title": "STEM Project Final Exam",
        "overview": "Build the product approved in Daily Grade 5, test it with the three fixed option-card inputs, improve it from evidence, and demonstrate the result.",
        "materials": ["Approved Daily Grade 5 proposal", "05 STEM Proposal Option Cards.md", "10 STEM Project Brief.md", "11 Project Logbook.md", "12 Milestone Checkpoints.md", "13 Controlled Testing Record.md", "14 Demonstration Notes.md", "15 Final Reflection.md", "16 Make-Up and Offline Path.md"],
        "directions": ["Use the assigned option, problem, users, product choices, knowledge source, and three test inputs. Do not search for a new topic.", "Save v1, final, and records files using the names in the proposal.", "Build a checkable product or approved paper interaction model.", "Record at least three controlled trials using C1–C3, identify one weakness, make one focused change, and repeat an equivalent test.", "Keep dated log entries and complete the milestone checks.", "Give a 1–2 minute demonstration and complete all seven reflection prompts."],
        "submit": "Submit the final product or approved paper equivalent, v1/final/records evidence, completed logbook, testing record, demonstration notes, and reflection.",
        "check": ["The approved option and proposal are present.", "The product is checkable with all three fixed inputs.", "Three trials and one before/after comparison are complete.", "The logbook, demo notes, and seven-part reflection are complete."],
        "criteria": EXAM_CRITERIA,
    }
    out = RUBRIC_ROOT / "Exam Projects/9th grade - IIIT - Final Project Packet and Rubric - STEM Project.docx"
    build_packet(spec, "Final Exam · Weeks 10–12", out, 90)


def build_checks():
    daily = RUBRIC_ROOT / "Daily"
    # Image and sound check and key.
    for key in (False, True):
        doc = Document()
        configure(doc, "Image and Sound Representation Check")
        add_header(doc, "Daily Grade 2 · Week 2", "Image and Sound Representation Check" + (" — Teacher Key" if key else ""), 40)
        doc.add_paragraph("Use only the provided Image and Sound Reference if the teacher says it may remain open. Show both image operations and units. No other formula is assessed.")
        doc.add_heading("Part A · Vocabulary meanings · 10 points", level=1)
        table = doc.add_table(rows=6, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Term", "Meaning in a real file"
        answers = {"pixel":"one small color square in an image","resolution":"pixels across and down","color depth":"bits used for one pixel's color","sample rate":"samples recorded each second","sample size":"bits used for one sound sample"}
        for i, term in enumerate(answers, 1):
            table.rows[i].cells[0].text = term
            table.rows[i].cells[1].text = answers[term] if key else ""
        prompts = [
            ("Part B · Image representation · 10 points", ["1. A 10 × 8 image contains how many pixels? Show the operation and unit.", "2. If each pixel uses 8 bits, how many bits are needed? Show the operation and unit.", "3. What happens to file size when color depth increases? Explain why."], ["10 × 8 = 80 pixels", "10 × 8 × 8 = 640 bits", "It increases because more bits are stored for every pixel."]),
            ("Part C · Sound representation · 12 points", ["4. What does sample rate measure? Include the per-second idea.", "5. What does sample size measure? Include the per-sample idea.", "6. Explain why a higher sample rate or sample size can increase detail and file size."], ["Samples recorded per second, measured in Hz.", "Bits used to store one sample.", "More samples or more bits store more sound detail and more data."]),
            ("Part D · Compression trade-off · 8 points", ["7. An audio file is too large to upload. Name one taught change, one file-size benefit, and one possible quality cost."], ["Example: compress the file; it becomes smaller, but some sound detail may be lost."]),
        ]
        for heading, qs, ans in prompts:
            doc.add_heading(heading, level=1)
            for q, a in zip(qs, ans):
                doc.add_paragraph(q)
                doc.add_paragraph(("Answer: " + a) if key else "________________________________________________________________________________\n________________________________________________________________________________")
        name = "9th grade - IIIT - Week 2 - Summative 2 - Image and Sound Representation Check" + (" - Teacher Key" if key else "") + ".docx"
        doc.save(daily / name)

    # Cybersecurity quiz and key.
    scenarios = [
        ("An email says your account will close today unless you click a link and enter your password.", "phishing", "Do not click; verify through the official site or trusted adult; report it."),
        ("A game download secretly installs software that encrypts files and demands payment.", "ransomware or trojan malware", "Avoid untrusted downloads; use verified updates and backups."),
        ("A website is overwhelmed by thousands of fake requests and real users cannot access it.", "DDoS", "Use firewall/service filtering and report to the administrator."),
        ("Someone tries many password combinations until one works.", "brute-force attack", "Use a strong unique password, MFA, and lockout/rate limits."),
        ("A message pretends to be from the principal and asks students to send private information.", "social engineering or phishing", "Verify identity and do not share unnecessary personal data."),
        ("A computer keeps spreading copies of malicious code through connected devices.", "worm malware", "Isolate the device, report it, scan it, and apply verified updates."),
    ]
    for key in (False, True):
        doc = Document(); configure(doc, "Cybersecurity Scenario Quiz")
        add_header(doc, "Daily Grade 3 · Week 6", "Cybersecurity Scenario Quiz" + (" — Teacher Key" if key else ""), 40)
        doc.add_paragraph("Use only the provided Cybersecurity Source Page if the teacher says it may remain open. Analyze the written records. Do not test any attack or use real credentials.")
        doc.add_heading("Part A · Six threats and protections · 24 points", level=1)
        table = doc.add_table(rows=7, cols=4); table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ["#", "Fixed scenario", "Threat", "Safe protection"]): cell.text = text
        for i, (scenario, threat, protection) in enumerate(scenarios, 1):
            vals = [str(i), scenario, threat if key else "", protection if key else ""]
            for cell, text in zip(table.rows[i].cells, vals): cell.text = text
        doc.add_heading("Part B · Three evidence justifications · 12 points", level=1)
        for n in range(7, 10):
            doc.add_paragraph(f"{n}. Choose a different scenario number. Name its clue, explain why the threat fits, and explain how the protection reduces risk.")
            doc.add_paragraph("Accept any three complete, accurate clue-threat-protection explanations." if key else "Scenario #: ____\n________________________________________________________________________________\n________________________________________________________________________________")
        doc.add_heading("Part C · School protection rule · 4 points", level=1)
        doc.add_paragraph("10. Choose one source-page protection as a school rule and explain why it helps many users.")
        doc.add_paragraph("Accept one safe taught protection with a broad, accurate reason." if key else "________________________________________________________________________________\n________________________________________________________________________________")
        name = "9th grade - IIIT - Week 6 - Summative 3 - Cyber Security Scenario Quiz" + (" - Teacher Key" if key else "") + ".docx"
        doc.save(daily / name)


def main():
    for spec in DAILY:
        name = f"9th grade - IIIT - Week {spec['week']} - Rubric for Summative {spec['number']} - {spec['title']}.docx"
        build_packet(spec, f"Daily Grade {spec['number']} · Week {spec['week']}", RUBRIC_ROOT / "Daily" / name)
    for spec in APPRECIATION:
        name = f"9th grade - IIIT - Week {spec['week']} - Rubric for Appreciation Grade {spec['number']} - {spec['title']}.docx"
        build_packet(spec, f"Appreciation Grade {spec['number']} · Week {spec['week']}", RUBRIC_ROOT / "Appreciation" / name)
    build_exam()
    build_checks()


if __name__ == "__main__":
    main()
