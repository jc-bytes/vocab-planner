from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")


def set_merged_row_text(table, row_idx, text):
    seen = set()
    for cell in table.rows[row_idx].cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        set_cell_text(cell, text)


def delete_rows_from(table, start_idx):
    while len(table.rows) > start_idx:
        row = table.rows[start_idx]
        table._tbl.remove(row._tr)


def update_docx(rel_path, *, weeks=None, meta=None, lessons=None, resources=None, delete_lesson_rows_from=None):
    path = ROOT / rel_path
    doc = Document(path)

    if weeks is not None:
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("WEEKS:"):
                set_paragraph_text(paragraph, f"WEEKS: {weeks}")
                break

    if meta:
        info = doc.tables[0]
        for row_idx, text in meta.items():
            set_merged_row_text(info, row_idx, text)

    if lessons:
        lesson_table = doc.tables[1]
        for (row_idx, col_idx), text in lessons.items():
            set_cell_text(lesson_table.rows[row_idx].cells[col_idx], text)

    if resources:
        lesson_table = doc.tables[1]
        for row_idx, text in resources.items():
            for cell in lesson_table.rows[row_idx].cells:
                set_cell_text(cell, text)

    if delete_lesson_rows_from is not None:
        delete_rows_from(doc.tables[1], delete_lesson_rows_from)

    doc.save(path)
    print(f"updated {rel_path}")


GRADE3_GENERIC_META = {
    1: (
        "Competences:\n"
        "Follow safe Grade 3 technology routines while working with monthly tools, vocabulary, STEM project evidence, and class materials.\n"
        "Identify, label, sort, create, test, and submit visible evidence using guided class steps."
    ),
    2: (
        "Learning Objectives:\n"
        "Identify key vocabulary and examples from the monthly topics.\n"
        "Complete the guided class evidence named in each lesson with teacher support.\n"
        "Check work for clear labels, correct steps, safe handling, and a simple reflection."
    ),
    3: (
        "Learning Outcomes:\n"
        "Complete the formal summative activity or activities scheduled for this month.\n"
        "Complete formative notebook, worksheet, drawing, digital, STEM project, or reflection evidence."
    ),
}


GRADE4_GENERIC_META = {
    1: (
        "Competences:\n"
        "Follow safe Grade 4 technology routines while working with monthly tools, vocabulary, STEM project evidence, and class materials.\n"
        "Identify, label, sort, create, test, and submit visible evidence using guided class steps."
    ),
    2: (
        "Learning Objectives:\n"
        "Identify key vocabulary and examples from the monthly topics.\n"
        "Complete the guided class evidence named in each lesson with teacher support.\n"
        "Check work for clear labels, correct steps, safe handling, and a simple reflection."
    ),
    3: (
        "Learning Outcomes:\n"
        "Complete the formal summative activity or activities scheduled for this month.\n"
        "Complete formative notebook, worksheet, drawing, digital, STEM project, or reflection evidence."
    ),
}


def main():
    update_docx(
        "plans/3rd Grade Technology/3rd Grade Monthly Planning/2nd Trimester/3° Technology - August.docx",
        meta=GRADE3_GENERIC_META,
        lessons={
            (4, 1): (
                "Topic:\nSTEM Sound/Action Practice\n\n"
                "Pre-Activities:\nPractice Keywords:\ninstrument: a tool that makes sound\ntest: try something to see what happens\n"
                "Review examples of simple digital instruments or sound buttons.\n\n"
                "While-Activities:\nBuild or trace 3-4 sound/action commands.\nTest and revise one sound or action if possible.\n"
                "Paper fallback: design an instrument button map.\nRecord one STEM idea: input, action/sound, and result.\n\n"
                "Post-Activities:\nSave practice evidence.\nWrite one change you made."
            ),
            (7, 0): (
                "Topic:\nSTEM Mini-Project Preparation\n\n"
                "Pre-Activities:\nPractice Keywords:\nreflection: thinking about your work\nsubmit: give work to the teacher\n"
                "Review the checklist for Summative #5.\n\n"
                "While-Activities:\nFinish a short sound/action sequence or paper instrument map.\n"
                "Write a reflection sentence: When I click/press ___, it ___.\n"
                "Check that the project has a chosen object, command order, test/trace, and result.\n\n"
                "Post-Activities:\nCheck evidence with a partner."
            ),
            (5, 1): (
                "Topic:\nScratch Sound STEM Mini-Project\n\n"
                "Pre-Activities:\nPractice Keywords:\ncomplete: finished with needed parts\nexplain: tell how something works\n"
                "Review submission options.\n\n"
                "While-Activities:\nComplete Summative Activity #5 (Summative): Submit a Scratch Sound STEM Mini-Project with a chosen sprite/object, 3-4 ordered sound/action commands, one test or trace, and a short reflection about what happened.\n\n"
                "Post-Activities:\nSubmit evidence.\nWrite one thing that worked well."
            ),
        },
    )

    update_docx(
        "plans/3rd Grade Technology/3rd Grade Monthly Planning/3rd Trimester/3° Technology - September.docx",
        meta=GRADE3_GENERIC_META,
        lessons={
            (1, 0): (
                "Topic:\nWhat Is Canva?\n\n"
                "Pre-Activities:\nPractice Keywords:\nCanva: a tool for making designs\ntemplate: a ready design that can be changed\n"
                "Look at sample STEM project display posters or invitations.\n\n"
                "While-Activities:\nLabel basic Canva areas or a printed Canva screenshot.\nChoose one template and identify title, image/icon, color, and text.\n\n"
                "Post-Activities:\nShare one Canva tool you noticed."
            ),
            (1, 1): (
                "Topic:\nSTEM Display Template Practice\n\n"
                "Pre-Activities:\nPractice Keywords:\ntool: something used to make or change work\nedit: change or improve\n"
                "Review title, image, and color.\n\n"
                "While-Activities:\nModify a practice template or paper template to show a Scratch Sound STEM idea: title, object/sprite, input/action, and result.\nSave screenshot or notebook version.\n\n"
                "Post-Activities:\nWrite one change you made."
            ),
            (4, 0): (
                "Topic:\nSTEM Display Tool Review\n\n"
                "Pre-Activities:\nPractice Keywords:\nsubmit: give work to the teacher\nevidence: work that shows learning\n"
                "Review the summative checklist.\n\n"
                "While-Activities:\nPractice selecting a template, editing the title, changing color/image/icon, and saving or showing work.\nPaper fallback: edit a printed template.\n\n"
                "Post-Activities:\nAsk one question before the check."
            ),
            (4, 1): (
                "Topic:\nSTEM Project Display/Template Check\n\n"
                "Pre-Activities:\nPractice Keywords:\nreadable: easy to see and understand\ncomplete: finished with needed parts\n"
                "Open Canva or the paper fallback.\n\n"
                "While-Activities:\nComplete Summative Activity #1 (Summative): Select a template, edit the title for a Scratch Sound STEM Mini-Project display, add or change one image/icon/color, and save/submit evidence.\nAccept screenshot, exported image, link, or paper template.\n\n"
                "Post-Activities:\nSubmit evidence.\nWrite one Canva tool you used."
            ),
        },
    )

    update_docx(
        "plans/4th Grade Technology/4th Grade Monthly Planning/2nd Trimester/4° Technology - August.docx",
        meta={
            1: (
                "Competences:\n"
                "Test an inertia car STEM investigation safely and explain the result with simple motion language.\n"
                "Navigate a block-programming environment with teacher support.\n"
                "Create or trace a simple STEM movement sequence and debug one issue.\n"
                "Reflect on movement, sequence, and responsible tool use."
            ),
            2: (
                "Learning Objectives:\n"
                "Record a simple inertia car test with prediction, result, and explanation.\n"
                "Use motion blocks or command cards to create a short movement sequence.\n"
                "Test and revise a movement sequence.\n"
                "Use sprite motion enrichment to reinforce movement blocks without adding a new formal grade."
            ),
            3: (
                "Learning Outcomes:\n"
                "Complete Summative Activity #4: inertia car STEM investigation evidence.\n"
                "Complete Summative Activity #5: block-programming STEM movement challenge.\n"
                "Complete formative app tutorial practice, guided motion enrichment, review, and reflection."
            ),
        },
        lessons={
            (1, 1): (
                "Topic:\nInertia Car STEM Investigation Evidence\n\n"
                "Pre-Activities:\nPractice Keywords:\nforce: a push or pull\nexplain: tell why something happened\n"
                "Review the summative checklist.\nSet up physical, video, or diagram-based test evidence.\n\n"
                "While-Activities:\nComplete Summative Activity #4 (Summative): Submit an inertia car STEM investigation with a car sketch/model, motion prediction, test result, and simple inertia explanation.\nAcceptable evidence: real test sheet, video-observation sheet, diagram test, or paper model with written result.\n\n"
                "Post-Activities:\nSubmit the test evidence.\nWrite one sentence: The car kept moving because __________."
            ),
            (7, 1): (
                "Topic:\nSTEM Movement Challenge\n\n"
                "Pre-Activities:\nPractice Keywords:\nsubmit: give work to the teacher\nreflect: think about what worked\n"
                "Review final evidence options: screenshot, file, notebook diagram, printed blocks, or command-card trace.\n\n"
                "While-Activities:\nComplete Summative Activity #5 (Summative): Create or trace a tested STEM movement challenge using motion blocks or command cards. Include the sequence, one debug note, and evidence of the final movement.\nAcceptable tools: Scratch, Blockly, robot app, printed block cards, or grid paper.\n\n"
                "Post-Activities:\nSubmit the evidence.\nWrite one sentence: I fixed __________ by __________."
            ),
        },
    )

    update_docx(
        "plans/4th Grade Technology/4th Grade Monthly Planning/3rd Trimester/4° Technology - September.docx",
        meta={
            1: (
                "Competences:\n"
                "Ask simple, safe STEM questions that can be answered with data.\n"
                "Collect and record a small data set accurately.\n"
                "Use tables to organize observations over time.\n"
                "Protect private information when collecting class data."
            ),
            2: (
                "Learning Objectives:\n"
                "Write or choose a STEM data question.\n"
                "Collect 8-10 safe responses or use a teacher-provided STEM data set.\n"
                "Organize data in a table with clear categories.\n"
                "Begin a simple log over time."
            ),
            3: (
                "Learning Outcomes:\n"
                "Complete Summative Activity #1: STEM data question and collection table.\n"
                "Complete formative data-question practice and data logging notes."
            ),
        },
        lessons={
            (1, 0): (
                "Topic:\nSTEM Data Questions\n\n"
                "Pre-Activities:\nPractice Keywords:\ndata: information collected to answer a question\nquestion: something we ask to learn information\n"
                "Look at sample class-safe STEM questions, such as Which car test went farther? Which movement command worked best? or Which surface slowed the car?\nSort questions into safe/not safe for class data.\n\n"
                "While-Activities:\nChoose or write one STEM data question connected to the inertia car or movement challenge.\nIdentify possible answer categories.\nPractice recording 3 sample answers in a table.\n\n"
                "Post-Activities:\nShare one safe STEM data question.\nRevise one question if it is too private or unclear."
            ),
            (1, 1): (
                "Topic:\nSTEM Data Practice\n\n"
                "Pre-Activities:\nPractice Keywords:\nsurvey: a way to collect answers from people\ncategory: a group used to organize answers\n"
                "Review safe survey behavior: ask politely, no private data, record honestly.\n\n"
                "While-Activities:\nPractice collecting 5 class-safe STEM responses or use teacher-provided inertia/movement data cards.\nOrganize responses into categories.\nCheck for clear handwriting or clean digital entry.\n\n"
                "Post-Activities:\nCount one category total.\nWrite one rule for honest data collection."
            ),
            (4, 1): (
                "Topic:\nSTEM Data Question and Collection Table\n\n"
                "Pre-Activities:\nPractice Keywords:\nevidence: work that shows learning\nsubmit: give work to the teacher\n"
                "Review the summative checklist: question, 8-10 data entries, categories, clear recording.\n\n"
                "While-Activities:\nComplete Summative Activity #1 (Summative): Create or complete a safe STEM data question and collection table with 8-10 entries and clear categories from inertia car, movement, or teacher-provided STEM data cards.\nUse Google Sheets, notebook table, printed worksheet, or teacher-provided data cards.\n\n"
                "Post-Activities:\nSubmit the table.\nWrite one answer your data can help explain."
            ),
        },
    )

    update_docx(
        "plans/5th Grade Technology/5th Grade Monthly Planning/2nd Trimester/5° Technology - August.docx",
        meta={
            1: (
                "Competences:\n"
                "Explain how inputs, conditions, choices, and outputs work in simple everyday and computer-style examples.\n"
                "Test and debug short if-then routines using cards, Scratch, Blockly, or paper fallback options.\n"
                "Design a low-risk simulated STEM routine with one input, one condition, one output, and one test result."
            ),
            2: (
                "Learning Objectives:\n"
                "Identify inputs, outputs, conditions, choices, true/false results, and simple logic errors.\n"
                "Complete and correct if-then examples with visible condition and output labels.\n"
                "Create a short Interactive Routine STEM Project and prepare Trimester 2 review or make-up evidence."
            ),
            3: (
                "Learning Outcomes:\n"
                "Complete Summative Activity #5: Interactive Routine STEM Project in Week 11.\n"
                "Complete formative if-then practice, debugging preparation, review game, and evidence workshop tasks.\n"
                "Use Week 13 only for exam-week quiet review, organization, or school-directed tasks."
            ),
        },
        lessons={
            (1, 1): (
                "Topic:\nSelection and Debugging Preparation\n"
                "Pre-Activities:\nPractice Keywords:\nselection: choosing an action based on a condition\ndebug: to find and correct an error\n"
                "Review condition, action, true, false, and debug.\nPractice finding the error in one sample command.\n"
                "While-Activities:\nComplete a formative selection and debugging practice check with 6 if-then examples, 4 input-output matches, and 2 simple command errors to correct.\nLabel each condition and output clearly.\n"
                "Post-Activities:\nSubmit practice evidence.\nWrite one debugging habit you used.\n"
                "Resources:\nnotebook, logic cards, worksheet, paper fallback."
            ),
            (3, 1): (
                "Topic:\nSelection Practice and Debugging Review\n"
                "Pre-Activities:\nPractice Keywords:\npractice: work done to improve a skill\nevidence: work that shows what you learned\n"
                "Review the Week 9 formative debugging preparation.\nChoose one error type to practice again.\n"
                "While-Activities:\nCorrect short if-then examples with a partner.\nCreate one new input-output example.\nExplain why the corrected logic works.\n"
                "Post-Activities:\nShare one corrected example.\nWrite one debugging habit you want to keep using.\n"
                "Resources:\nnotebook, logic cards, worksheet, paper fallback."
            ),
            (5, 1): (
                "Topic:\nInteractive Routine STEM Project\n"
                "Pre-Activities:\nPractice Keywords:\nlogic: clear thinking about steps and rules\noutcome: what happens at the end\n"
                "Review your routine drawing.\nPractice explaining condition and output.\n"
                "While-Activities:\nComplete Summative Activity #5 (Summative): Create a short paper, Scratch, or Blockly-style Interactive Routine STEM Project with one input, one condition, one output, one test result, and one improvement note.\nHardware is optional; paper, Scratch, or Blockly simulation is the default evidence.\n"
                "Post-Activities:\nPresent your STEM routine to a small group.\nWrite one change that made the routine clearer.\n"
                "Resources:\nnotebook, logic cards, worksheet, paper fallback."
            ),
            (7, 1): (
                "Topic:\nMake-Up and Evidence Workshop\n"
                "Pre-Activities:\nPractice Keywords:\nevidence: work that shows what you learned\nchecklist: a list used to check finished work\n"
                "Review your Trimester 2 evidence checklist.\nChoose one missing or weak product to improve.\n"
                "While-Activities:\nComplete make-up work for systems, search, storyboard, visual product, or the Interactive Routine STEM Project.\nAsk a peer to review one product with the checklist.\nRevise one item before submitting.\n"
                "Post-Activities:\nSubmit updated evidence.\nWrite one topic you understand better now.\n"
                "Resources:\nnotebook, logic cards, worksheet, paper fallback."
            ),
        },
    )

    update_docx(
        "plans/5th Grade Technology/5th Grade Monthly Planning/3rd Trimester/5° Technology - September.docx",
        meta={
            1: (
                "Competences:\n"
                "Organize STEM project information with forms, fields, records, data cards, rows, columns, sorting, and grouping.\n"
                "Compare paper databases and computer-style tables using simple evidence.\n"
                "Answer questions from organized data using clear field, record, sort, group, filter, and value language."
            ),
            2: (
                "Learning Objectives:\n"
                "Create data cards from Interactive Routine STEM Project information and label fields and records.\n"
                "Sort and group data cards by different fields or values.\n"
                "Use a paper table or approved tool to search, filter, and answer simple data questions."
            ),
            3: (
                "Learning Outcomes:\n"
                "Complete Summative Activity #1: STEM project fields, records, and sorting check.\n"
                "Complete formative paper database creation, card sorting, computer-style table practice, and search practice.\n"
                "Use database vocabulary in notebook evidence and short explanations."
            ),
        },
        lessons={
            (1, 0): (
                "Topic:\nSTEM Fields and Records\n"
                "Pre-Activities:\nPractice Keywords:\nfield: one type of information in a database\nrecord: one complete set of information\n"
                "Match database words with examples: field = category, record = one complete item, form = place to enter data, data = information.\nAnswer one question about why people organize STEM project information.\n"
                "While-Activities:\nFill out a simple paper form about the Interactive Routine STEM Project with fields such as input, condition, output, test result, and improvement.\nLabel the fields and records in the example.\n"
                "Post-Activities:\nShare one STEM project field.\nWrite one reason forms help organize data.\n"
                "Resources:\nnotebook, forms, data cards, sorting checklist."
            ),
            (1, 1): (
                "Topic:\nCreate STEM Project Data Cards\n"
                "Pre-Activities:\nPractice Keywords:\ndatabase: an organized collection of information\ndata card: a card that stores one record\n"
                "Review field and record examples.\nChoose classroom-safe STEM project data cards.\n"
                "While-Activities:\nCreate a small set of at least 6 STEM project data cards using the same fields.\nCheck that every card has complete information.\nSort the cards by one field.\n"
                "Post-Activities:\nExplain your sorting rule.\nStore your data cards for next class.\n"
                "Resources:\nnotebook, forms, data cards, sorting checklist."
            ),
            (3, 1): (
                "Topic:\nSTEM Fields, Records, and Sorting Check\n"
                "Pre-Activities:\nPractice Keywords:\nfield: one category of information\nrecord: one row or item in a database\n"
                "Review field, record, sort, group, and question.\nPractice one sorting example.\n"
                "While-Activities:\nComplete Summative Activity #1 (Summative): Organize 12 STEM project data cards, label 4 fields, identify 3 records, sort the cards by one field, group them by one value, and answer 3 questions using the data.\nCheck that each answer uses evidence from the cards.\n"
                "Post-Activities:\nSubmit your sorting check.\nWrite one database word you understand well.\n"
                "Resources:\nnotebook, forms, data cards, sorting checklist."
            ),
        },
    )

    grade6_aug_resources = "Resources: Pen, pencil, technology notebook, student's email and password, charged and numbered mBots, mBot Knowledge Cards Part 1, final challenge planning sheet, mBot readiness checklist, route/test data sheet, floor tape/markers, projector, computers/software."
    update_docx(
        "plans/6th Grade Technology/Planning/Monthly/2nd Trimester/6° Technology - August.docx",
        meta={
            1: (
                "Competences:\n"
                "Use mBot Knowledge Cards to review movement, output, sensor, and debugging tasks.\n"
                "Combine movement with one output or sensor in a controlled mBot STEM Challenge Project.\n"
                "Demonstrate, explain, and document a robot program using clear testing and improvement evidence.\n"
                "Leave mBots ready with numbered kits, saved code/screenshot, challenge sheet, route/test data, and sensor/output notes."
            ),
            2: (
                "Learning Objectives:\n"
                "Read and apply selected mBot Knowledge Cards as formative practice stations.\n"
                "Plan, build, and test a realistic mBot STEM challenge with a goal and success rule.\n"
                "Explain the robot action, one card, block, or part used, and one improvement made after testing.\n"
                "Prepare robot readiness evidence for the first weeks of Trimester 3."
            ),
            3: (
                "Learning Outcomes:\n"
                "Complete mBot card practice with recorded robot action and explanation as formative preparation.\n"
                "Submit a final STEM challenge plan and station notes showing testing and debugging.\n"
                "Demonstrate a final mBot STEM challenge and complete a reflection about goal, action, improvement, and readiness for the next project."
            ),
        },
        lessons={
            (1, 0): (
                "Topic:\nKnowledge Cards review\n\n"
                "Pre-activities\nAnswer the warm-up question: Which mBot card helped you the most?\nReview August Week 1 practice vocabulary (part 1).\nmovement: changing position\noutput: what a device shows or does\nsensor: part that detects something\ndebug: find and fix a problem\nReview the assigned mBot Knowledge Card Part 1.\n\n"
                "While activities\nPractice reading one mBot Knowledge Card and predicting what robot action it helps create.\nReview how the card could support a STEM challenge.\n\n"
                "Post activities\nShare the robot action from the assigned card.\nWrite one way the card could help in the final challenge."
            ),
            (1, 1): (
                "Topic:\nCard practice stations\n\n"
                "Pre-activities\nReview August Week 1 practice vocabulary and the selected cards (part 2).\ncard: small guide\nstation: place for one activity\naction: something someone does\nChoose two stations to complete first.\n\n"
                "While activities\nComplete formative mBot card practice: use one assigned mBot Knowledge Card Part 1, complete the card practice task, record the robot action, and explain what the card helps the robot do.\nComplete two mBot practice stations.\nTest the robot action at each station.\nRecord what worked and what was hard.\nAsk for help after trying one debugging step.\n\n"
                "Post activities\nSubmit your station notes.\nWrite one card you want to use in the STEM challenge."
            ),
            (4, 0): (
                "Topic: Plan the mBot STEM challenge\n\n"
                "Pre-activities\nAnswer the warm-up question: What task can your robot complete successfully?\nReview August Week 2 practice vocabulary.\ngoal: what you are trying to do\nsuccess rule: rule that says the task worked\ntest area: space used for trying work\nChoose one challenge type: path, obstacle, signal, or line/sensor task.\n\n"
                "While activities\nPlan the mBot STEM challenge.\nWrite the goal and success rule.\nList the cards, blocks, or parts you will use.\nCheck that the plan can be completed in class time.\n\n"
                "Post activities\nSubmit the STEM challenge plan.\nExplain your success rule to a partner."
            ),
            (7, 1): (
                "Topic: Demonstrate, reflect, and leave mBots ready\n\n"
                "Pre-activities\nReview your demonstration notes.\nPrepare the mBot and test area.\nComplete a vocabulary review activity using August Week 3 practice vocabulary.\naudience: people who see or hear work\nexplain: make clear with words\nreflection: thinking about what worked\n\n"
                "While activities\nComplete Summative Activity #5: Demonstrate your mBot STEM Challenge Project.\nExplain the goal, route/behavior, main robot action, output or sensor, test result, and one improvement.\nSave code or screenshot evidence.\nComplete the readiness checklist: numbered kit, charged battery, challenge sheet, route/test data, and sensor/output notes.\n\n"
                "Post activities\nComplete the final robotics reflection.\nSubmit the reflection, project evidence, and readiness checklist."
            ),
        },
        resources={2: grade6_aug_resources, 5: grade6_aug_resources, 8: grade6_aug_resources},
    )

    grade6_sep_resources = "Resources: Pen, pencil, technology notebook, student's email and password, ready mBots, mBot readiness checklist, mBot STEM data collection sheet, prepared mBot data-set template, formula practice file, Google Sheets or spreadsheet software access, projector, internet/accounts if using Google Sheets."
    update_docx(
        "plans/6th Grade Technology/Planning/Monthly/3rd Trimester/6° Technology - September.docx",
        weeks=3,
        meta={
            1: (
                "Competences:\n"
                "Inspect and reuse ready mBots safely for a second controlled STEM challenge.\n"
                "Organize mBot challenge data in a spreadsheet using rows, columns, cells, headers, and tables.\n"
                "Format spreadsheet information and use simple formulas to calculate values and explain how inputs affect outputs."
            ),
            2: (
                "Learning Objectives:\n"
                "Check robot readiness evidence from Trimester 2 and plan a second challenge.\n"
                "Collect mBot challenge data and enter it into a clean spreadsheet table.\n"
                "Practice basic formulas and cell references using the mBot data set."
            ),
            3: (
                "Learning Outcomes:\n"
                "Submit a clean mBot STEM data table based on the second challenge.\n"
                "Complete formula practice tasks using the mBot data set.\n"
                "Explain one input/output or quantity/result change from the project data."
            ),
        },
        lessons={
            (1, 0): (
                "Topic:\nInspect ready mBots\n\n"
                "Pre-activities\nAnswer the warm-up question: What evidence did we leave ready from Trimester 2?\nReview September Week 1 practice vocabulary (part 1).\ndata: information from a test\nrecord: one row\nheader: column label\ntable: rows and columns\n\n"
                "While activities\nInspect numbered mBots using the readiness checklist.\nMatch each robot with its challenge sheet, saved code/screenshot, route/test data, and sensor/output notes.\nIdentify what data could be collected in a second challenge.\n\n"
                "Post activities\nMark the robot ready, needs charging, or needs teacher help.\nWrite one data field you could record."
            ),
            (1, 1): (
                "Topic:\nPlan second mBot STEM challenge\n\n"
                "Pre-activities\nReview September Week 1 practice vocabulary (part 2).\nresult: what happened after a test\nreadiness: prepared to work safely\nChoose one safe challenge type: route time, obstacle success, signal accuracy, or sensor response.\n\n"
                "While activities\nPlan the second mBot STEM challenge.\nChoose the data fields, such as team, challenge type, attempts, time, success, output/sensor used, and improvement.\nCreate or prepare the spreadsheet table with title, headers, and units/categories.\n\n"
                "Post activities\nSubmit the challenge plan and table setup.\nExplain one field to a partner."
            ),
            (4, 0): (
                "Topic:\nCollect mBot data\n\n"
                "Pre-activities\nAnswer the warm-up question: What makes test data fair?\nReview September Week 2 practice vocabulary (part 1).\nformula: spreadsheet calculation\ncell reference: cell name like A1\nsum: answer from adding\naverage: usual or middle value\n\n"
                "While activities\nRun or observe the second mBot STEM challenge in short turns.\nRecord at least 6 clean records using the data collection sheet or teacher data cards.\n\n"
                "Post activities\nCheck one record for missing information.\nShare one result."
            ),
            (4, 1): (
                "Topic:\nmBot STEM Data Table\n\n"
                "Pre-activities\nComplete a vocabulary review activity using September Week 2 practice vocabulary (part 2).\ninput: information or action going in\noutput: what a device shows or does\nvalue: number or data\nOpen the mBot STEM data table.\n\n"
                "While activities\nComplete Summative Activity #1: Create a clean mBot STEM data table with at least 6 records, headers, units/categories, readable formatting, and one sentence explaining what the data shows.\nSubmit the table in class or on Google Classroom.\n\n"
                "Post activities\nSubmit the clean data table.\nWrite one sentence explaining how formatting helped."
            ),
            (7, 0): (
                "Topic:\nFormulas with mBot data\n\n"
                "Pre-activities\nAnswer the warm-up question: Which mBot result can we calculate?\nReview formula, cell reference, sum, average, input, output, and value.\n\n"
                "While activities\nUse formulas on the mBot data table.\nCalculate one total, count, or average.\nPredict the answer before calculating.\nRecord one input and output example.\n\n"
                "Post activities\nWrite one sentence explaining what the formula calculated.\nShare one result."
            ),
            (7, 1): (
                "Topic:\nFormula practice with project data\n\n"
                "Pre-activities\nOpen the prepared mBot formula practice file or class data table.\nReview how to copy a formula without changing the headers.\n\n"
                "While activities\nCalculate totals and averages from mBot data.\nChange one input value and observe what changes.\nRecord one input/output change and one question a chart could answer next week.\n\n"
                "Post activities\nSubmit the formula practice task.\nComplete a reflection: What changed when the input changed?"
            ),
        },
        resources={2: grade6_sep_resources, 5: grade6_sep_resources, 8: grade6_sep_resources},
        delete_lesson_rows_from=9,
    )

    grade6_oct_resources = "Resources: Pen, pencil, technology notebook, student's email and password, mBot STEM data table, mBot chart checklist, Tinkercad or 3D tool accounts, 3D model design-plan template, screenshot/export instructions, projector, internet access."
    update_docx(
        "plans/6th Grade Technology/Planning/Monthly/3rd Trimester/6° Technology - October.docx",
        meta={
            1: (
                "Competences:\n"
                "Use spreadsheets to present mBot STEM project data with charts.\n"
                "Interpret chart information and write a simple conclusion from project data.\n"
                "Create and improve a basic 3D model from a labeled design plan."
            ),
            2: (
                "Learning Objectives:\n"
                "Review a mBot STEM data table and choose a chart that answers a project question.\n"
                "Create a chart from the mBot STEM data set with title, labels, answers, and a conclusion.\n"
                "Use basic 3D modelling actions such as move, resize, rotate, duplicate, group, and improve."
            ),
            3: (
                "Learning Outcomes:\n"
                "Submit a mBot STEM chart task with title, labels, answers, and conclusion.\n"
                "Submit 3D modelling practice evidence showing basic shape tools.\n"
                "Complete a 3D model design plan with sketch, at least three labeled shapes, purpose, and planned improvement."
            ),
        },
        lessons={
            (1, 0): (
                "Topic:\nmBot charts tell a story\n\n"
                "Pre-activities\nAnswer the warm-up question: When is a chart easier than a table?\nReview October Week 1 practice vocabulary (part 1).\nchart: visual way to show data\ntitle: name of a page or project\nlabel: word that names something\ncategory: group\nOpen the mBot STEM data table from September.\n\n"
                "While activities\nIdentify the title, labels, and categories in a sample chart.\nChoose one mBot data question that could become a chart.\nExplain what the chart should show.\n\n"
                "Post activities\nWrite one sentence explaining what charts help people understand.\nShare your chart idea with a partner."
            ),
            (1, 1): (
                "Topic:\nmBot Data Chart and Conclusion\n\n"
                "Pre-activities\nComplete a vocabulary review activity using October Week 1 practice vocabulary (part 2).\naxis: chart line for values or groups\ncompare: look for similarities or differences\nconclusion: idea from evidence\nOpen the mBot STEM data set.\n\n"
                "While activities\nComplete Summative Activity #2: Create one chart from the mBot STEM data set, add a title and labels, answer 3 questions about the chart, and write one conclusion sentence.\nSubmit the chart task in class or on Google Classroom.\n\n"
                "Post activities\nShare one conclusion from your chart.\nWrite one situation where a chart is better than a table."
            ),
            (4, 0): (
                "Topic:\n3D modelling basics\n\n"
                "Pre-activities\nAnswer the warm-up question: What objects are made from simple shapes?\nReview October Week 2 practice vocabulary (part 1).\n3D: height, width, and depth\nshape: form like cube or sphere\nmove: change place\nresize: make bigger or smaller\nIdentify basic shapes in a real object.\n\n"
                "While activities\nOpen the 3D modelling tool.\nCreate a basic shape.\nPractice moving, resizing, and rotating the shape.\nRecord one tool you used.\n\n"
                "Post activities\nWrite one sentence explaining what 3D means.\nSave or screenshot your first shape practice."
            ),
            (4, 1): (
                "Topic:\nShape practice\n\n"
                "Pre-activities\nComplete a word soup or matching activity using October Week 2 practice vocabulary (part 2).\nrotate: turn around\nduplicate: make a copy\nOpen your shape practice file.\n\n"
                "While activities\nCreate a simple object using basic shapes.\nMove, resize, rotate, and duplicate shapes.\nChallenge yourself to make the object more useful or interesting.\nSave or screenshot your progress.\n\n"
                "Post activities\nShow your object to a partner.\nWrite one thing you changed while building."
            ),
            (7, 0): (
                "Topic:\nPlan a simple model\n\n"
                "Pre-activities\nAnswer the warm-up question: Why should we plan before building?\nReview October Week 3 practice vocabulary (part 1).\nmodel: simple version or design\ndesign: plan for how something works\nprototype: first version for testing\nsketch: quick drawing\nSketch a small model, such as a name badge, desk item, or project part.\n\n"
                "While activities\nPractice labeling at least three shapes on a draft sketch.\nDiscuss the purpose of the model and one possible improvement.\nReview the 3D model design plan expectations.\n\n"
                "Post activities\nShare the purpose of your model with a partner.\nCircle the part of your sketch you plan to improve."
            ),
            (7, 1): (
                "Topic:\nDesign-plan workshop\n\n"
                "Pre-activities\nReview October Week 3 practice vocabulary (part 2).\ngroup: join objects together\nimprove: make better\npurpose: reason something is made\nfeedback: helpful comments\nOpen your model sketch and 3D file.\n\n"
                "While activities\nCreate your model from the sketch.\nUse at least three shapes.\nGroup or combine shapes when useful.\nImprove one part after viewing or testing the model.\nPrepare your design-plan evidence for the summative check next week.\n\n"
                "Post activities\nSave or screenshot your model.\nWrite one planned improvement."
            ),
            (10, 0): (
                "Topic:\n3D model evidence review\n\n"
                "Pre-activities\nReview the 3D model design plan checklist: sketch, 3 labeled shapes, model purpose, and one planned improvement.\n\n"
                "While activities\nCheck your sketch and model evidence.\nFix missing labels, unclear purpose, or missing improvement note.\n\n"
                "Post activities\nAsk one final question before the summative."
            ),
            (10, 1): (
                "Topic:\n3D Model Design Plan\n\n"
                "Pre-activities\nReview October Week 3 practice vocabulary with Blooket or a crossword puzzle.\nOpen your design plan and model evidence.\n\n"
                "While activities\nComplete Summative Activity #3: Create a 3D model design plan with a sketch, at least 3 labeled shapes, the model purpose, and one planned improvement.\nSubmit the design plan in class or on Google Classroom.\nSubmit or screenshot your model as practice evidence.\n\n"
                "Post activities\nReflect on what you improved and why.\nPrepare to share the project in November."
            ),
        },
        resources={2: grade6_oct_resources, 5: grade6_oct_resources, 8: grade6_oct_resources, 11: grade6_oct_resources},
        delete_lesson_rows_from=12,
    )


if __name__ == "__main__":
    main()
