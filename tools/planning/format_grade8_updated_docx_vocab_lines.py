from pathlib import Path
import json
import re

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX_DIR = ROOT / "plans/8th grade/8th Grade Technology - Updated"
VOCAB_DIR = ROOT / "apps" / "sparks" / "vocabularies" / "grade8"


TOPIC_WEEKS = {
    "8° Technology - March.docx": {
        "Alignment, distribution, and grouping": ("Week 2", "grade8_it_march_week2_alignment_grouping.json"),
        "Paths and nodes": ("Week 2", "grade8_it_march_week2_alignment_grouping.json"),
        "Campaign graphic planning": ("Week 3", "grade8_it_march_week3_campaign_graphics.json"),
        "Campaign icon or logo draft": ("Week 3", "grade8_it_march_week3_campaign_graphics.json"),
        "SVG and markup basics": ("Week 4", "grade8_it_march_week4_svg_markup.json"),
        "Design review and improvement": ("Week 4", "grade8_it_march_week4_svg_markup.json"),
    },
    "8° Technology - April.docx": {
        "Computing systems": ("Week 5", "grade8_it_april_week5_computing_systems.json"),
        "Operating systems and logical thinking": ("Week 5", "grade8_it_april_week5_computing_systems.json"),
        "AI and machine learning": ("Week 6", "grade8_it_april_week6_ai_responsible_remix.json"),
        "Open source and responsible remixing": ("Week 6", "grade8_it_april_week6_ai_responsible_remix.json"),
        "HTML structure": ("Week 7", "grade8_it_april_week7_html_css.json"),
        "Images and CSS styling": ("Week 7", "grade8_it_april_week7_html_css.json"),
        "Search engines and keywords": ("Week 8", "grade8_it_april_week8_search_portfolio.json"),
        "Portfolio content and video organization": ("Week 8", "grade8_it_april_week8_search_portfolio.json"),
    },
    "8° Technology - May.docx": {
        "Project evidence check": ("Week 9", "grade8_it_may_week9_portfolio_planning.json"),
        "Google Sites setup": ("Week 9", "grade8_it_may_week9_portfolio_planning.json"),
        "Exam project launch": ("Week 10", "grade8_it_may_week10_homepage_media.json"),
        "Portfolio build and media integration": ("Week 10", "grade8_it_may_week10_homepage_media.json"),
        "Testing and feedback": ("Week 11", "grade8_it_may_week11_testing_revision.json"),
        "Final revision and reflection": ("Week 11", "grade8_it_may_week11_testing_revision.json"),
        "Presentation setup": ("Week 12", "grade8_it_may_week12_presentation.json"),
        "Presentation continuation": ("Week 12", "grade8_it_may_week12_presentation.json"),
        "Exam/admin buffer": ("Week 13", "grade8_it_may_week13_archive_reflection.json"),
        "Trimester reflection and archive": ("Week 13", "grade8_it_may_week13_archive_reflection.json"),
    },
    "8° Technology - June.docx": {
        "Flowing LED sequence": ("Week 2", "grade8_iit_june_week2_led_button.json"),
        "Digital input with a button": ("Week 2", "grade8_iit_june_week2_led_button.json"),
        "Button logic and state changes": ("Week 3", "grade8_iit_june_week3_state_serial.json"),
        "Serial monitor basics": ("Week 3", "grade8_iit_june_week3_state_serial.json"),
        "Debugging circuits": ("Week 4", "grade8_iit_june_week4_safe_debugging.json"),
        "Safe build challenge": ("Week 4", "grade8_iit_june_week4_safe_debugging.json"),
    },
    "8° Technology - July.docx": {
        "Digital vs analog signals": ("Week 5", "grade8_iit_july_week5_analog_values.json"),
        "Potentiometer data and LED brightness": ("Week 5", "grade8_iit_july_week5_analog_values.json"),
        "Photoresistor and environmental input": ("Week 6", "grade8_iit_july_week6_light_rgb.json"),
        "RGB LED and color output": ("Week 6", "grade8_iit_july_week6_light_rgb.json"),
        "Buzzer output": ("Week 7", "grade8_iit_july_week7_buzzer_servo.json"),
        "Servo control": ("Week 7", "grade8_iit_july_week7_buzzer_servo.json"),
        "Conditional logic with sensors": ("Week 8", "grade8_iit_july_week8_integrated_challenge.json"),
        "Integrated practice challenge": ("Week 8", "grade8_iit_july_week8_integrated_challenge.json"),
    },
    "8° Technology - August.docx": {
        "Ultrasonic sensing": ("Week 9", "grade8_iit_august_week9_ultrasonic_design.json"),
        "Project design plan": ("Week 9", "grade8_iit_august_week9_ultrasonic_design.json"),
        "Exam project launch": ("Week 10", "grade8_iit_august_week10_prototype_build.json"),
        "Prototype building and code": ("Week 10", "grade8_iit_august_week10_prototype_build.json"),
        "Testing and debugging": ("Week 11", "grade8_iit_august_week11_reliability.json"),
        "Improvement and documentation": ("Week 11", "grade8_iit_august_week11_reliability.json"),
        "Demonstration setup": ("Week 12", "grade8_iit_august_week12_demonstration.json"),
        "Demonstration continuation": ("Week 12", "grade8_iit_august_week12_demonstration.json"),
        "Exam/admin buffer": ("Week 13", "grade8_iit_august_week13_cleanup.json"),
        "Trimester reflection and cleanup": ("Week 13", "grade8_iit_august_week13_cleanup.json"),
    },
    "8° Technology - September.docx": {
        "Events and screen flow": ("Week 2", "grade8_iiit_september_week2_events_flow.json"),
        "Tappy Tap-style app practice": ("Week 2", "grade8_iiit_september_week2_events_flow.json"),
        "User needs and feedback": ("Week 3", "grade8_iiit_september_week3_user_needs.json"),
        "App plan and event map": ("Week 3", "grade8_iiit_september_week3_user_needs.json"),
    },
    "8° Technology - October.docx": {
        "Algorithms and first Python output": ("Week 4", "grade8_iiit_october_week4_first_python.json"),
        "Variables and input": ("Week 4", "grade8_iiit_october_week4_first_python.json"),
        "Numbers and arithmetic": ("Week 5", "grade8_iiit_october_week5_numbers_selection.json"),
        "Selection with if/else": ("Week 5", "grade8_iiit_october_week5_numbers_selection.json"),
        "Random numbers and multi-branch choices": ("Week 6", "grade8_iiit_october_week6_loops_branches.json"),
        "Loops and counters": ("Week 6", "grade8_iiit_october_week6_loops_branches.json"),
        "Debugging selection and loops": ("Week 7", "grade8_iiit_october_week7_debugging.json"),
        "Scratch + Micro:bit setup": ("Week 7", "grade8_iiit_october_week8_microbit_setup.json"),
        "Sensor-controlled game planning": ("Week 8", "grade8_iiit_october_week8_microbit_setup.json"),
        "Prototype and habits check": ("Week 8", "grade8_iiit_october_week8_microbit_setup.json"),
    },
    "8° Technology - November.docx": {
        "Representation and binary digits": ("Week 9", "grade8_iiit_november_week9_binary_representation.json"),
        "Project readiness and final planning": ("Week 9", "grade8_iiit_november_week9_binary_representation.json"),
        "Exam project launch": ("Week 10", "grade8_iiit_november_week10_game_build.json"),
        "Game building and sensor control": ("Week 10", "grade8_iiit_november_week10_game_build.json"),
        "Debugging and user testing": ("Week 11", "grade8_iiit_november_week11_user_testing.json"),
        "Documentation and presentation prep": ("Week 11", "grade8_iiit_november_week11_user_testing.json"),
        "Presentation setup": ("Week 12", "grade8_iiit_november_week12_presentation.json"),
        "Presentation continuation": ("Week 12", "grade8_iiit_november_week12_presentation.json"),
    },
}


def load_vocab(file_name):
    data = json.loads((VOCAB_DIR / file_name).read_text())
    return [(item["word"], item["definition"].rstrip(".")) for item in data["words"]]


def split_words(words):
    midpoint = (len(words) + 1) // 2
    return words[:midpoint], words[midpoint:]


def definition_lines(words):
    return "\n".join(f"{word}: {definition}" for word, definition in words)


def get_topic(text):
    match = re.search(r"Topic:\n+([^\n]+)", text)
    return match.group(1).strip() if match else ""


def insert_vocab_block(text, week, words, part_number):
    if definition_lines(words) in text:
        return text, False

    sentence = f"Review {week} practice vocabulary (part {part_number})."
    block = f"{sentence}\n{definition_lines(words)}"
    marker = "Pre-activities\n"

    if marker not in text:
        return text, False

    if f"{week} practice vocabulary" in text:
        pattern = re.compile(rf"([^\n.]*\b{re.escape(week)} practice vocabulary[^\n.]*\.)", re.I)
        match = pattern.search(text)
        if match:
            return text[:match.start(1)] + block + text[match.end(1):], True

    return text.replace(marker, marker + block + "\n", 1), True


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def find_docx(doc_name):
    direct = DOCX_DIR / doc_name
    if direct.exists():
        return direct
    matches = sorted(DOCX_DIR.glob(f"*/{doc_name}"))
    if matches:
        return matches[0]
    return direct


def main():
    total = 0
    for doc_name, topic_weeks in TOPIC_WEEKS.items():
        path = find_docx(doc_name)
        doc = Document(path)
        paragraphs = list(iter_paragraphs(doc))
        week_seen = {}
        doc_count = 0

        for paragraph in paragraphs:
            topic = get_topic(paragraph.text)
            if topic not in topic_weeks:
                continue
            week, vocab_file = topic_weeks[topic]
            words = load_vocab(vocab_file)
            first_half, second_half = split_words(words)
            occurrence = week_seen.get(week, 0) + 1
            week_seen[week] = occurrence
            part_number = 1 if occurrence == 1 else 2
            part_words = first_half if occurrence == 1 else second_half
            text, changed = insert_vocab_block(paragraph.text, week, part_words, part_number)
            if changed:
                set_paragraph_text(paragraph, text)
                doc_count += 1

        if doc_count:
            doc.save(path)
        total += doc_count
        print(f"{doc_name}: {doc_count} paragraphs updated")

    print(f"total paragraphs updated: {total}")


if __name__ == "__main__":
    main()
