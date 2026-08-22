from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX_DIR = ROOT / "plans/6th grade/6th Grade Technology - Updated"
VOCAB_DIR = ROOT / "apps" / "sparks" / "vocabularies" / "grade6"


DEFINITIONS = {
    "internet": "worldwide computer network",
    "website": "online place with pages",
    "browser": "app used to open websites",
    "address": "where to find a place online",
    "domain name": "readable website name",
    "DNS": "system that matches names to addresses",
    "IP address": "number address for a device or website",
    "packet": "small part of a message",
    "message": "information sent to someone",
    "network": "connected devices",
    "send": "move a message to someone",
    "receive": "get a message",
    "part": "one piece of something bigger",
    "order": "correct sequence",
    "missing": "not there when needed",
    "route": "path to a place",
    "delivery": "getting something to the right place",
    "communicate": "share information",
    "collaborate": "work together",
    "share": "let others use or see",
    "shared file": "file more than one person can use",
    "chat": "short online messages",
    "comment": "note added to work",
    "respectful": "kind and careful with others",
    "help": "make work easier for someone",
    "public": "seen by many people",
    "private": "kept for only certain people",
    "audience": "people who see or hear work",
    "permission": "saying it is okay",
    "web page": "one page on a website",
    "title": "name of a page or project",
    "menu": "list of choices",
    "link": "clickable connection",
    "image": "picture",
    "section": "part of a page",
    "copyright": "rule that protects created work",
    "credit": "say who made or owns something",
    "source": "where information or an image came from",
    "creator": "person who made something",
    "safe": "not likely to cause harm",
    "header": "top label or section",
    "text": "written words",
    "layout": "how things are arranged",
    "preview": "look before final",
    "readable": "easy to read",
    "home page": "main page of a website",
    "navigation": "moving around a site",
    "button": "thing you click or press",
    "page": "one screen or sheet of information",
    "test": "try and check",
    "feedback": "helpful comments",
    "Scratch": "block coding tool",
    "sprite": "character or object in Scratch",
    "block": "coding piece",
    "script": "set of coding blocks",
    "variable": "stores a changing value",
    "value": "number or data",
    "score": "points in a game",
    "change": "make different",
    "debug": "find and fix a problem",
    "kind": "friendly and caring",
    "reliable": "can be trusted",
    "rule": "what should be followed",
    "privacy": "keeping information protected",
    "campaign": "message to help people act",
    "poster": "visual message",
    "recycle": "turn waste into usable material",
    "reuse": "use again",
    "repair": "fix so it works again",
    "e-waste": "old electronic waste",
    "device": "technology tool",
    "battery": "part that stores power",
    "trash": "waste thrown away",
    "sorting": "putting into groups",
    "waste": "things thrown away",
    "environment": "the natural world around us",
    "fact": "true information",
    "tip": "helpful advice",
    "action": "something someone does",
    "slide": "presentation page",
    "infographic": "visual facts or tips",
    "present": "show and explain",
    "robot": "machine that follows instructions",
    "mBot": "small classroom robot",
    "program": "set of instructions",
    "instruction": "one step to follow",
    "safety": "protecting people and materials",
    "motor": "part that makes movement",
    "wheel": "round part for moving",
    "sensor": "part that detects something",
    "LED": "small light",
    "buzzer": "part that makes sound",
    "sequence": "steps in order",
    "forward": "toward the front",
    "backward": "toward the back",
    "turn": "change direction",
    "stop": "end movement",
    "speed": "how fast something moves",
    "time": "how long something lasts",
    "improve": "make better",
    "path": "way from start to finish",
    "problem": "something to solve",
    "fix": "correct a problem",
    "output": "what a device shows or does",
    "signal": "sign that gives information",
    "state": "what something is doing now",
    "start": "beginning",
    "finish": "end",
    "accurate": "correct and close to target",
    "marker": "object or line showing a place",
    "repeat": "do again",
    "loop": "code that repeats",
    "pattern": "repeated design or action",
    "square": "shape with four equal sides",
    "zigzag": "path that turns left and right",
    "shorter": "using less space or time",
    "detect": "notice or find",
    "obstacle": "thing in the way",
    "line": "long mark or path",
    "input": "information or action going in",
    "response": "what happens after",
    "if": "starts a condition",
    "condition": "rule that is checked",
    "true": "correct or happening",
    "false": "not correct or not happening",
    "flowchart": "diagram of steps or choices",
    "challenge": "task that takes effort",
    "goal": "what you are trying to do",
    "success rule": "rule that says the task worked",
    "test area": "space used for trying work",
    "commands": "program instructions",
    "blocks": "coding pieces",
    "movement": "changing position",
    "card": "small guide",
    "station": "place for one activity",
    "improvement": "change that makes better",
    "parts": "pieces of a device",
    "explain": "make clear with words",
    "reflection": "thinking about what worked",
    "spreadsheet": "digital table",
    "data": "collected facts or values",
    "information": "data that has meaning",
    "row": "line across a table",
    "column": "line down a table",
    "cell": "one spreadsheet box",
    "table": "rows and columns",
    "formula": "spreadsheet calculation",
    "chart": "visual way to show data",
    "cell reference": "cell name like A1",
    "sum": "answer from adding",
    "average": "usual or middle value",
    "budget": "money plan",
    "quantity": "how many",
    "cost": "price",
    "total": "whole amount",
    "item": "one thing in a list",
    "estimate": "careful guess",
    "materials": "things used for a project",
    "label": "word that names something",
    "category": "group",
    "axis": "chart line for values or groups",
    "compare": "look for similarities or differences",
    "conclusion": "idea from evidence",
    "3D": "height, width, and depth",
    "shape": "form like cube or sphere",
    "move": "change place",
    "resize": "make bigger or smaller",
    "rotate": "turn around",
    "duplicate": "make a copy",
    "model": "simple version or design",
    "design": "plan for how something works",
    "prototype": "first version for testing",
    "sketch": "quick drawing",
    "group": "join objects together",
    "purpose": "reason something is made",
    "artifact": "thing made in a project",
    "micro:bit": "tiny coding computer",
    "display": "screen or lights",
    "icon": "small picture or symbol",
    "increase": "make larger",
    "decrease": "make smaller",
    "counter": "program or value that counts",
    "reset": "set back to the start",
    "if/then": "rule where one action follows a condition",
    "light": "brightness",
    "temperature": "how hot or cold",
    "project": "work made to show or solve something",
    "demonstrate": "show how something works",
}


SUMMATIVE = {
    "t1": ["internet", "website", "browser", "address", "domain name", "DNS", "IP address", "packet", "message", "network"],
    "t2": ["robot", "mBot", "program", "instruction", "safety", "motor", "wheel", "sensor", "LED", "buzzer"],
    "t3": ["spreadsheet", "data", "information", "row", "column", "cell", "table", "header", "formula", "chart"],
}


PRACTICE = {
    "March Week 2": ["send", "receive", "part", "order", "missing", "route", "delivery"],
    "March Week 3": ["communicate", "collaborate", "share", "shared file", "chat", "comment", "respectful", "help"],
    "March Week 4": ["public", "private", "audience", "permission", "web page", "title", "menu", "link", "image", "section"],
    "April Week 1": ["copyright", "credit", "source", "creator", "safe", "permission", "image"],
    "April Week 2": ["title", "header", "section", "image", "text", "layout", "preview", "readable"],
    "April Week 3": ["link", "menu", "home page", "navigation", "button", "page", "test", "feedback"],
    "April Week 4": ["Scratch", "sprite", "block", "script", "variable", "value", "score", "change", "test", "debug"],
    "May Week 1": ["safe", "kind", "respectful", "reliable", "rule", "privacy", "campaign", "poster"],
    "May Week 2": ["recycle", "reuse", "repair", "e-waste", "device", "battery", "trash", "sorting", "waste", "environment"],
    "May Week 3": ["audience", "message", "title", "fact", "tip", "action", "poster", "slide", "infographic", "present"],
    "June Week 2": ["sequence", "forward", "backward", "turn", "stop", "speed", "time"],
    "June Week 3": ["test", "debug", "improve", "path", "route", "problem", "fix"],
    "June Week 4": ["output", "signal", "LED", "buzzer", "state", "start", "finish"],
    "July Week 1": ["speed", "time", "turn", "sequence", "route", "accurate", "marker"],
    "July Week 2": ["repeat", "loop", "pattern", "square", "zigzag", "shorter"],
    "July Week 3": ["sensor", "detect", "obstacle", "line", "input", "response"],
    "July Week 4": ["if", "condition", "response", "true", "false", "flowchart"],
    "July Week 5": ["challenge", "goal", "success rule", "test area", "commands", "blocks"],
    "August Week 1": ["movement", "output", "sensor", "debug", "card", "station", "action"],
    "August Week 2": ["goal", "success rule", "test area", "improvement", "challenge", "parts"],
    "August Week 3": ["challenge", "test", "debug", "improve", "audience", "explain", "reflection"],
    "September Week 2": ["formula", "cell reference", "sum", "average", "input", "output", "value"],
    "October Week 1": ["budget", "quantity", "cost", "total", "item", "estimate", "materials"],
    "October Week 2": ["chart", "title", "label", "category", "axis", "compare", "conclusion"],
    "October Week 3": ["3D", "shape", "move", "resize", "rotate", "duplicate", "model"],
    "October Week 4": ["design", "prototype", "sketch", "group", "improve", "purpose", "feedback"],
    "November Week 1": ["present", "feedback", "improve", "goal", "artifact", "reflection"],
    "November Week 2": ["micro:bit", "input", "output", "LED", "button", "display", "icon"],
    "November Week 3": ["variable", "value", "increase", "decrease", "counter", "reset"],
    "December Week 1": ["sensor", "condition", "if/then", "light", "temperature", "message", "flowchart"],
    "December Week 2": ["input", "output", "variable", "test", "improve", "success rule", "project"],
    "December Week 3": ["input", "output", "condition", "variable", "improve", "demonstrate", "explain"],
}


def words_only(words):
    return ", ".join(words)


def practice_text(label):
    words = PRACTICE[label]
    defs = "; ".join(f"{word} = {DEFINITIONS[word]}" for word in words)
    return f"Practice vocabulary ({label}): {defs}."


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return 0
    text = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return 1


def replace_in_doc(doc, replacements):
    count = 0
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for paragraph in all_paragraphs:
        for old, new in replacements:
            count += replace_in_paragraph(paragraph, old, new)
    return count


REPLACEMENTS = {
    "6° Technology - March.docx": [
        ("Match today's vocabulary words with their definitions: address, domain name, DNS.", f"Review Trimester 1 summative vocabulary: {words_only(SUMMATIVE['t1'])}."),
        ("Review last class vocabulary words with a quick Blooket or matching game: address, domain name, DNS.", f"Review Trimester 1 summative vocabulary with a quick Blooket or matching game: {words_only(SUMMATIVE['t1'])}."),
        ("Add today's vocabulary word to your notebook: IP address.", "No practice vocabulary this week; focus on the summative vocabulary table."),
        ("Create a vocabulary table with definition, illustration, and example sentence for these words: address, domain name, DNS, IP address, packet.", f"Create a vocabulary table with definition, illustration, and example sentence for these words: {words_only(SUMMATIVE['t1'])}."),
        ("Match today's vocabulary words with their definitions: packet, header, payload.", practice_text("March Week 2")),
        ("Complete a word soup or crossword puzzle using last class words: packet, header, payload.", f"Complete a word soup or crossword puzzle using March Week 2 practice vocabulary: {words_only(PRACTICE['March Week 2'])}."),
        ("Match today's words with examples: collaborate, shared file, chat.", practice_text("March Week 3")),
        ("Review last class vocabulary with Blooket, Gimkit, or a matching activity: collaborate, shared file, chat.", f"Review March Week 3 practice vocabulary with Blooket, Gimkit, or a matching activity: {words_only(PRACTICE['March Week 3'])}."),
        ("Match today's words with their definitions: public, private, audience, permission.", practice_text("March Week 4")),
        ("Complete a quick review puzzle using website words: website, web page, link, menu.", f"Complete a quick review puzzle using March Week 4 practice vocabulary: {words_only(PRACTICE['March Week 4'])}."),
    ],
    "6° Technology - April.docx": [
        ("Match today's words with definitions: copyright, credit, copyright-free.", practice_text("April Week 1")),
        ("Review last class words with a word soup or Blooket: copyright, credit, copyright-free.", f"Review April Week 1 practice vocabulary with a word soup or Blooket: {words_only(PRACTICE['April Week 1'])}."),
        ("Add today's word to your notebook: fair use.", "Use the simple definitions from the practice vocabulary list."),
        ("Match today's words with definitions: title, header, section, image.", practice_text("April Week 2")),
        ("Match today's words with definitions: link, menu, navigation, home page.", practice_text("April Week 3")),
        ("Review last class words with Blooket, Gimkit, or matching cards: link, menu, navigation, home page.", f"Review April Week 3 practice vocabulary with Blooket, Gimkit, or matching cards: {words_only(PRACTICE['April Week 3'])}."),
        ("Match today's words with definitions: variable, value, score, change.", practice_text("April Week 4")),
        ("Complete a vocabulary review activity using: variable, value, score, change.", f"Complete a vocabulary review activity using April Week 4 practice vocabulary: {words_only(PRACTICE['April Week 4'])}."),
    ],
    "6° Technology - May.docx": [
        ("Match today's words with definitions: safe, respectful, private, reliable.", practice_text("May Week 1")),
        ("Review last class words with Blooket or a crossword puzzle: safe, respectful, private, reliable.", f"Review May Week 1 practice vocabulary with Blooket or a crossword puzzle: {words_only(PRACTICE['May Week 1'])}."),
        ("Match today's words with definitions: recycle, reuse, repair, e-waste.", practice_text("May Week 2")),
        ("Complete a word soup or matching activity using last class words: recycle, reuse, repair, e-waste.", f"Complete a word soup or matching activity using May Week 2 practice vocabulary: {words_only(PRACTICE['May Week 2'])}."),
        ("Plan your audience, message, and three key words.", f"Plan your audience, message, and three key words from May Week 3 practice vocabulary: {words_only(PRACTICE['May Week 3'])}."),
        ("Review your topic words with a quick vocabulary check.", f"Review May Week 3 practice vocabulary with a quick check: {words_only(PRACTICE['May Week 3'])}."),
    ],
    "6° Technology - June.docx": [
        ("Match today's words with definitions: robot, program, instruction, safety.", f"Review Trimester 2 summative vocabulary: {words_only(SUMMATIVE['t2'])}."),
        ("Review last class words with Blooket, Gimkit, or matching cards: robot, program, instruction, safety.", f"Review Trimester 2 summative vocabulary with Blooket, Gimkit, or matching cards: {words_only(SUMMATIVE['t2'])}."),
        ("Create a vocabulary table with definition, illustration, and example sentence for these words: robot, mBot, program, instruction, safety.", f"Create a vocabulary table with definition, illustration, and example sentence for these words: {words_only(SUMMATIVE['t2'])}."),
        ("Match today's words with definitions: sequence, forward, backward, turn, stop.", practice_text("June Week 2")),
        ("Complete a quick review puzzle using last class words: sequence, forward, backward, turn, stop.", f"Complete a quick review puzzle using June Week 2 practice vocabulary: {words_only(PRACTICE['June Week 2'])}."),
        ("Match today's words with definitions: test, debug, improve.", practice_text("June Week 3")),
        ("Review the words test, debug, and improve with Blooket or a matching activity.", f"Review June Week 3 practice vocabulary with Blooket or a matching activity: {words_only(PRACTICE['June Week 3'])}."),
        ("Match today's words with definitions: output, LED, buzzer, signal.", practice_text("June Week 4")),
        ("Complete a vocabulary review activity using: output, LED, buzzer, signal.", f"Complete a vocabulary review activity using June Week 4 practice vocabulary: {words_only(PRACTICE['June Week 4'])}."),
    ],
    "6° Technology - July.docx": [
        ("Match today's words with definitions: speed, time, turn, sequence.", practice_text("July Week 1")),
        ("Review movement words with Blooket or a crossword puzzle: speed, time, turn, sequence.", f"Review July Week 1 practice vocabulary with Blooket or a crossword puzzle: {words_only(PRACTICE['July Week 1'])}."),
        ("Match today's words with definitions: repeat, loop, pattern.", practice_text("July Week 2")),
        ("Complete a vocabulary review activity using: repeat, loop, pattern.", f"Complete a vocabulary review activity using July Week 2 practice vocabulary: {words_only(PRACTICE['July Week 2'])}."),
        ("Match today's words with definitions: sensor, detect, obstacle, line.", practice_text("July Week 3")),
        ("Complete a word soup or matching activity using: sensor, detect, obstacle, line.", f"Complete a word soup or matching activity using July Week 3 practice vocabulary: {words_only(PRACTICE['July Week 3'])}."),
        ("Match today's words with definitions: if, condition, response.", practice_text("July Week 4")),
        ("Review condition words with Blooket or a crossword puzzle: if, condition, response.", f"Review July Week 4 practice vocabulary with Blooket or a crossword puzzle: {words_only(PRACTICE['July Week 4'])}."),
        ("Review key words from July with a quick matching activity.", f"Review July Week 5 practice vocabulary with a quick matching activity: {words_only(PRACTICE['July Week 5'])}."),
    ],
    "6° Technology - August.docx": [
        ("Match review words with definitions: movement, output, sensor, debug.", practice_text("August Week 1")),
        ("Review the selected cards for movement, LED/buzzer, sensor, and debugging.", f"Review August Week 1 practice vocabulary and the selected cards: {words_only(PRACTICE['August Week 1'])}."),
        ("Review challenge words: goal, success rule, test area, improvement.", practice_text("August Week 2")),
        ("Complete a vocabulary review activity using: challenge, test, debug, improve.", f"Complete a vocabulary review activity using August Week 3 practice vocabulary: {words_only(PRACTICE['August Week 3'])}."),
    ],
    "6° Technology - September.docx": [
        ("Match today's words with definitions: data, information, row, column, cell.", f"Review Trimester 3 summative vocabulary: {words_only(SUMMATIVE['t3'])}."),
        ("Review last class words with a word soup or crossword puzzle: data, information, row, column, cell.", f"Review Trimester 3 summative vocabulary with a word soup or crossword puzzle: {words_only(SUMMATIVE['t3'])}."),
        ("Create a vocabulary table with definition, illustration, and example sentence for these words: data, information, row, column, cell.", f"Create a vocabulary table with definition, illustration, and example sentence for these words: {words_only(SUMMATIVE['t3'])}."),
        ("Match today's words with definitions: formula, cell reference, sum, average.", practice_text("September Week 2")),
        ("Complete a vocabulary review activity using: formula, cell reference, sum, average.", f"Complete a vocabulary review activity using September Week 2 practice vocabulary: {words_only(PRACTICE['September Week 2'])}."),
    ],
    "6° Technology - October.docx": [
        ("Match today's words with definitions: budget, quantity, cost, total.", practice_text("October Week 1")),
        ("Review last class words with Blooket or a crossword puzzle: budget, quantity, cost, total.", f"Review October Week 1 practice vocabulary with Blooket or a crossword puzzle: {words_only(PRACTICE['October Week 1'])}."),
        ("Match today's words with definitions: chart, title, label, category.", practice_text("October Week 2")),
        ("Complete a vocabulary review activity using: chart, title, label, category.", f"Complete a vocabulary review activity using October Week 2 practice vocabulary: {words_only(PRACTICE['October Week 2'])}."),
        ("Match today's words with definitions: 3D, shape, move, resize, rotate.", practice_text("October Week 3")),
        ("Complete a word soup or matching activity using: 3D, shape, move, resize, rotate.", f"Complete a word soup or matching activity using October Week 3 practice vocabulary: {words_only(PRACTICE['October Week 3'])}."),
        ("Match today's words with definitions: design, prototype, duplicate, group, improve.", practice_text("October Week 4")),
        ("Review last class words with Blooket or a crossword puzzle: design, prototype, duplicate, group, improve.", f"Review October Week 4 practice vocabulary with Blooket or a crossword puzzle: {words_only(PRACTICE['October Week 4'])}."),
    ],
    "6° Technology - November.docx": [
        ("Match today's words with definitions: present, feedback, improve.", practice_text("November Week 1")),
        ("Review the words present, feedback, and improve with a quick matching activity.", f"Review November Week 1 practice vocabulary with a quick matching activity: {words_only(PRACTICE['November Week 1'])}."),
        ("Match today's words with definitions: micro:bit, input, output, LED, button.", practice_text("November Week 2")),
        ("Complete a word soup or crossword puzzle using last class words: micro:bit, input, output, LED, button.", f"Complete a word soup or crossword puzzle using November Week 2 practice vocabulary: {words_only(PRACTICE['November Week 2'])}."),
        ("Match today's words with definitions: variable, value, increase, decrease.", practice_text("November Week 3")),
        ("Review variable words with Blooket or matching cards: variable, value, increase, decrease.", f"Review November Week 3 practice vocabulary with Blooket or matching cards: {words_only(PRACTICE['November Week 3'])}."),
    ],
    "6° Technology - December.docx": [
        ("Match today's words with definitions: sensor, condition, if/then, light, temperature.", practice_text("December Week 1")),
        ("Complete a word soup or crossword puzzle using last class words: sensor, condition, if/then, light, temperature.", f"Complete a word soup or crossword puzzle using December Week 1 practice vocabulary: {words_only(PRACTICE['December Week 1'])}."),
        ("Review project words: input, output, variable, test, improve.", practice_text("December Week 2")),
        ("Review the words input, output, condition, variable, improve.", f"Review December Week 3 practice vocabulary: {words_only(PRACTICE['December Week 3'])}."),
    ],
}


def main():
    total = 0
    for name, replacements in REPLACEMENTS.items():
        path = DOCX_DIR / name
        if not path.exists():
            matches = sorted(DOCX_DIR.glob(f"*/{name}"))
            if matches:
                path = matches[0]
        doc = Document(path)
        count = replace_in_doc(doc, replacements)
        doc.save(path)
        total += count
        print(f"{name}: {count} replacements")
    print(f"total replacements: {total}")


if __name__ == "__main__":
    main()
