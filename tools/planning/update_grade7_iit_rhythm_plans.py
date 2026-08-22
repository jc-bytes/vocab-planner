#!/usr/bin/env python3
"""Update the Grade 7 IIT official Word plans from the revised Markdown plans."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
DRAFT = ROOT / "plans" / "7th Grade Technology" / "Planning" / "Drafts" / "2nd Trimester"
WORD = ROOT / "plans" / "7th Grade Technology" / "Planning" / "Monthly" / "2nd Trimester"

FILES = {
    "7° Technology - July.md": "7° Technology - July.docx",
    "7° Technology - August.md": "7° Technology - August.docx",
}

FONT = "Arial"
SIZE = Pt(12)


def section(text: str, heading: str, next_heading: str) -> str:
    match = re.search(
        rf"^## {re.escape(heading)}\n\n(.*?)\n\n## {re.escape(next_heading)}$",
        text,
        flags=re.M | re.S,
    )
    if not match:
        raise ValueError(f"Could not find section {heading}")
    return match.group(1).strip()


def list_items(block: str) -> list[str]:
    return [line[2:].strip() for line in block.splitlines() if line.startswith("- ")]


def monthly_rows(text: str) -> list[list[str]]:
    marker = "## Monthly Plan"
    block = text.split(marker, 1)[1]
    rows = []
    for line in block.splitlines():
        if not re.match(r"^\| Week \d+ \|", line):
            continue
        parts = [part.strip().replace("<br>", "\n") for part in line.strip().strip("|").split("|")]
        if len(parts) != 7:
            raise ValueError(f"Unexpected monthly-plan row: {line}")
        rows.append(parts)
    return rows


def set_run_font(run, *, bold: bool = False, highlight: bool = False) -> None:
    run.font.name = FONT
    run.font.size = SIZE
    run.bold = bold
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def add_labeled_line(cell, label: str, value: str, *, highlight: bool = False) -> None:
    paragraph = cell.paragraphs[0] if not cell.text and len(cell.paragraphs) == 1 else cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run(label), bold=True)
    set_run_font(paragraph.add_run(value), highlight=highlight)


def set_summary_cell(cell, label: str, values: list[str]) -> None:
    clear_cell(cell)
    add_labeled_line(cell, f"{label}: ", " ".join(values))


def set_lesson_cell(cell, row: list[str]) -> None:
    _week, _duration, topic, objective, pre, during, post = row
    clear_cell(cell)
    add_labeled_line(cell, "Topic: ", topic)
    add_labeled_line(cell, "Class Objective: ", objective)
    add_labeled_line(cell, "Pre-activities: ", pre)
    assessed = any(
        phrase in during
        for phrase in (
            "Complete Daily Summative #5",
            "Begin Exam Project:",
            "Continue Exam Project:",
            "Complete Exam Project Presentation:",
        )
    )
    add_labeled_line(cell, "While activities: ", during, highlight=assessed)
    add_labeled_line(cell, "Post-activities: ", post)


def update_document(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document(docx_path)

    competences = list_items(section(text, "Competences", "Learning Objectives"))
    objectives = list_items(section(text, "Learning Objectives", "Learning Outcomes"))
    outcomes = list_items(section(text, "Learning Outcomes", "Evaluation"))
    set_summary_cell(doc.tables[0].cell(1, 0), "Competences", competences)
    set_summary_cell(doc.tables[0].cell(2, 0), "Learning Objectives", objectives)
    set_summary_cell(doc.tables[0].cell(3, 0), "Learning Outcomes", outcomes)

    rows = monthly_rows(text)
    if len(rows) != 8:
        raise ValueError(f"Expected 8 lesson rows in {md_path.name}; found {len(rows)}")
    plan = doc.tables[1]
    for index, row in enumerate(rows):
        content_row = 1 + (index // 2) * 3
        column = index % 2
        set_lesson_cell(plan.cell(content_row, column), row)

    doc.save(docx_path)
    print(f"updated {docx_path}")


def main() -> None:
    for md_name, docx_name in FILES.items():
        update_document(DRAFT / md_name, WORD / docx_name)


if __name__ == "__main__":
    main()
