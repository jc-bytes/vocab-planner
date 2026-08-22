#!/usr/bin/env python3
"""Rebuild Grade 6 Trimester 3 rubrics from the official school rubric layout."""

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
FOLDER = ROOT / "plans/6th Grade Technology/Assessments/Rubrics/IIIT"
TEMPLATE = FOLDER / "6th grade - IIIT - Week 1 - Rubric for Summative 1.docx"
FONT = "Arial"

RUBRICS = [
    {
        "week": 2, "number": 2, "title": "School Data Table",
        "criteria": [
            ("Complete records", "Includes at least 8 complete records with values matching the teacher data.", "Includes 6-7 complete records; one or two values may need correction.", "Includes 3-5 complete records or several values need correction.", "Includes 0-2 complete records, or the table cannot be checked."),
            ("Headers and consistency", "All columns have correct headers; categories and number formats are consistent.", "Headers are present; one category or number-format inconsistency remains.", "One header is unclear or several entries use inconsistent formats.", "Headers are missing or most entries are inconsistent."),
            ("Readable formatting", "The table is readable with visible headers, suitable column widths, and no hidden data.", "The table is readable with one spacing, width, or emphasis issue.", "The table can be read, but several cells or labels are difficult to check.", "The table is unreadable, incomplete, or not submitted."),
            ("Data sentence", "One complete sentence states a specific fact supported by the table.", "The sentence matches the table but is missing one useful detail.", "The sentence is vague or only partly supported by the table.", "The sentence is missing or contradicts the table."),
        ],
    },
    {
        "week": 5, "number": 3, "title": "STEAM Preparation and Work Process",
        "criteria": [
            ("Preparation", "Arrives ready for both work weeks, knows the assigned project and role, and begins promptly.", "Is ready in both weeks with one minor reminder or missing item.", "Needs repeated reminders or is unprepared during one work session.", "Is unprepared for most observed work or does not begin the assigned task."),
            ("Participation and role", "Completes the assigned role in both weeks and contributes useful work at each checkpoint.", "Completes the role with one incomplete checkpoint or teacher reminder.", "Completes part of the role; several checkpoints need teacher support.", "Does not complete the assigned role or provides no usable contribution evidence."),
            ("Safe, responsible work", "Uses materials, files, class time, and group communication safely and responsibly in every observation.", "Shows safe, responsible work with one corrected reminder.", "Needs repeated reminders about materials, files, time, or group conduct.", "Uses materials or time unsafely, or prevents the group from completing assigned work."),
            ("Individual log evidence", "Records at least 3 dated entries naming the task, result, and next step or improvement.", "Records 2 complete entries and one partial entry.", "Records 1 complete entry or several unclear fragments.", "Provides no checkable individual log evidence."),
        ],
    },
    {
        "week": 6, "number": 4, "title": "STEAM Expo Participation and Closure",
        "criteria": [
            ("Assigned expo role", "Completes the assigned speaking or demonstration role at the expo.", "Completes the role with one teacher prompt or small missing part.", "Completes only part of the role or needs several prompts.", "Does not complete the assigned role and provides no approved make-up evidence."),
            ("Project explanation", "Clearly explains what the group created, the problem, how it works, what was learned, and how it can help.", "Explains 4 of the 5 required points clearly.", "Explains 2-3 required points or several points are unclear.", "Explains 0-1 required point or gives no checkable explanation."),
            ("Respect and cleanup", "Listens respectfully, follows expo directions, and completes the assigned cleanup task.", "Meets all three expectations with one corrected reminder.", "Meets one or two expectations and needs repeated reminders.", "Does not meet the participation or cleanup expectations."),
            ("Individual reflection", "Writes one specific result, one challenge, and one improvement supported by the project experience.", "Includes all three reflection parts; one needs more detail.", "Includes one or two clear reflection parts.", "Reflection is missing or does not describe the student's project experience."),
        ],
    },
    {
        "week": 7, "number": 5, "title": "School Data Chart and Conclusion",
        "criteria": [
            ("Data and chart type", "Uses the assigned data range and creates one accurate column chart.", "Creates the column chart with one minor range or data-selection error.", "Creates a chart, but several values are missing or the chart type is wrong.", "Chart is missing or does not use the assigned data."),
            ("Title and axis labels", "Includes the assigned chart title and correct labels for both axes.", "Includes the title and both labels; one has a minor wording issue.", "Includes only one or two of the three required text elements.", "Title and axis labels are missing or do not match the data."),
            ("Three answers", "Answers all 3 questions correctly using values or categories from the chart.", "Answers all 3; one answer needs a clearer chart reference.", "Answers 1-2 questions correctly.", "Answers are missing or are not supported by the chart."),
            ("Conclusion", "Writes one complete conclusion that states a specific pattern or comparison supported by the chart.", "The conclusion matches the chart but needs one specific detail.", "The conclusion is vague or only partly supported.", "The conclusion is missing or contradicts the chart."),
        ],
    },
]


def font_all(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT
                        run.font.size = Pt(12)
                        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                            fonts.set(qn(f"w:{key}"), FONT)


def set_cell(cell, text):
    cell.text = text
    font_all_cell(cell)


def font_all_cell(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT
            run.font.size = Pt(12)
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{key}"), FONT)


def build(spec):
    document = Document(TEMPLATE)
    table = document.tables[0]
    header = table.cell(0, 0)
    header.text = "\n".join([
        "Academia Internacional David", "Robotics and Technology", "3rd Trimester",
        f"Summative #{spec['number']}", "6th A & B & C",
        "Name: ________________________________   Date: __________________   Group: 6° A B C",
        "Teacher: Porfirio Rios                  Score: _____ / 40pts", spec["title"],
    ])
    for index, criterion in enumerate(spec["criteria"], start=2):
        for cell, text in zip(table.rows[index].cells[:5], criterion):
            set_cell(cell, text)
    readiness = table.rows[7].cells
    readiness_text = (
        "Punctuality, Readiness & Respect",
        "Submits the required evidence on time, arrives ready with the assigned computer and login, and works respectfully.",
        "Submits on time with one minor readiness or respect reminder.",
        "Submits the evidence, but lateness or repeated readiness/respect reminders affect the class.",
        "Does not submit on time, is not ready to work, or does not follow respectful-work expectations.",
    )
    for cell, text in zip(readiness[:5], readiness_text):
        set_cell(cell, text)
    font_all(document)
    name = f"6th grade - IIIT - Week {spec['week']} - Rubric for Summative {spec['number']}.docx"
    document.save(FOLDER / name)
    print(name)


for rubric in RUBRICS:
    build(rubric)
